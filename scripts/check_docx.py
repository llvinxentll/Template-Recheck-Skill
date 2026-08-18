#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
check_docx.py — TULIBS thesis .docx format checker (rev.2024 / rev.2023 Times)

Checks a student's Word file against the Thammasat University Library (TULIBS)
thesis template. Works directly on the .docx internals (styles.xml, sectPr,
runs) so it verifies things a PDF cannot: real font names, half-point sizes,
paragraph styles, indentation in twips, line spacing settings.

Usage:
    python3 check_docx.py student.docx                 # auto-detect profile
    python3 check_docx.py student.docx --profile thai
    python3 check_docx.py student.docx --profile english
    python3 check_docx.py student.docx --profile english-times
    python3 check_docx.py student.docx --json report.json

Profiles:
    thai            TULIBS_Thesis-template-Thai_rev_2024  (TH Sarabun New 16pt)
    english         TULIBS-Thesis-template-English_rev_2024 (TH Sarabun New 16pt)
    english-times   TULIBS-Thesis-template-English-Times_rev_2023
                    (Times New Roman 12pt Latin + TH Sarabun New 16pt Thai)

Exit code 0 always (this is a reporting tool, not a gate).
Requires: python-docx  (pip install python-docx --break-system-packages)
"""
import sys, os, json, argparse, re, time, shutil, subprocess, tempfile, unicodedata

# ให้ import โมดูลข้าง ๆ (pagemap_service) ได้เสมอ แม้ถูก import จาก cwd อื่น
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

EMU_PER_INCH = 914400
TWIPS_PER_INCH = 1440

# ---- allowed fonts per profile -------------------------------------------
SARABUN = {"TH Sarabun New", "TH SarabunPSK", "THSarabunNew", "TH Sarabun PSK"}
# symbol fonts may be legitimate for bullets; mathematical fonts are exempt
# only inside Office Math (a normal prose run set to Cambria Math is still wrong).
SYMBOL_OK = {"Symbol", "Wingdings", "Marlett"}
MATH_ONLY = {"Cambria Math", "MT Extra"}

PROFILES = {
    "thai": {
        "label": "TULIBS Thai rev.2024",
        "latin_font": SARABUN, "thai_font": SARABUN,
        "body_pt": 16.0, "latin_body_pt": 16.0,
        "main_heading_left_in": 0.0,   # flush left
        "line_spacing": 1.0,           # single
    },
    "english": {
        "label": "TULIBS English rev.2024",
        "latin_font": SARABUN, "thai_font": SARABUN,
        "body_pt": 16.0, "latin_body_pt": 16.0,
        "main_heading_left_in": 0.25,  # 0.25" left indent
        "line_spacing": 1.0,           # single
    },
    "english-times": {
        "label": "TULIBS English Times rev.2023",
        "latin_font": {"Times New Roman"}, "thai_font": SARABUN,
        "body_pt": 16.0, "latin_body_pt": 12.0,   # Thai 16pt CS, Latin 12pt
        "main_heading_left_in": 0.25,
        "line_spacing": 1.5,           # this older variant uses 1.5 lines
    },
}

# Expected named styles that must exist and stay unchanged (template integrity)
# key: style name -> dict of expected attrs (checked leniently)
TU_STYLES = [
    "TU_Chapter", "TU_Paragraph_Normal",
    "TU_Main Heading _Chapter1", "TU_Main Heading_Chapter2",
    "TU_Sub-heading 1", "TU_Sub-heading 2", "TU_Sub-heading 3",
    "TU_Para_Sub-heading 1", "TU_Para_Sub-heading 2", "TU_Para_Sub-heading 3",
]

TH_RANGE = (0x0E00, 0x0E7F)

def is_thai(ch):
    return TH_RANGE[0] <= ord(ch) <= TH_RANGE[1]

def script_of(text):
    """Return 'thai' if the run holds any Thai char, else 'latin'."""
    for ch in text:
        if is_thai(ch):
            return "thai"
    return "latin"


def scripts_in_text(text):
    """คืนชุดสคริปต์ที่มีตัวอักษรจริงใน run เดียวกัน

    Word เก็บข้อความไทยและละตินไว้ใน run เดียวกันได้ แต่ใช้ฟอนต์คนละช่อง
    (w:cs/eastAsia กับ w:ascii/hAnsi) จึงต้องตรวจทั้งสองช่อง ไม่ใช่เลือกภาษาเดียว
    จากการพบอักษรไทยเพียงตัวเดียว.
    """
    out = set()
    if any(is_thai(ch) for ch in text):
        out.add("thai")
    if any(("A" <= ch <= "Z") or ("a" <= ch <= "z") for ch in text):
        out.add("latin")
    return out or {script_of(text)}


def _on_off(el):
    if el is None:
        return None
    val = el.get(qn("w:val"))
    return val not in ("0", "false", "off")


def _rpr_values(rpr, script):
    """อ่าน font/size/bold ที่ระบุใน rPr ชั้นเดียว (ยังไม่ resolve inheritance)."""
    if rpr is None:
        return {"font": None, "size": None, "bold": None, "color": None}
    font = None
    rf = rpr.find(qn("w:rFonts"))
    if rf is not None:
        keys = (("w:cs", "w:eastAsia", "w:ascii", "w:hAnsi")
                if script == "thai" else
                ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"))
        font = next((rf.get(qn(k)) for k in keys if rf.get(qn(k))), None)
    size = None
    size_keys = ("w:szCs", "w:sz") if script == "thai" else ("w:sz", "w:szCs")
    for key in size_keys:
        el = rpr.find(qn(key))
        if el is not None:
            try:
                size = int(el.get(qn("w:val"))) / 2
            except (TypeError, ValueError):
                size = None
            break
    bold_keys = ("w:bCs", "w:b") if script == "thai" else ("w:b", "w:bCs")
    bold = None
    for key in bold_keys:
        el = rpr.find(qn(key))
        if el is not None:
            bold = _on_off(el)
            break
    color = None
    color_el = rpr.find(qn("w:color"))
    if color_el is not None:
        raw = color_el.get(qn("w:val"))
        theme = color_el.get(qn("w:themeColor"))
        color = (f"theme:{theme}" if theme else raw)
        if color in (None, "auto", "000000"):
            color = None
    return {"font": font, "size": size, "bold": bold, "color": color}


def _style_chain(style):
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        yield style
        style = style.base_style


def _resolved_style_values(style, script, cache):
    if style is None:
        return {"font": None, "size": None, "bold": None, "color": None}
    key = ("style", style.style_id, script)
    if cache is not None and key in cache:
        return cache[key]
    resolved = {"font": None, "size": None, "bold": None, "color": None}
    for item in _style_chain(style):
        vals = _rpr_values(item._element.find(qn("w:rPr")), script)
        for name in resolved:
            if resolved[name] is None and vals[name] is not None:
                resolved[name] = vals[name]
    if cache is not None:
        cache[key] = resolved
    return resolved


def effective_run_format(run, paragraph, doc, script, cache=None, paragraph_style=None):
    """Resolve direct → character style → paragraph style → docDefaults.

    python-docx คืน ``None`` เมื่อค่ามาจาก style; การเอา ``None`` ไปแปลว่า
    "ไม่หนา/ไม่มีฟอนต์" ทำให้เกิด false positive จำนวนมาก โดยเฉพาะ Keywords.
    """
    rpr = run._element.find(qn("w:rPr"))
    sources = [_rpr_values(rpr, script)]
    # อย่าเรียก run.style ทุก run: ถ้าไม่มี w:rStyle ย่อมไม่มี character-style override.
    if rpr is not None and rpr.find(qn("w:rStyle")) is not None:
        try:
            sources.append(_resolved_style_values(run.style, script, cache))
        except (AttributeError, KeyError):
            pass
    try:
        style = paragraph_style if paragraph_style is not None else paragraph.style
        sources.append(_resolved_style_values(style, script, cache))
    except (AttributeError, KeyError):
        pass
    default_key = ("defaults", script)
    if cache is not None and default_key in cache:
        sources.append(cache[default_key])
    else:
        defaults = doc.styles.element.find(qn("w:docDefaults"))
        default_rpr = (defaults.find(qn("w:rPrDefault") + "/" + qn("w:rPr"))
                       if defaults is not None else None)
        default_values = _rpr_values(default_rpr, script)
        if cache is not None:
            cache[default_key] = default_values
        sources.append(default_values)

    resolved = {"font": None, "size": None, "bold": None, "color": None}
    for vals in sources:
        for key in resolved:
            if resolved[key] is None and vals[key] is not None:
                resolved[key] = vals[key]
    return resolved


def run_is_bold(run, paragraph, doc):
    """True เมื่อ run แสดงผลหนาจาก direct formatting หรือ style inheritance."""
    scripts = scripts_in_text(run.text or "")
    return all(effective_run_format(run, paragraph, doc, s)["bold"] is True
               for s in scripts)


def iter_document_paragraphs(doc):
    """Yield (paragraph, logical location) ตามลำดับจริง รวม table cells."""
    seen_cells = set()
    body_paragraphs = doc.paragraphs  # property นี้สร้าง list ใหม่ทุกครั้ง — cache กัน O(n²)

    def count_breaks(el):
        total = 0
        for node in el.iter():
            if node.tag == qn("w:lastRenderedPageBreak"):
                total += 1
            elif node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
                total += 1
        return total

    def walk_table(table, table_path, start_page):
        local_page = start_page
        for ri, row in enumerate(table.rows, 1):
            for ci, cell in enumerate(row.cells, 1):
                cell_key = cell._tc
                if cell_key in seen_cells:       # merged cell อาจถูกคืนซ้ำ
                    continue
                seen_cells.add(cell_key)
                for pi, paragraph in enumerate(cell.paragraphs, 1):
                    yield paragraph, (f"{table_path} แถว {ri} "
                                      f"คอลัมน์ {ci} ย่อหน้า {pi}")
                    local_page += count_breaks(paragraph._p)
                for ni, nested in enumerate(cell.tables, 1):
                    yield from walk_table(nested, f"{table_path}.{ni}", local_page)

    page, body_i, table_i = 1, 0, 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield body_paragraphs[body_i], f"para {body_i}"
            body_i += 1
            page += count_breaks(child)
        elif child.tag == qn("w:tbl"):
            table_i += 1
            yield from walk_table(Table(child, doc._body), f"ตาราง {table_i}", page)
            page += count_breaks(child)


def run_is_math(run):
    node = run._element
    math_tags = {qn("m:oMath"), qn("m:oMathPara")}
    while node is not None:
        if node.tag in math_tags:
            return True
        node = node.getparent()
    return False

# ---------------------------------------------------------------------------
class Finding:
    def __init__(self, sev, cat, msg, detail="", loc="", correct=""):
        self.sev, self.cat, self.msg, self.detail = sev, cat, msg, detail
        self.loc = loc            # clear location: section / style / paragraph
        self.correct = correct    # the concrete correct value ("ค่าที่ถูกต้อง")
    def row(self):
        domain = "spelling" if self.cat == "Spelling" else "template"
        return {"domain": domain, "severity": self.sev,
                "category": self.cat, "location": self.loc,
                "issue": self.msg, "detail": self.detail, "correct": self.correct}

SEV = {"critical": "🔴", "major": "🟠", "minor": "🟡", "ok": "🟢", "info": "ℹ️"}

# ---------------------------------------------------------------------------
def _rpr_font_size(rpr):
    """ดึง (ascii font, ขนาด pt) จาก <w:rPr> ใด ๆ — คืน (None, None) ถ้าไม่ได้ระบุ"""
    if rpr is None:
        return None, None
    font = sz = None
    rf = rpr.find(qn("w:rFonts"))
    if rf is not None:
        font = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi")) or rf.get(qn("w:cs"))
    s = rpr.find(qn("w:sz"))
    if s is not None:
        try:
            sz = int(s.get(qn("w:val"))) / 2
        except (TypeError, ValueError):
            sz = None
    return font, sz


def detect_profile(doc):
    """เดาโปรไฟล์เทมเพลตจากลายเซ็นฟอนต์ในไฟล์

    ตัวชี้ขาดคือ "ละติน Times New Roman 12pt" ซึ่งมีเฉพาะเทมเพลต rev.2023.
    ต้องดู 3 ชั้นเพราะเล่มนักศึกษาจริงมักไม่ได้เริ่มจากไฟล์เทมเพลต — บางเล่ม
    docDefaults ยังเป็นค่า Word มาตรฐาน (Calibri 11) แต่ตัวเนื้อหาเป็น Times 12
    ทั้งเล่ม. ถ้าดูแค่ docDefaults จะตัดสินเป็น rev.2024 แล้ว flag ฟอนต์ผิด
    ทั้งเล่มทั้งที่นักศึกษาทำถูกตามเทมเพลต 2023 — false positive ที่แพงมาก.
    """
    el = doc.styles.element
    # ชั้นที่ 1: docDefaults (ไฟล์ที่ทำจากเทมเพลตจริงจะตรงตั้งแต่ชั้นนี้)
    dd = el.find(qn("w:docDefaults"))
    font = sz = None
    if dd is not None:
        font, sz = _rpr_font_size(dd.find(qn("w:rPrDefault") + "/" + qn("w:rPr")))
    if font and "Times" in font and sz == 12.0:
        return "english-times"

    # ชั้นที่ 2: สไตล์ Normal — เล่มที่ตั้งฟอนต์ผ่าน "Set as Default" ของ Word
    for st in el.findall(qn("w:style")):
        if st.get(qn("w:styleId")) in ("Normal", "TU_Paragraph_Normal"):
            f2, s2 = _rpr_font_size(st.find(qn("w:rPr")))
            if f2 and "Times" in f2 and (s2 or sz) == 12.0:
                return "english-times"

    # ชั้นที่ 3: ฟอนต์ที่ใช้จริงมากที่สุดใน run — เล่มที่ format ด้วยมือทั้งเล่ม
    tally = {}
    for p in doc.paragraphs:
        for r in p.runs:
            f3, s3 = _rpr_font_size(r._element.find(qn("w:rPr")))
            if f3 and s3:
                tally[(f3, s3)] = tally.get((f3, s3), 0) + len(r.text or "")
    if tally:
        (top_font, top_sz), _ = max(tally.items(), key=lambda kv: kv[1])
        if "Times" in top_font and top_sz == 12.0:
            return "english-times"

    # rev.2024: แยกไทย/อังกฤษจากภาษาของข้อความหน้าปก
    joined = " ".join(p.text for p in doc.paragraphs[:12])
    if any(is_thai(c) for c in joined):
        return "thai"
    return "english"

def inch(emu):
    return None if emu is None else round(emu / EMU_PER_INCH, 3)

def approx(a, b, tol=0.03):
    return a is not None and b is not None and abs(a - b) <= tol

# ---------------------------------------------------------------------------
# Page numbers from Word's last saved pagination (fallback).
# Word records where each page broke during its last save as
# <w:lastRenderedPageBreak/> (plus any explicit <w:br w:type="page"/>).
# Counting those in document order gives a fallback page estimate. Files never
# opened/saved by Word may lack these markers; a later edit can also make them
# stale. Exact physical sheets therefore come from resolve_page_data() rendering
# a fresh PDF; marker-derived locations retain a "~" warning.
def build_page_index(doc):
    """Return (page_by_para_index, has_markers).
    page_by_para_index[i] = 1-based page where doc.paragraphs[i] starts.

    We walk the body's direct children in document order. doc.paragraphs is
    exactly the sequence of body-level <w:p> in that same order, so counting
    them as we go maps each paragraph to a page without relying on element
    id()/identity (lxml hands out fresh proxy objects on each access, so id()
    keys are unreliable). Page breaks inside tables and inside a paragraph's
    runs are both counted, so pagination stays aligned with what Word shows.
    """
    P    = qn("w:p")
    TBL  = qn("w:tbl")
    LRP  = qn("w:lastRenderedPageBreak")
    BR   = qn("w:br")
    TYPE = qn("w:type")

    def count_breaks(el):
        n = 0
        for e in el.iter():
            if e.tag == LRP:
                n += 1
            elif e.tag == BR and e.get(TYPE) == "page":
                n += 1
        return n

    page = 1
    markers = 0
    page_by_index = []
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == P:
            page_by_index.append(page)          # paragraph starts on current page
            b = count_breaks(child); page += b; markers += b
        elif tag == TBL:
            b = count_breaks(child); page += b; markers += b
        # sectPr and other body-level elements don't hold paragraphs — skip
    return page_by_index, (markers > 0)

def _normalize_page_text(text):
    """Normalize DOCX/PDF text for page matching without losing Thai letters."""
    text = unicodedata.normalize("NFC", text or "")
    text = text.replace("\u200b", "").replace("\ufeff", "")
    return re.sub(r"\s+", "", text).casefold()


def _document_text_records(doc):
    """Return document-order records for body and table-cell paragraphs."""
    records, seen_cells = [], set()
    body_paragraphs = doc.paragraphs

    def add_table(table, table_path):
        for ri, row in enumerate(table.rows, 1):
            for ci, cell in enumerate(row.cells, 1):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for pi, paragraph in enumerate(cell.paragraphs, 1):
                    key = f"{table_path} แถว {ri} คอลัมน์ {ci} ย่อหน้า {pi}"
                    style = paragraph.style.name if paragraph.style else ""
                    records.append((key, paragraph.text or "", style))
                for ni, nested in enumerate(cell.tables, 1):
                    add_table(nested, f"{table_path}.{ni}")

    body_i = table_i = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = body_paragraphs[body_i]
            style = paragraph.style.name if paragraph.style else ""
            records.append((f"para {body_i}", paragraph.text or "", style))
            body_i += 1
        elif child.tag == qn("w:tbl"):
            table_i += 1
            add_table(Table(child, doc._body), f"ตาราง {table_i}")
    return records


# ระยะมองไปข้างหน้าตอนจับคู่ย่อหน้ากับแผ่นใน PDF
_PAGE_MATCH_WINDOW = 3      # ปกติดูแค่แผ่นปัจจุบัน + 3
_PAGE_MATCH_RESYNC = 40     # พลาดติดกันเกินนี้ = หลุดจริง ค่อยขยายระยะค้น
_PAGE_MATCH_WIDE = 25       # ระยะค้นตอนพยายามกลับเข้าที่
_PAGE_MATCH_TRUST = 12      # anchor สั้นกว่านี้ไม่พอจะเชื่อว่าไม่ซ้ำที่อื่น


def _match_records_to_pdf_pages(records, pdf_pages):
    """Map records to physical PDF sheets; never move backwards in the document."""
    normalized_pages = [_normalize_page_text(page) for page in pdf_pages]
    assigned = [None] * len(records)
    exact = [False] * len(records)
    cursor = cursor_offset = 0
    matched = eligible = 0
    misses = 0
    for index, (_, text, _) in enumerate(records):
        compact = _normalize_page_text(text)
        if len(compact) < 4:
            continue
        eligible += 1
        # The start of a paragraph determines its physical sheet. Shorter
        # fallbacks tolerate line wrapping and PDF text extraction quirks.
        sizes = [min(len(compact), 100), min(len(compact), 60), min(len(compact), 30)]
        anchors = []
        for size in sizes:
            if size >= 4 and compact[:size] not in anchors:
                anchors.append(compact[:size])
        # ค้นหาแบบ "มองไปข้างหน้าไม่กี่แผ่น" เท่านั้น — ห้ามค้นถึงท้ายเล่ม
        # เพราะข้อความที่ซ้ำได้ทั้งเล่ม (หัวข้อที่โผล่ในสารบัญ, ข้อคำถามที่ซ้ำใน
        # ภาคผนวก) จะทำให้เคอร์เซอร์กระโดดไปท้ายเล่มแล้วไม่มีวันกลับ — หลังจากนั้น
        # ทุก record จะจับคู่ไม่ได้และเลขแผ่นทั้งรายงานเพี้ยนยกชุด
        span = _PAGE_MATCH_WINDOW if misses < _PAGE_MATCH_RESYNC else _PAGE_MATCH_WIDE
        limit = min(cursor + span + 1, len(normalized_pages))
        found = found_offset = anchor_length = None
        for page_index in range(cursor, limit):
            search_from = cursor_offset if page_index == cursor else 0
            for anchor in anchors:
                offset = normalized_pages[page_index].find(anchor, search_from)
                if offset >= 0:
                    found = page_index + 1
                    found_offset = offset
                    anchor_length = len(anchor)
                    break
            if found is not None:
                break
        if found is None:
            # จับคู่ไม่ได้ = ปล่อยว่างไว้ให้ขั้นเติมช่องว่างจัดการ **ห้ามขยับเคอร์เซอร์**
            misses += 1
            continue
        # ขยับเคอร์เซอร์เฉพาะเมื่อ anchor ยาวพอจะเป็นเอกลักษณ์ หรืออยู่บนแผ่นเดิม
        if anchor_length >= _PAGE_MATCH_TRUST or found == cursor + 1:
            cursor = found - 1
            cursor_offset = found_offset + anchor_length
        assigned[index] = found
        exact[index] = True
        matched += 1
        misses = 0

    # A paragraph bracketed by two confirmed records on the same sheet is
    # also confirmed. Other gaps inherit the preceding sheet only as fallback.
    previous = None
    next_exact = [None] * len(records)
    upcoming = None
    for i in range(len(records) - 1, -1, -1):
        if exact[i]:
            upcoming = assigned[i]
        next_exact[i] = upcoming
    for i in range(len(records)):
        if exact[i]:
            previous = assigned[i]
            continue
        if previous is not None and next_exact[i] == previous:
            assigned[i] = previous
            exact[i] = True
        elif previous is not None:
            assigned[i] = previous
    ratio = (matched / eligible) if eligible else 1.0
    return assigned, exact, ratio


def _render_pdf_pages(path):
    """Render DOCX and return extracted text per physical sheet.

    ลำดับการทำงาน (ดูรายละเอียดใน scripts/pagemap_service.py):
      1. มี pagemap worker เปิดอยู่นอก sandbox → ส่งงานผ่านคิวไฟล์ (ไม่ใช้เน็ต)
      2. ไม่มี worker → เรียก soffice ตรง ๆ ในโปรเซสนี้ (พฤติกรรมเดิม)
      3. เรียกตรงแล้วถูก sandbox ฆ่า (SIGABRT บน macOS/Codex) → ยก RuntimeError
         พร้อมวิธีเปิด worker แทนที่จะรายงานว่า "ไฟล์เสีย"
    """
    from pagemap_service import (START_HINT, SandboxBlockedError,
                                 convert_docx_to_pages, submit_job, worker_alive)
    if worker_alive():
        return submit_job(path)
    try:
        return convert_docx_to_pages(path)
    except SandboxBlockedError as exc:
        raise RuntimeError(f"{exc}\n{START_HINT}") from exc


def _section_ranges(doc, body_pages, total_pages):
    ranges, start, body_i = [], 0, -1

    def append_range(end):
        values = [p for p in body_pages[start:end] if p]
        if values:
            ranges.append((min(values), max(values)))
        elif total_pages:
            ranges.append((1, total_pages))
        else:
            ranges.append((None, None))

    for child in doc.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        body_i += 1
        ppr = child.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            append_range(body_i + 1)
            start = body_i + 1
    append_range(len(body_pages))
    return ranges


NO_MARKER_HINT = (
    "ไฟล์นี้ไม่มีมาร์กเกอร์แบ่งหน้าของ Word จึงระบุแผ่นงานไม่ได้ — "
    "ให้เปิดไฟล์ใน Microsoft Word แล้วกดบันทึก (Save) ทับหนึ่งครั้ง Word จะเขียนตำแหน่งแบ่งหน้าลงไฟล์ให้ "
    "จากนั้นรันตรวจใหม่ (ไม่ต้องติดตั้งหรือเปิดโปรแกรมอื่นเพิ่ม)")


def resolve_page_data(path, doc=None, engine="markers"):
    """Return physical-sheet mapping.

    engine="markers" (ค่าเริ่มต้น) — อ่าน `w:lastRenderedPageBreak` ในไฟล์ ไม่พึ่ง
        โปรแกรมอื่น เร็วและไม่มีทางล้ม. เลขที่ได้คลาดสะสมได้ (วัดกับเล่มจริง 135 หน้า:
        ท้ายเล่มคลาดถึง +8 แผ่น) **แต่ไม่กระทบผลลัพธ์** เพราะรายงานไม่แสดงเลขแผ่นแล้ว
        ใช้แค่จัดลำดับแถว ซึ่งความคลาดแบบสะสมไม่ทำให้ลำดับสลับ
    engine="auto" / "rendered" — จัดหน้าใหม่เป็น PDF ด้วย LibreOffice ได้เลขแผ่นตรงจริง
        เก็บไว้เผื่อกรณีที่อยากได้เลขแผ่นแม่น ๆ ไปใช้นอกรายงาน; ช้ากว่าและบน macOS
        ที่รันใน sandbox ของ Codex จะถูกฆ่า (สคริปต์ดักและถอยไป markers ให้เอง)
    """
    if engine not in ("auto", "rendered", "markers"):
        raise ValueError("page engine must be auto, rendered, or markers")
    doc = doc or Document(path)
    marker_pages, has_markers = build_page_index(doc)
    warning = ""
    if engine in ("auto", "rendered"):
        try:
            pdf_pages = _render_pdf_pages(path)
            records = _document_text_records(doc)
            assigned, exact, ratio = _match_records_to_pdf_pages(records, pdf_pages)
            location_pages = {record[0]: assigned[i] for i, record in enumerate(records)
                              if assigned[i] is not None}
            location_exact = {record[0]: exact[i] for i, record in enumerate(records)}
            style_pages = {}
            for i, (_, _, style) in enumerate(records):
                if style and assigned[i]:
                    style_pages.setdefault(style, set()).add(assigned[i])
            body_pages, body_exact = [], []
            for i in range(len(doc.paragraphs)):
                key = f"para {i}"
                page = location_pages.get(key)
                is_exact = location_exact.get(key, False)
                if page is None and i < len(marker_pages):
                    page = marker_pages[i]
                body_pages.append(page)
                body_exact.append(is_exact)
            total = len(pdf_pages)
            return {
                "pages": body_pages, "body_exact": body_exact,
                "has_pages": bool(total), "source": "rendered_pdf",
                "total_pages": total, "match_ratio": round(ratio, 4),
                "location_pages": location_pages,
                "location_exact": location_exact,
                "style_pages": {k: sorted(v) for k, v in style_pages.items()},
                "section_ranges": _section_ranges(doc, body_pages, total),
                "warning": "",
            }
        except Exception as exc:
            if engine == "rendered":
                raise
            warning = str(exc)
    total = max(marker_pages) if has_markers and marker_pages else 0
    if not has_markers:
        warning = (warning + " · " + NO_MARKER_HINT) if warning else NO_MARKER_HINT
    return {
        "pages": marker_pages, "body_exact": [False] * len(marker_pages),
        "has_pages": has_markers, "source": "word_markers" if has_markers else "unavailable",
        "total_pages": total, "match_ratio": 0.0, "location_pages": {},
        "location_exact": {}, "style_pages": {},
        "section_ranges": _section_ranges(doc, marker_pages, total),
        "warning": warning,
    }


def unpack_page_data(page_data, doc=None):
    if isinstance(page_data, dict):
        return page_data.get("pages", []), bool(page_data.get("has_pages"))
    if isinstance(page_data, tuple) and len(page_data) == 2:
        return page_data
    if isinstance(page_data, list):
        return page_data, bool(page_data)
    if page_data is not None:
        raise TypeError("unsupported page_data value")
    return build_page_index(doc)


_PARA_LOC_RE = re.compile(r"(?:\bpara|¶)\s*(\d+)", re.I)
_TABLE_LOC_RE = re.compile(
    r"ตาราง\s+\d+(?:\.\d+)*\s+แถว\s+\d+\s+คอลัมน์\s+\d+\s+ย่อหน้า\s+\d+")
_OLD_PAGE_PREFIX_RE = re.compile(
    r"(?:หน้า|แผ่น(?:งาน)?ที่)\s*~?\d+(?:\s*[–-]\s*\d+)?(?:\s*,\s*\d+)*\s*·\s*")


def _compact_ranges(pages):
    values = sorted(set(int(p) for p in pages if p))
    if not values:
        return ""
    groups, start = [], values[0]
    previous = start
    for value in values[1:]:
        if value == previous + 1:
            previous = value
            continue
        groups.append(str(start) if start == previous else f"{start}–{previous}")
        start = previous = value
    groups.append(str(start) if start == previous else f"{start}–{previous}")
    return ", ".join(groups)


def enrich_location_with_pages(location, page_data):
    """Prefix a logical location with physical sheet(s), preserving uncertainty."""
    if not location:
        return location
    if not isinstance(page_data, dict):
        pages, has_pages = unpack_page_data(page_data)
        page_data = {"pages": pages, "body_exact": [False] * len(pages),
                     "has_pages": has_pages, "total_pages": max(pages or [0]),
                     "location_pages": {}, "location_exact": {},
                     "style_pages": {}, "section_ranges": []}
    clean = _OLD_PAGE_PREFIX_RE.sub("", str(location)).strip()
    found_pages, exact_flags = [], []
    location_pages = page_data.get("location_pages", {})
    location_exact = page_data.get("location_exact", {})
    body_pages = page_data.get("pages", [])
    body_exact = page_data.get("body_exact", [])

    for match in _TABLE_LOC_RE.finditer(clean):
        key = match.group(0)
        if location_pages.get(key):
            found_pages.append(location_pages[key])
            exact_flags.append(location_exact.get(key, False))
    for match in _PARA_LOC_RE.finditer(clean):
        index = int(match.group(1))
        if 0 <= index < len(body_pages) and body_pages[index]:
            found_pages.append(body_pages[index])
            exact_flags.append(body_exact[index] if index < len(body_exact) else False)

    section = re.search(r"\bSection\s*(\d+)", clean, re.I)
    if section and not found_pages:
        index = int(section.group(1))
        ranges = page_data.get("section_ranges", [])
        if 0 <= index < len(ranges):
            start, end = ranges[index]
            if start:
                found_pages.extend(range(start, (end or start) + 1))
                exact_flags.append(page_data.get("source") == "rendered_pdf")

    style = re.match(r"style\s+(.+)$", clean, re.I)
    if style and not found_pages:
        found_pages.extend(page_data.get("style_pages", {}).get(style.group(1).strip(), []))
        if found_pages:
            exact_flags.append(page_data.get("source") == "rendered_pdf")

    total = int(page_data.get("total_pages") or 0)
    if not found_pages and total:
        if (clean.lower() in ("ทั้งเอกสาร", "header", "footer", "styles.xml") or
                style is not None):
            found_pages.extend(range(1, total + 1))
            exact_flags.append(page_data.get("source") == "rendered_pdf")
        elif clean.startswith("หน้าปก") or clean.startswith("cover"):
            found_pages.append(1)
            exact_flags.append(page_data.get("source") == "rendered_pdf")
    if not found_pages:
        return clean
    rendered_exact = bool(exact_flags) and all(exact_flags)
    approximation = "" if rendered_exact else "~"
    return f"แผ่นที่ {approximation}{_compact_ranges(found_pages)} · {clean}"


def enrich_locations_with_pages(findings, page_data):
    for finding in findings:
        finding.loc = enrich_location_with_pages(finding.loc, page_data)


def enrich_dict_locations_with_pages(findings, page_data):
    for finding in findings:
        finding["location"] = enrich_location_with_pages(finding.get("location", ""), page_data)

# Canonical section headings we surface a page number for, so a manual/APA
# pass can quote real pages (e.g. where the reference list begins).
_SECTION_MARKERS = [
    r"^บทคัดย่อ", r"^ABSTRACT", r"^สารบัญ$", r"^สารบัญตาราง", r"^สารบัญภาพ",
    r"^กิตติกรรมประกาศ", r"^บทที่\s*\d", r"^CHAPTER\s*\d",
    r"^รายการอ้างอิง", r"^บรรณานุกรม", r"^REFERENCES", r"^ภาคผนวก",
    r"^ประวัติผู้เขียน",
]
def section_page_map(doc, page_by_index):
    """List (heading_text, page) for the real (non-TOC) section headings —
    TOC lines carry a trailing tab+number, so we skip those."""
    out = []
    seen = set()
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if (not t or "\t" in para.text or
                re.search(r"\.{3,}\s*\(?\d+\)?\s*$", para.text)):
            # skip TOC rows (tabs or dot leaders)
            continue
        for pat in _SECTION_MARKERS:
            if re.match(pat, t, re.I):
                pg = page_by_index[i] if i < len(page_by_index) else None
                key = (t[:30], pg)
                if pg and key not in seen:
                    seen.add(key)
                    out.append((t[:40], pg))
                break
    return out

# ---------------------------------------------------------------------------
def section_anchors(doc):
    """{ดัชนี section: ข้อความแรกที่อ่านออกในส่วนนั้น} — ใช้ชี้ว่า "ส่วนนี้" คือหน้าไหน

    ปัญหาระดับ section (ระยะขอบ/ขนาดกระดาษ) เดิมรายงานว่า "Section 7" ซึ่งผู้อ่าน
    เปิด Word แล้วหาไม่เจอ เพราะ Word ไม่ได้แสดงเลข section ที่ไหนเลย. การยกข้อความ
    จริงบรรทัดแรกของส่วนนั้นมาให้ ทำให้ Ctrl+F ไปถึงหน้าที่ต้องแก้ได้ทันที
    """
    anchors = {}
    index = 0
    pending = ""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:tbl") and not pending:
            # section ที่มีแต่ตาราง (หน้าแนวนอนสำหรับตารางใหญ่) ไม่มีย่อหน้าให้ยึด
            # ต้องดึงข้อความจากช่องแรกที่มีตัวอักษร ไม่งั้นจะชี้จุดไม่ได้เลย
            for cell_text in child.itertext():
                cell_text = (cell_text or "").strip()
                if len(cell_text) >= 3:
                    pending = cell_text[:80]
                    break
            continue
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext()).strip()
        if text and not pending:
            # ตัดที่ขึ้นบรรทัดใหม่: ข้อความที่คร่อม line break ก๊อปไปวางใน Ctrl+F
            # แล้วค้นไม่เจอ เพราะ Word ถือว่าเป็นคนละบรรทัด
            pending = re.split(r"[\r\n\v]", text)[0].strip()[:80]
        # sectPr ที่ฝังใน pPr = จุดจบของ section นั้น
        ppr = child.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            anchors[index] = pending
            index += 1
            pending = ""
    anchors.setdefault(index, pending)
    return anchors


def check_sections(doc, F):
    anchors = section_anchors(doc)
    for i, s in enumerate(doc.sections):
        anchor = anchors.get(i, "")
        where = f" · ส่วนที่ขึ้นต้นด้วย “{anchor}”" if anchor else ""
        land = str(s.orientation).startswith("LANDSCAPE") or s.orientation == 1
        pw, ph = inch(s.page_width), inch(s.page_height)
        # A4 = 8.27 x 11.69
        exp = (11.69, 8.27) if land else (8.27, 11.69)
        if not (approx(pw, exp[0], .05) and approx(ph, exp[1], .05)):
            F.append(Finding("critical", "Page size",
                f"page {pw}x{ph}\" is not A4 ({exp[0]}x{exp[1]}\")",
                "A4 = 8.27 x 11.69 in (210 x 297 mm)",
                loc=f"Section {i}",
                correct=f"A4 = {exp[0]} x {exp[1]} นิ้ว"))
        T, B = inch(s.top_margin), inch(s.bottom_margin)
        L, Rm = inch(s.left_margin), inch(s.right_margin)
        H = inch(s.header_distance)
        eT, eB = 1.5, 1.0
        eL, eR = (1.0, 1.5) if land else (1.5, 1.0)
        for name, got, want in (("top", T, eT), ("bottom", B, eB),
                                ("left", L, eL), ("right", Rm, eR)):
            if not approx(got, want, .02):
                F.append(Finding("critical", "Margin",
                    f"{name} margin {got}\" — required {want}\"",
                    f"ตั้งขอบ{name} = {want} นิ้ว (Layout > Margins > Custom)",
                    loc=f"Section {i} ({'landscape' if land else 'portrait'}){where}",
                    correct=f"{want} นิ้ว"))
        # header distance (page number sits 1.0" from top) — divider pages use 0.49
        if H is not None and not (approx(H, 1.0, .03) or approx(H, 0.49, .03)):
            F.append(Finding("minor", "Header",
                f"header distance {H}\" (expected 1.0\")",
                "ตั้ง Header from top = 1.0 นิ้ว",
                loc=f"Section {i}{where}", correct="1.0 นิ้ว"))

def get_style_attrs(style_el):
    rpr = style_el.find(qn("w:rPr"))
    ppr = style_el.find(qn("w:pPr"))
    d = {"font": None, "fontCs": None, "sz": None, "szCs": None, "bold": None,
         "align": None, "firstLine": None, "left": None}
    if rpr is not None:
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            d["font"] = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi"))
            d["fontCs"] = rf.get(qn("w:cs")) or rf.get(qn("w:eastAsia"))
        s = rpr.find(qn("w:sz"));  d["sz"] = int(s.get(qn("w:val")))/2 if s is not None else None
        sc = rpr.find(qn("w:szCs")); d["szCs"] = int(sc.get(qn("w:val")))/2 if sc is not None else None
        b = rpr.find(qn("w:b"))
        if b is None:
            b = rpr.find(qn("w:bCs"))
        d["bold"] = _on_off(b)
    if ppr is not None:
        j = ppr.find(qn("w:jc"));  d["align"] = j.get(qn("w:val")) if j is not None else None
        ie = ppr.find(qn("w:ind"))
        if ie is not None:
            fl = ie.get(qn("w:firstLine")); d["firstLine"] = round(int(fl)/1440,3) if fl else None
            lf = ie.get(qn("w:left")) or ie.get(qn("w:start")); d["left"] = round(int(lf)/1440,3) if lf else None
    return d

def style_usage(doc):
    """{ชื่อสไตล์: (จำนวนย่อหน้าที่ใช้, ข้อความตัวอย่างแรก)} รวมการสืบทอด base style.

    ทำไมต้องรู้: ไฟล์ที่ก๊อปเนื้อหามาวางบนเทมเพลตมักมีสไตล์ TU_* ค้างอยู่ใน
    styles.xml ครบทุกตัว แต่ข้อความจริงกลับใช้ Heading 1/Normal. การรายงานว่า
    "สไตล์ TU_Main Heading_Chapter2 ขนาดเพี้ยน" ทั้งที่ไม่มีข้อความไหนใช้สไตล์นั้น
    = สั่งให้นักศึกษาแก้สิ่งที่ไม่มีผลกับหน้ากระดาษ และไม่มีทางชี้ได้ว่าอยู่ตรงไหน
    ในเล่ม ซึ่งเป็นข้อร้องเรียนหลักของผู้ใช้รายงานฉบับก่อน
    """
    usage = {}

    def bump(style, text):
        while style is not None:
            name = getattr(style, "name", None)
            if name:
                count, sample = usage.get(name, (0, ""))
                usage[name] = (count + 1, sample or (text or "").strip()[:120])
            style = getattr(style, "base_style", None)

    def visit(paragraph):
        text = paragraph.text
        bump(paragraph.style, text)
        for run in paragraph.runs:
            if run.style is not None and run.style is not paragraph.style:
                bump(run.style, run.text or text)

    for paragraph in doc.paragraphs:
        visit(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    visit(paragraph)
    return usage


def check_template_integrity(doc, F, profile):
    """Confirm the TU_* styles still exist (i.e. the student really started
    from the template and did not rebuild the file from a blank document)."""
    names = {s.name for s in doc.styles}
    required = ["TU_Chapter", "TU_Paragraph_Normal", "TU_Sub-heading 1",
                "TU_Sub-heading 2", "TU_Sub-heading 3",
                "TU_Para_Sub-heading 1", "TU_Para_Sub-heading 2",
                "TU_Para_Sub-heading 3"]
    missing = [n for n in required if n not in names]
    if not any(n.startswith("TU_Main Heading") for n in names):
        missing.append("TU_Main Heading _ChapterN")
    if missing:
        F.append(Finding("major", "Template integrity",
            "Core TULIBS styles are missing: " + ", ".join(missing),
            "คัดลอกเนื้อหาไปวางบนเทมเพลตต้นฉบับ TULIBS เพื่อให้ได้สไตล์ที่ถูกต้อง",
            loc="styles.xml",
            correct="มีสไตล์ TU_Chapter / TU_Paragraph_Normal / TU_Main Heading / TU_Sub-heading ครบ"))
    # ตรวจทุก TU style ที่รู้เกณฑ์ แต่ flag เฉพาะค่าที่ระบุชัดแล้วผิด;
    # ค่า None อาจ inherit จาก base style/docDefaults จึงไม่ควรถูกตัดสินว่าผิด.
    body_ascii = 12.0 if profile == "english-times" else 16.0
    expected = {
        "TU_Chapter": {"sz": 14.0 if profile == "english-times" else 18.0,
                       "szCs": 18.0, "bold": True},
        "TU_Paragraph_Normal": {"sz": body_ascii, "szCs": 16.0,
                                "bold": False, "firstLine": 0.8},
        "TU_Sub-heading 1": {"sz": body_ascii, "szCs": 16.0,
                             "bold": True, "firstLine": 0.8},
        "TU_Sub-heading 2": {"sz": body_ascii, "szCs": 16.0,
                             "bold": True, "firstLine": 1.1},
        "TU_Sub-heading 3": {"sz": body_ascii, "szCs": 16.0,
                             "bold": True, "firstLine": 1.4},
        "TU_Para_Sub-heading 1": {"sz": body_ascii, "szCs": 16.0,
                                  "bold": False, "firstLine": 1.19},
        "TU_Para_Sub-heading 2": {"sz": body_ascii, "szCs": 16.0,
                                  "bold": False, "firstLine": 1.63},
        "TU_Para_Sub-heading 3": {"sz": body_ascii, "szCs": 16.0,
                                  "bold": False, "firstLine": 1.63},
    }
    allowed_latin = PROFILES[profile]["latin_font"]
    allowed_thai = PROFILES[profile]["thai_font"]
    usage = style_usage(doc)
    dormant = {}          # ชื่อสไตล์ที่เพี้ยนแต่ไม่มีข้อความใช้ -> รายการที่เพี้ยน
    el = doc.styles.element
    for st in el.findall(qn("w:style")):
        nm = st.find(qn("w:name"))
        if nm is None: continue
        name = nm.get(qn("w:val"))
        rule_name = name
        if name and name.startswith("TU_Main Heading"):
            rule_name = "TU_Main Heading"
        rules = expected.get(rule_name)
        if rule_name == "TU_Main Heading":
            rules = {"sz": body_ascii, "szCs": 16.0, "bold": True,
                     "left": PROFILES[profile]["main_heading_left_in"]}
        if not rules:
            continue
        a = get_style_attrs(st)
        uses, sample = usage.get(name, (0, ""))
        labels = {"sz": "ขนาดละติน", "szCs": "ขนาดไทย",
                  "bold": "ตัวหนา", "firstLine": "เยื้องบรรทัดแรก",
                  "left": "เยื้องซ้าย"}

        def report(severity, key, got, want, unit):
            label = labels.get(key, key)
            if not uses:
                # ไม่มีข้อความใช้สไตล์นี้ → เก็บไว้สรุปรวมแถวเดียวท้ายสุด
                dormant.setdefault(name, []).append(f"{label} {got}{unit}")
                return
            where = (f"เห็นได้ที่ “{sample}”" if sample else "")
            F.append(Finding(
                severity, "Style drift",
                f"{name}: {label} = {got}{unit} (เทมเพลต = {want}{unit}) "
                f"— มีผลกับข้อความ {uses} จุดในเล่ม {where}",
                f"แก้ Modify Style > {name} ให้{label}ตรงกับเทมเพลต",
                loc=f"style {name}", correct=f"{want}{unit}"))

        for key, want in rules.items():
            got = a.get(key)
            if got is None:
                continue
            same = approx(got, want, .02) if isinstance(want, float) else got == want
            if same:
                continue
            unit = " นิ้ว" if key in ("firstLine", "left") else (" pt" if key.startswith("sz") else "")
            report("major" if key in ("sz", "szCs") else "minor", key, got, want, unit)
        for key, allowed, label in (("font", allowed_latin, "ฟอนต์ละติน"),
                                    ("fontCs", allowed_thai, "ฟอนต์ไทย")):
            got = a.get(key)
            if got and got not in allowed:
                labels[key] = label
                report("major", key, got, sorted(allowed)[0], "")

    if dormant:
        # แถวเดียวจบ: บอกว่ามีสไตล์เพี้ยนอยู่จริง แต่ยังไม่กระทบหน้ากระดาษ
        # ถ้าปล่อยให้กระจายเป็นหลายสิบแถว รายงานจะเต็มไปด้วยงานที่แก้แล้วไม่เห็นอะไรเปลี่ยน
        names_list = ", ".join(sorted(dormant)[:6])
        more = f" และอีก {len(dormant) - 6} สไตล์" if len(dormant) > 6 else ""
        F.append(Finding(
            "minor", "Style drift",
            f"มีสไตล์ TU_ ที่ค่าไม่ตรงเทมเพลต {len(dormant)} สไตล์ ({names_list}{more}) "
            f"แต่**ยังไม่มีข้อความไหนในเล่มใช้สไตล์เหล่านี้** จึงไม่มีผลกับหน้ากระดาษที่พิมพ์ออกมา",
            "ไม่ต้องแก้ก็ได้ถ้าไม่ได้ตั้งใจใช้สไตล์เหล่านี้ — สิ่งที่ควรดูแทนคือ "
            "หัวข้อ/ย่อหน้าที่ใช้อยู่จริงในเล่มมีขนาดและการเยื้องถูกต้องไหม "
            "(ดูรายการอื่นในตารางนี้). จะแก้ให้ครบก็ได้ที่ Home > กล่อง Styles > "
            "คลิกขวาที่ชื่อสไตล์ > Modify",
            loc="styles.xml", correct="ค่าตามเทมเพลต TULIBS"))

def check_fonts_and_sizes(doc, F, profile):
    p = PROFILES[profile]
    bad_fonts = {}     # font -> count
    bad_loc = {}       # font -> first occurrence
    bad_sizes = {}     # (style, script, got, want) -> [count, first location]
    color_issues = 0
    color_loc = ""
    seen_sizes = {}
    format_cache = {}

    def expected_size(style_name, script):
        if not style_name:
            return None
        latin_body = p["latin_body_pt"]
        if style_name == "TU_Chapter":
            return 18.0 if script == "thai" else (14.0 if profile == "english-times" else 18.0)
        if (style_name == "TU_Paragraph_Normal" or
                style_name.startswith("TU_Main Heading") or
                style_name.startswith("TU_Sub-heading") or
                style_name.startswith("TU_Para_Sub-heading")):
            return 16.0 if script == "thai" else latin_body
        return None

    for para, para_loc in iter_document_paragraphs(doc):
        paragraph_style = para.style
        style_name = paragraph_style.name if paragraph_style is not None else ""
        for run in para.runs:
            txt = run.text
            if not txt.strip():
                continue
            first_fmt = None
            for scr in scripts_in_text(txt):
                fmt = effective_run_format(run, para, doc, scr, format_cache, paragraph_style)
                first_fmt = first_fmt or fmt
                fname = fmt["font"]
                allowed = p["thai_font"] if scr == "thai" else p["latin_font"]
                exempt = (fname in SYMBOL_OK or (fname in MATH_ONLY and run_is_math(run)))
                if fname and fname not in allowed and not exempt:
                    bad_fonts[fname] = bad_fonts.get(fname, 0) + 1
                    if fname not in bad_loc:
                        bad_loc[fname] = f"{para_loc}: “{txt.strip()[:30]}” ({scr})"
                size_pt = fmt["size"]
                if size_pt:
                    seen_sizes[size_pt] = seen_sizes.get(size_pt, 0) + 1
                    want = expected_size(style_name, scr)
                    if want is not None and not approx(size_pt, want, .01):
                        key = (style_name, scr, size_pt, want)
                        if key not in bad_sizes:
                            bad_sizes[key] = [0, f"{para_loc}: “{txt.strip()[:30]}”"]
                        bad_sizes[key][0] += 1
            # ---- color (should be black/auto) ----
            if first_fmt and first_fmt["color"]:
                color_issues += 1
                if not color_loc:
                    color_loc = f"{para_loc}: “{txt.strip()[:30]}” ({first_fmt['color']})"
    if bad_fonts:
        top = sorted(bad_fonts.items(), key=lambda x: -x[1])[:8]
        total_bad = sum(bad_fonts.values())
        # widespread foreign fonts = critical; isolated (<=2 runs) = major/artifact
        sev = "critical" if total_bad > 2 else "major"
        loc = "; ".join(f"{n} → {bad_loc.get(n,'')}" for n, _ in top[:4])
        latin = sorted(p['latin_font'])[0]
        thai = "TH Sarabun New"
        correct = (f"{latin} {p['latin_body_pt']:g}pt (ละติน) / {thai} {p['body_pt']:g}pt (ไทย)"
                   if p['latin_body_pt'] != p['body_pt']
                   else f"{thai} {p['body_pt']:g}pt")
        F.append(Finding(sev, "Font",
            "Foreign fonts found (not allowed for this template): " +
            ", ".join(f"{n} x{c}" for n, c in top),
            "เลือกข้อความแล้วเปลี่ยนฟอนต์ให้ถูกต้อง (มัก copy-paste มาจากที่อื่น)",
            loc=loc, correct=correct))
    if bad_sizes:
        items = sorted(bad_sizes.items(), key=lambda kv: -kv[1][0])
        total = sum(meta[0] for _, meta in items)
        shown = items[:6]
        issue = "; ".join(
            f"{style}/{scr}: {got:g}pt → {want:g}pt x{meta[0]}"
            for (style, scr, got, want), meta in shown)
        locations = "; ".join(meta[1] for _, meta in shown[:4])
        F.append(Finding("critical" if total > 2 else "major", "Font size",
            f"พบขนาดฟอนต์ผิดจาก named style {total} run: {issue}",
            "ล้าง direct formatting ที่ทับสไตล์ หรือกำหนดขนาด run ให้ตรง named style",
            loc=locations,
            correct="เนื้อหา/หัวข้อย่อย: ไทย 16pt; ละติน 16pt (Times profile = 12pt); หัวบท 18pt"))
    # size distribution note
    if seen_sizes:
        common = sorted(seen_sizes.items(), key=lambda x: -x[1])[:6]
        F.append(Finding("info", "Font size distribution",
            "Explicit run sizes found: " +
            ", ".join(f"{s}pt x{c}" for s, c in common),
            f"Body should be {p['body_pt']}pt"
            + (f" (Thai) / {p['latin_body_pt']}pt (Latin)"
               if p['latin_body_pt'] != p['body_pt'] else "")
            + "; chapter/section headings 18pt; cover title 20pt."))
    if color_issues:
        F.append(Finding("minor", "Color",
            f"{color_issues} run(s) use non-black text colour",
            "เปลี่ยนสีตัวอักษรเป็นสีดำ (Automatic/Black) ทั้งเล่ม",
            loc=color_loc, correct="ดำล้วน (000000)"))

def check_page_numbering(doc, F):
    """Best-effort: confirm page-number fields exist in headers and flag if
    front-matter numbering style cannot be verified."""
    found_page_field = False
    for sec in doc.sections:
        for hdr in (sec.header, sec.first_page_header, sec.even_page_header):
            if hdr is None:
                continue
            xml = hdr._element.xml if hdr._element is not None else ""
            if "PAGE" in xml or "w:fldSimple" in xml or "instrText" in xml:
                found_page_field = True
    if not found_page_field:
        F.append(Finding("minor", "Page numbers",
            "Could not confirm automatic page-number fields in headers",
            "ใส่เลขหน้าอัตโนมัติ มุมขวาบน; ตรวจรูปแบบใน Word",
            loc="header", correct="ส่วนนำ (1)(2)… เริ่มที่บทคัดย่อ / เนื้อหา 1 2 3 เริ่มที่บทที่ 1, มุมขวาบน"))
    else:
        F.append(Finding("info", "Page numbers",
            "Page-number field(s) present in header — verify numbering scheme "
            "visually: front matter (1)(2)…, body 1 2 3, top-right corner."))

def check_structure(doc, F, profile):
    chapters = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        stl = para.style.name if para.style else ""
        if stl in ("TU_Chapter", "Heading 1") or re.match(
                r"^(บทที่|CHAPTER)\s+\d+", t, re.I):
            chapters.append(t[:60])
    if chapters:
        F.append(Finding("info", "Structure",
            f"{len(chapters)} chapter-level heading(s) detected",
            " | ".join(chapters[:12])))
    else:
        F.append(Finding("major", "Structure",
            "No chapter headings (TU_Chapter / 'CHAPTER n' / 'บทที่ n') found",
            "ใช้สไตล์ TU_Chapter กับบรรทัด 'บทที่ n'/'CHAPTER n' กึ่งกลาง หนา",
            loc="ทั้งเอกสาร",
            correct="'บทที่ n'/'CHAPTER n' 18pt หนา กึ่งกลาง (สไตล์ TU_Chapter)"))

def _cover_end_index(doc):
    """Index of the first non-cover paragraph (abstract/TOC/first chapter).
    Cover + second-language cover live before this."""
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        stl = para.style.name if para.style else ""
        if stl in ("TU_Chapter", "Heading 1"):
            return i
        if re.match(r"^(บทคัดย่อ|ABSTRACT|สารบัญ|CONTENTS|กิตติกรรมประกาศ"
                    r"|ACKNOWLEDG|บทที่\s*\d|CHAPTER\s*\d)", t, re.I):
            return i
    return min(len(doc.paragraphs), 40)

# Boundaries on the cover boilerplate that MUST be space/line separated.
# (left, right) — flagged only when they appear glued (no whitespace between).
_COVER_GLUE = [
    ("มหาวิทยาลัยธรรมศาสตร์", "ปีการศึกษา"),
    ("มหาวิทยาลัยธรรมศาสตร์", "ลิขสิทธิ์"),
    ("THAMMASAT UNIVERSITY", "ACADEMIC YEAR"),
    ("THAMMASAT UNIVERSITY", "COPYRIGHT"),
]
_COVER_GLUE_RE = [
    # Thai Buddhist year glued to the copyright line
    (re.compile(r"(25\d\d)(ลิขสิทธิ์)"), r"\1 \2", "ปี พ.ศ. ต้องเว้นวรรค/ขึ้นบรรทัดใหม่ก่อน 'ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์'"),
    # Gregorian year glued to COPYRIGHT
    (re.compile(r"(20\d\d)(COPYRIGHT)"), r"\1 \2", "ปี ค.ศ. ต้องเว้นวรรค/ขึ้นบรรทัดใหม่ก่อน 'COPYRIGHT OF THAMMASAT UNIVERSITY'"),
]

def check_cover(doc, F, profile):
    """Deterministic cover-page wording checks so results are identical on every
    machine/account (do NOT rely on a manual visual pass for these):
      - required boilerplate elements present
      - glued text / missing spaces at known cover boundaries
      - author name run together (data-driven: a spaced form exists elsewhere)
      - academic-year language (Thai cover = พ.ศ., English cover = ค.ศ.)
    """
    end = _cover_end_index(doc)
    cover = [(i, doc.paragraphs[i].text) for i in range(end)]
    cover_join = "\n".join(t for _, t in cover)

    # --- required boilerplate present ---
    if "มหาวิทยาลัยธรรมศาสตร์" not in cover_join and "THAMMASAT UNIVERSITY" not in cover_join.upper():
        F.append(Finding("major", "Cover",
            "ไม่พบข้อความ 'มหาวิทยาลัยธรรมศาสตร์'/'THAMMASAT UNIVERSITY' ในหน้าปก",
            "หน้าปกต้องมีชื่อมหาวิทยาลัยครบตามเทมเพลต",
            loc="หน้าปก", correct="…คณะ… มหาวิทยาลัยธรรมศาสตร์ ปีการศึกษา 25XX"))

    # --- glued boundaries (fixed token pairs) ---
    for i, t in cover:
        for left, right in _COVER_GLUE:
            if (left + right) in t.replace(" ", "") and (left + right) in t:
                F.append(Finding("major", "Cover",
                    f"คำติดกัน (ขาดเว้นวรรค/ขึ้นบรรทัด): '…{left}{right}…'",
                    f"แยก '{left}' กับ '{right}' ด้วยการเว้นวรรคหรือขึ้นบรรทัดใหม่",
                    loc=f"หน้าปก (para {i})",
                    correct=f"{left}  |  {right}"))
        for rgx, repl, hint in _COVER_GLUE_RE:
            m = rgx.search(t)
            if m:
                F.append(Finding("major", "Cover",
                    f"คำติดกันบนหน้าปก: '{m.group(0)}'",
                    hint, loc=f"หน้าปก (para {i})",
                    correct=rgx.sub(repl, m.group(0))))

    # --- author name run together (use the document's own spaced form as truth) ---
    title_re = re.compile(r"^(นางสาว|นาย|นาง|เด็กหญิง|เด็กชาย)\S")
    spaced_names = {}
    for _, t in cover:
        t = t.strip()
        if title_re.match(t) and " " in t and 2 <= len(t.split()) <= 4:
            spaced_names[t.replace(" ", "")] = t
    for i, t in cover:
        t = t.strip()
        if title_re.match(t) and " " not in t and t in spaced_names:
            F.append(Finding("major", "Cover",
                f"ชื่อผู้เขียนบนหน้าปกเขียนติดกัน: '{t}'",
                "เว้นวรรคระหว่างชื่อและนามสกุลให้ตรงกับหน้าอื่น",
                loc=f"หน้าปก (para {i})", correct=spaced_names[t]))

    # --- academic-year language (Thai=พ.ศ. / English=ค.ศ.) ---
    for i, t in cover:
        m = re.search(r"ปีการศึกษา\s*(\d{4})", t)
        if m and int(m.group(1)) < 2400:
            F.append(Finding("major", "Cover",
                f"ปกไทยใช้ปี ค.ศ. ({m.group(1)}) — ต้องเป็น พ.ศ.",
                f"แปลงเป็น พ.ศ.: {int(m.group(1)) + 543}",
                loc=f"หน้าปก (para {i})", correct=f"ปีการศึกษา {int(m.group(1)) + 543}"))
        m = re.search(r"ACADEMIC YEAR\s*(\d{4})", t, re.I)
        if m and int(m.group(1)) >= 2400:
            F.append(Finding("major", "Cover",
                f"English cover uses พ.ศ. ({m.group(1)}) — must be ค.ศ.",
                f"Convert to ค.ศ.: {int(m.group(1)) - 543}",
                loc=f"cover (para {i})", correct=f"ACADEMIC YEAR {int(m.group(1)) - 543}"))

# ---------------------------------------------------------------------------
# Word/character correctness checks (high precision, deterministic).
# Scope is deliberately narrow: unambiguously wrong words (curated list),
# malformed/extra characters, and canonical heading labels. No dictionary
# spell-check — anything requiring judgement is out of scope; the priority
# of this tool is template-format compliance.
# ---------------------------------------------------------------------------

# Curated Thai misspellings -> correct (Royal Institute standard). Only add
# forms that are unambiguously wrong. Where the wrong form is a prefix of the
# correct one, a negative look-ahead guard is added automatically below.
THAI_MISSPELL = {
    "อนุญาติ": "อนุญาต", "กระทันหัน": "กะทันหัน", "ผลลัพท์": "ผลลัพธ์",
    "สังเกตุ": "สังเกต", "ลายเซ็นต์": "ลายเซ็น", "เว็บไซท์": "เว็บไซต์",
    "เว็ปไซต์": "เว็บไซต์", "เวบไซต์": "เว็บไซต์", "อีเมล์": "อีเมล",
    "อัพเดท": "อัปเดต", "อัพเดต": "อัปเดต", "อัปเดท": "อัปเดต",
    "ซอฟท์แวร์": "ซอฟต์แวร์", "ปรากฎ": "ปรากฏ",
    # NOTE: สมมุติฐาน/สมมติฐาน ถูกทั้งคู่ตามพจนานุกรมราชบัณฑิตฯ — ห้ามใส่
    "อินเตอร์เน็ต": "อินเทอร์เน็ต",
    "อินเตอร์เน็ท": "อินเทอร์เน็ต", "แอพพลิเคชั่น": "แอปพลิเคชัน",
    "แอพพลิเคชัน": "แอปพลิเคชัน", "แอปพลิเคชั่น": "แอปพลิเคชัน",
    "แอพลิเคชั่น": "แอปพลิเคชัน", "เปอร์เซนต์": "เปอร์เซ็นต์",
    "โน๊ต": "โน้ต", "เบรค": "เบรก", "ค็อมพิวเตอร์": "คอมพิวเตอร์",
    "กราฟฟิก": "กราฟิก", "ประสิทธิ์ภาพ": "ประสิทธิภาพ",
    "โปรโมชั่น": "โปรโมชัน", "อัตราส่วน": "อัตราส่วน",
}
# drop entries whose "wrong" == "right" (kept above only for documentation)
THAI_MISSPELL = {w: r for w, r in THAI_MISSPELL.items() if w != r}

def _thai_misspell_patterns():
    pats = []
    for wrong, right in THAI_MISSPELL.items():
        if right.startswith(wrong) and len(right) > len(wrong):
            guard = re.escape(right[len(wrong)])
            pats.append((re.compile(re.escape(wrong) + f"(?!{guard})"), wrong, right))
        else:
            pats.append((re.compile(re.escape(wrong)), wrong, right))
    return pats

_THAI_MISSPELL_PATS = _thai_misspell_patterns()
_THAI_MISSPELL_META = {f"m{i}": (wrong, right)
                       for i, (_, wrong, right) in enumerate(_THAI_MISSPELL_PATS)}
_THAI_MISSPELL_RE = re.compile("|".join(
    f"(?P<m{i}>{rgx.pattern})" for i, (rgx, _, _) in enumerate(_THAI_MISSPELL_PATS)))

# Malformed Thai character sequences — always wrong.
_TH_MARK_ABOVE_BELOW = "ัิ-ฺ็-๎"  # sara/tone above & below
_MALFORMED = [
    (re.compile(r"[่-๋]{2,}"),
     "วรรณยุกต์ซ้อนกัน (มีวรรณยุกต์ ≥2 ตัวติดกัน)"),
    (re.compile(r"([ัิ-ฺ็-๎])\1"),
     "สระ/วรรณยุกต์ซ้ำตัวเดียวกันติดกัน"),
    (re.compile(r"(?:^|[\s\d\"“”'()\[\]A-Za-z])([" + _TH_MARK_ABOVE_BELOW + r"])"),
     "สระ/วรรณยุกต์ลอย (ไม่มีพยัญชนะนำหน้า)"),
    (re.compile(r"ํา"),
     "ใช้ ' ํา' (นิคหิต+สระอา) แทน 'ำ' (สระอำ)"),
]

# Zero-width / problematic invisible characters.
_INVISIBLE = {
    "​": "zero-width space (U+200B)",
    "﻿": "zero-width no-break space (U+FEFF)",
    " ": "non-breaking space (U+00A0)",
    "‎": "left-to-right mark (U+200E)",
    "‏": "right-to-left mark (U+200F)",
}

# Canonical thesis section headings — used to catch a misspelled heading.
_CANON_HEADINGS = [
    "บทคัดย่อ", "กิตติกรรมประกาศ", "สารบัญ", "สารบัญตาราง", "สารบัญภาพ",
    "รายการอ้างอิง", "บรรณานุกรม", "ภาคผนวก", "ประวัติผู้เขียน",
    "สารบัญตารางและภาพ",
]

def _lev(a, b):
    if a == b:
        return 0
    if abs(len(a) - len(b)) > 2:
        return 3
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        cur = [i]
        for j, cb in enumerate(b, 1):
            cur.append(min(prev[j] + 1, cur[j - 1] + 1,
                           prev[j - 1] + (ca != cb)))
        prev = cur
    return prev[-1]

def check_spelling(doc, F, profile):
    """Word/character correctness across the whole document (deterministic)."""
    # One occurrence per finding: the report can then give one exact sheet and
    # sort spelling corrections in physical document order.
    def group(sev, cat, title, hits, correct=""):
        for i, preview in hits:
            F.append(Finding(sev, cat, title, "",
                             loc=f"para {i}: “{preview}”", correct=correct))

    misspell_hits = {}      # right-form -> list[(i, context)]
    malformed_hits = {}     # desc -> list
    invis_hits = {}         # desc -> list
    double_space = []
    space_before = []       # space before ฯ (ไปยาลน้อย)
    dbl_word = []           # doubled latin word
    heading_typo = []

    # context around the exact match — so the report quotes the actual spot,
    # not the start of the paragraph
    def _ctx(t, start, end, pad=15):
        a, b = max(0, start - pad), min(len(t), end + pad)
        return ("…" if a else "") + t[a:b] + ("…" if b < len(t) else "")

    for i, para in enumerate(doc.paragraphs):
        t = para.text
        if not t.strip():
            continue

        # 1) curated Thai misspellings
        for m in _THAI_MISSPELL_RE.finditer(t):
            wrong, right = _THAI_MISSPELL_META[m.lastgroup]
            misspell_hits.setdefault(f"{wrong} → {right}", []).append(
                (i, _ctx(t, m.start(), m.end())))

        # 2) malformed Thai sequences
        for rgx, desc in _MALFORMED:
            m = rgx.search(t)
            if m:
                malformed_hits.setdefault(desc, []).append(
                    (i, _ctx(t, m.start(), m.end())))

        # 3) invisible / problematic characters
        for ch, desc in _INVISIBLE.items():
            p = t.find(ch)
            if p != -1:
                invis_hits.setdefault(desc, []).append(
                    (i, _ctx(t, p, p + 1)))

        # 4) double spaces (skip tab-aligned rows: TOC / tables pad with spaces)
        if "\t" not in t:
            m = re.search(r"  +", t)
            if m:
                double_space.append((i, _ctx(t, m.start(), m.end())))
        # space before ฯ (ไปยาลน้อยต้องติดคำ) — ยกเว้น ฯลฯ / ฯพณฯ ซึ่งเว้นหน้าได้.
        # ไม้ยมก (ๆ) ไม่ตรวจ: ราชบัณฑิตฯ ให้เว้นวรรคหน้า-หลัง ("ต่าง ๆ" ถูกต้อง)
        m = re.search(r"\sฯ(?!ลฯ|พณ)", t)
        if m:
            space_before.append((i, _ctx(t, m.start(), m.end())))

        # 5) doubled Latin word (the the / data data)
        m = re.search(r"\b([A-Za-z]{2,})\s+\1\b", t)
        if m:
            dbl_word.append((i, _ctx(t, m.start(), m.end())))

        # 6) misspelled section heading (near-miss of a canonical label)
        stl = para.style.name if para.style else ""
        head_like = stl in ("TU_Chapter", "Heading 1", "Heading 2") or (
            len(t.strip()) <= 20 and re.match(
                r"^(บท|สารบัญ|ภาคผนวก|ประวัติ|กิตติ|รายการ|บรรณาน)", t.strip()))
        if head_like:
            s = t.strip()
            # a valid heading may extend a canonical label (e.g. "บทคัดย่อภาษาไทย",
            # "ภาคผนวก ก", "สารบัญตาราง") — only flag a genuine near-miss where
            # neither string is a prefix of the other.
            if any(s.startswith(c) or c.startswith(s) for c in _CANON_HEADINGS):
                pass
            else:
                for canon in _CANON_HEADINGS:
                    if 1 <= _lev(s, canon) <= 2:
                        heading_typo.append((i, f"“{s}” ≈ “{canon}”"))
                        break

    for right_form, hits in misspell_hits.items():
        group("major", "Spelling", f"คำสะกดผิด: แก้เป็น “{right_form.split(' → ')[1]}”",
              hits, correct=right_form.split(" → ")[1])
    for desc, hits in malformed_hits.items():
        group("major", "Spelling", f"อักขระผิดรูป: {desc}", hits)
    for desc, hits in invis_hits.items():
        group("minor", "Spelling", f"อักขระแปลกปลอม: {desc} (ลบออก)", hits)
    group("minor", "Spelling", "เว้นวรรคซ้อน (เคาะ space เกิน 1 ครั้ง)",
          double_space, correct="เว้นวรรคครั้งเดียว")
    group("minor", "Spelling", "เว้นวรรคก่อน ฯ (ไปยาลน้อยต้องติดกับคำ; ฯลฯ ไม่นับ)",
          space_before)
    group("minor", "Spelling", "คำภาษาอังกฤษซ้ำติดกัน (doubled word)", dbl_word)
    if heading_typo:
        for i, preview in heading_typo:
            F.append(Finding("major", "Spelling",
                "หัวข้อ section อาจสะกดผิดจากรูปมาตรฐาน",
                "แก้ให้ตรงชื่อหัวข้อมาตรฐาน", loc=f"para {i}: {preview}"))

def check_spacing(doc, F, profile):
    """Flag paragraphs whose line spacing differs from the template value.
    2024 templates use single (1.0); the 2023 Times variant uses 1.5."""
    expected = PROFILES[profile]["line_spacing"]
    off = []
    for para, loc in iter_document_paragraphs(doc):
        ls = para.paragraph_format.line_spacing
        if ls is None:
            try:
                for style in _style_chain(para.style):
                    ls = style.paragraph_format.line_spacing
                    if ls is not None:
                        break
            except (AttributeError, KeyError):
                pass
        if ls is None:
            continue
        try:
            # multiplier เป็น float; exact/at-least spacing เป็น Length (มี .pt)
            got = float(ls) if not hasattr(ls, "pt") else None
            if got is None or abs(got - expected) > 0.1:
                shown = f"{ls.pt:g}pt (exact/at least)" if hasattr(ls, "pt") else str(ls)
                off.append((loc, shown))
        except (TypeError, ValueError):
            off.append((loc, str(ls)))
    if off:
        locations = "; ".join(f"{loc} = {got}" for loc, got in off[:6])
        more = f" (+{len(off) - 6} จุดอื่น)" if len(off) > 6 else ""
        F.append(Finding("minor", "Line spacing",
            f"{len(off)} paragraph(s) differ from the template line spacing "
            f"({expected:g}){more}",
            f"ตั้งระยะห่างบรรทัด = {expected:g} ทั้งเล่ม (Paragraph > Line spacing)",
            loc=locations, correct=f"{expected:g} เท่า"))

# ---------------------------------------------------------------------------
def build_payload(path, profile, F):
    crit = sum(1 for x in F if x.sev == "critical")
    maj = sum(1 for x in F if x.sev == "major")
    minr = sum(1 for x in F if x.sev == "minor")
    return {"file": os.path.basename(path), "profile": profile,
            "profile_label": PROFILES[profile]["label"],
            "counts": {"critical": crit, "major": maj, "minor": minr},
            "findings": [x.row() for x in F]}


def inspect_document(doc, path, profile=None, started_at=None, page_data=None):
    """ตรวจ Document ที่โหลดแล้วและคืน (payload, findings).

    แยกการโหลดไฟล์ออกจากการตรวจเพื่อให้ ``check_all.py`` ใช้ Document instance
    เดียวกับ deep checks ได้ ลดการ unzip/parse OOXML ซ้ำหนึ่งรอบ.
    """
    t0 = started_at if started_at is not None else time.time()
    if profile is None:
        try:
            profile = detect_profile(doc)
        except Exception:
            profile = "thai"   # safe default; user can override with --profile
    F = []
    # Run each check in isolation. A single check that hits an unexpected
    # document quirk should never abort the whole run and lose the other
    # findings — that used to force a full re-run. Instead we record the
    # failure as an info finding and keep going, so the report is still useful.
    checks = [
        ("Page setup", check_sections),
        ("Template integrity", lambda d, f: check_template_integrity(d, f, profile)),
        ("Cover page", lambda d, f: check_cover(d, f, profile)),
        ("Spelling", lambda d, f: check_spelling(d, f, profile)),
        ("Fonts & sizes", lambda d, f: check_fonts_and_sizes(d, f, profile)),
        ("Line spacing", lambda d, f: check_spacing(d, f, profile)),
        ("Page numbers", check_page_numbering),
        ("Structure", lambda d, f: check_structure(d, f, profile)),
    ]
    for name, fn in checks:
        try:
            fn(doc, F)
        except Exception as e:
            F.append(Finding("info", name,
                f"ตรวจส่วนนี้อัตโนมัติไม่สำเร็จ ({type(e).__name__}: {e}) — "
                "ข้ามไปตรวจส่วนอื่นต่อ",
                "ตรวจส่วนนี้ด้วยตาเพิ่มเติมจากในไฟล์",
                loc=name))

    # ---- attach physical sheet numbers to every finding location ----
    try:
        resolved_pages = page_data if page_data is not None else resolve_page_data(path, doc)
        page_by_index, has_pages = unpack_page_data(resolved_pages, doc)
        enrich_locations_with_pages(F, resolved_pages)
        sec_pages = section_page_map(doc, page_by_index) if has_pages else []
    except Exception:
        resolved_pages = {}
        page_by_index, has_pages, sec_pages = [], False, []

    elapsed = round(time.time() - t0, 2)
    payload = build_payload(path, profile, F)
    payload["elapsed_sec"] = elapsed
    payload["has_page_numbers"] = has_pages
    payload["section_pages"] = sec_pages
    if isinstance(resolved_pages, dict):
        payload["page_source"] = resolved_pages.get("source", "unavailable")
        payload["total_pages"] = resolved_pages.get("total_pages", 0)
        payload["page_match_ratio"] = resolved_pages.get("match_ratio", 0.0)
        payload["page_map_warning"] = resolved_pages.get("warning", "")
    return payload, F


def print_summary(doc, path, payload, F):
    profile = payload["profile"]
    has_pages = payload.get("has_page_numbers", False)
    sec_pages = payload.get("section_pages", [])
    elapsed = payload.get("elapsed_sec", 0)
    crit, maj, minr = (payload["counts"][k] for k in ("critical", "major", "minor"))

    print("=" * 70)
    print(f"TULIBS thesis .docx check — {os.path.basename(path)}")
    print(f"Profile: {profile}  ({PROFILES[profile]['label']})")
    print(f"Sections: {len(doc.sections)}  Paragraphs: {len(doc.paragraphs)}  "
          f"Tables: {len(doc.tables)}")
    print(f"Findings: 🔴 {crit} critical  🟠 {maj} major  🟡 {minr} minor")
    print(f"Auto-scan time: {elapsed} s   "
          f"Physical sheets: {payload.get('page_source', 'unavailable')} "
          f"({payload.get('total_pages', 0)} sheets)")
    if sec_pages:
        print("Section pages (for APA/visual page-tagging): "
              + "; ".join(f"{h}=แผ่นที่ {p}" for h, p in sec_pages[:16]))
    print("=" * 70)
    for x in F:
        loc = f" @ {x.loc}" if x.loc else ""
        print(f"{SEV.get(x.sev,'')} [{x.cat}]{loc} {x.msg}")
        if x.detail:
            print(f"      → {x.detail}")
    print("=" * 70)
    print("NOTE: primary scope is template-format compliance (page setup, "
          "fonts/sizes, styles, line spacing, structure, cover-page wording). "
          "Word/character checks are limited to unambiguous errors only: "
          "curated Thai misspellings, malformed Thai sequences, invisible/extra "
          "chars, doubled words, section-heading spelling. Run check_deep.py (or "
          "check_all.py) for TOC/caption/APA mechanical checks; approval-page "
          "committee layout and semantic APA details still need a reading pass.")


def run(path, profile=None, json_out=None, report_out=None):
    t0 = time.time()
    doc = Document(path)
    payload, F = inspect_document(doc, path, profile, started_at=t0)
    print_summary(doc, path, payload, F)

    if json_out:
        with open(json_out, "w", encoding="utf-8") as fh:
            json.dump(payload, fh, ensure_ascii=False, indent=2)
        print(f"JSON written: {json_out}")
    if report_out:
        try:
            from make_report import build_report
        except ImportError:
            sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
            from make_report import build_report
        build_report(payload, report_out)
        print(f"DOCX report written: {report_out}")
    return F

def main():
    ap = argparse.ArgumentParser(description="TULIBS thesis .docx format checker")
    ap.add_argument("docx")
    ap.add_argument("--profile", choices=list(PROFILES.keys()), default=None)
    ap.add_argument("--json", default=None)
    ap.add_argument("--report", default=None,
                    help="write a .docx correction-table report to this path")
    a = ap.parse_args()
    if not os.path.exists(a.docx):
        sys.exit(f"File not found: {a.docx}")
    run(a.docx, a.profile, a.json, a.report)

if __name__ == "__main__":
    main()
