#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""check_deep.py — ตรวจอัตโนมัติชั้นที่สอง: โครงสร้างเอกสาร + APA 7 เชิงกลไก

ทำไมต้องมีไฟล์นี้
-----------------
`check_docx.py` ตรวจ "ค่าตั้งค่า" (ฟอนต์/ระยะขอบ/สไตล์) ส่วนงานที่เหลืออีกหลายอย่าง
เดิมโยนให้ LLM ไล่อ่านเอง — ซึ่งช้าที่สุดในกระบวนการทั้งหมด และผลไม่คงที่ระหว่างรอบ
แต่พอดูจริง ๆ งานพวกนั้นจำนวนมาก **ไม่ต้องใช้วิจารณญาณเลย** เช่น "เลขหน้าในสารบัญ
ตรงกับหน้าจริงไหม" เป็นการเทียบตัวเลขล้วน ๆ, "caption ตารางอยู่เหนือตารางไหม" เป็นการ
ดูลำดับ element ใน XML

ไฟล์นี้จึงดึงงานกลุ่มนั้นมาทำด้วยสคริปต์ ได้ทั้งเร็วขึ้น (วินาที → เศษวินาที) และแม่นขึ้น
(deterministic ตอบเท่ากันทุกรอบ) เหลือให้ LLM เฉพาะสิ่งที่ต้องอ่านแล้วตัดสินจริง ๆ เช่น
ปกครบองค์ประกอบไหม, ชื่อบทความควรเป็น sentence case หรือยัง, รายการนี้เป็นแหล่งประเภทอะไร

หลักที่ยึด: **ตรวจเฉพาะสิ่งที่ผิดแบบไม่ต้องตีความ** ถ้ากฎไหนมีข้อยกเว้นที่ต้องใช้
วิจารณญาณ ปล่อยให้ LLM ทำ อย่าใส่มาที่นี่ — false positive ในงานตรวจวิทยานิพนธ์
แพงกว่าการตรวจไม่เจอ เพราะนักศึกษาจะเสียเวลาแก้สิ่งที่ถูกอยู่แล้ว

ใช้:
    python3 scripts/check_deep.py thesis.docx --json deep.json
    python3 scripts/check_deep.py thesis.docx --json deep.json --only template
"""
import sys, os, re, json, argparse, time

try:
    from docx import Document
except ImportError:
    sys.exit("ต้องติดตั้ง python-docx ก่อน:  pip install python-docx")

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from check_docx import (build_page_index, detect_profile, enrich_dict_locations_with_pages,
                        PROFILES, is_thai, qn, resolve_page_data, run_is_bold,
                        unpack_page_data)

SEV_ICON = {"critical": "🔴", "major": "🟠", "minor": "🟡", "info": "ℹ️"}

TABLE_CAP = re.compile(r"^\s*(ตารางที่|Table)\s*([\d๐-๙]+[\.\-][\d๐-๙]+|[\d๐-๙]+)", re.I)
FIG_CAP   = re.compile(r"^\s*(ภาพที่|รูปที่|Figure)\s*([\d๐-๙]+[\.\-][\d๐-๙]+|[\d๐-๙]+)", re.I)
CHAPTER   = re.compile(r"^\s*(บทที่\s*([\d๐-๙]+)|CHAPTER\s*([\dIVX]+))\s*$", re.I)
TOC_ROW   = re.compile(r"^(.*?)[\s\.\t]*\(?(\d+)\)?\s*$")
KEYWORDS  = re.compile(r"^\s*(คำสำคัญ|คำสําคัญ|Keywords?)\s*[:：]", re.I)


def F(sev, cat, issue, detail="", loc="", correct="", domain="template"):
    return {"domain": domain, "severity": sev, "category": cat,
            "location": loc, "issue": issue, "detail": detail, "correct": correct}


def loc_of(i, page):
    return f"para {i}" if i is not None else "ตำแหน่งในเอกสาร"


# ---------------------------------------------------------------------------
# TEMPLATE — โครงสร้างเอกสาร
# ---------------------------------------------------------------------------
def check_toc_pages(doc, pages, out):
    """เทียบเลขหน้าในสารบัญกับหน้าที่หัวข้อนั้นอยู่จริง

    สารบัญคือบรรทัดที่ลงท้ายด้วยตัวเลข และมี tab หรือจุดไข่ปลาคั่น เราจับคู่ข้อความ
    หัวข้อกับหัวข้อจริงในเล่ม (บรรทัดเดียวกันแต่ไม่มีเลขหน้าต่อท้าย) แล้วเทียบตัวเลข
    ทำให้ครอบคลุม *ทุก* รายการ ไม่ใช่สุ่ม 5 รายการแบบที่คนทำไหว
    """
    toc_rows, real = [], {}
    paragraphs = doc.paragraphs
    body = [(i, (p.text or "").strip(), pages[i] if i < len(pages) else None)
            for i, p in enumerate(paragraphs)]
    nonempty = [b for b in body if b[1]]

    for n, (i, t, pg) in enumerate(nonempty):
        raw = paragraphs[i].text or ""
        if ("\t" in raw or re.search(r"\.{3,}", raw)) and re.search(r"\d\s*$", t):
            m = TOC_ROW.match(re.sub(r"\.{3,}", "\t", raw).strip())
            if m and m.group(1).strip():
                toc_rows.append((i, pg, m.group(1).strip(" .\t"), int(m.group(2))))
            continue
        real.setdefault(re.sub(r"\s+", " ", t), pg)
        # หัวบทในเล่มมักแยกสองบรรทัด ("บทที่ 1" / "บทนำ") แต่สารบัญพิมพ์รวมบรรทัดเดียว
        # ("บทที่ 1 บทนำ") — ถ้าไม่รวมให้ จะจับคู่ไม่ได้แล้วข้ามการตรวจไปเงียบ ๆ
        if n + 1 < len(nonempty):
            joined = re.sub(r"\s+", " ", f"{t} {nonempty[n+1][1]}")
            real.setdefault(joined, pg)

    if not toc_rows:
        return
    real_prefix = {}
    for key, value in real.items():
        if len(key) >= 8:
            real_prefix.setdefault(key[:20], value)
    checked = mismatch = 0
    for i, pg, title, claimed in toc_rows:
        key = re.sub(r"\s+", " ", title)
        actual = real.get(key)
        if actual is None:                      # หาหัวข้อจริงไม่เจอ — อาจสะกดต่างกันเล็กน้อย
            actual = real_prefix.get(key[:20]) if len(key) >= 8 else None
        if actual is None:
            continue
        checked += 1
        if actual != claimed:
            mismatch += 1
            out.append(F("major", "สารบัญ",
                         f"สารบัญระบุ '{title}' อยู่หน้า {claimed} แต่หัวข้อจริงอยู่หน้า {actual}",
                         "inspection-checklist B4: เลขหน้าในสารบัญต้องตรงกับหน้าจริง — "
                         "มักเพี้ยนเพราะแก้เนื้อหาแล้วไม่ได้กด Update Table",
                         loc_of(i, pg), f"{title} ... {actual}"))
    if checked and not mismatch:
        out.append(F("info", "สารบัญ",
                     f"เทียบเลขหน้าในสารบัญกับหน้าจริงครบ {checked} รายการ — ตรงทั้งหมด",
                     "", "สารบัญ"))
    elif not checked:
        out.append(F("info", "สารบัญ",
                     "พบบรรทัดสารบัญแต่จับคู่กับหัวข้อจริงไม่ได้ — ต้องเทียบด้วยตา",
                     "ชื่อหัวข้อในสารบัญอาจสะกดไม่ตรงกับหัวข้อในเล่ม", "สารบัญ"))


# หมายเหตุ: การตรวจตำแหน่ง caption (ตารางเหนือ/ภาพใต้) และ "ภาพ/ตารางที่ไม่มี
# คำบรรยาย" ถูกถอดออกโดยตั้งใจ.
#
# เหตุผล: ตำแหน่ง caption ไม่ได้ถูกกำหนดตายตัวในเทมเพลต TULIBS และเมื่อวัดกับ
# วิทยานิพนธ์จริง 10 เล่ม กฎชุดนี้เป็นแหล่ง false positive ที่ใหญ่ที่สุดที่เหลืออยู่
# (~150 แถว) เพราะแยกไม่ออกระหว่าง caption กับประโยคเนื้อหาที่ขึ้นต้นเหมือนกัน —
# "ตารางที่ 4.3 พบว่า ผู้ตอบแบบสอบถามส่วนใหญ่…" คือเนื้อความ ไม่ใช่ caption ที่วางผิด
# ที่ — และยังไปฟ้องตาราง layout ในหน้าสารบัญ/หน้าอนุมัติว่าขาดคำบรรยาย.
#
# ถ้าจะเอากลับมา ต้องแยก caption จริงออกจากร้อยแก้วให้ได้ก่อน (เช่นดูจากสไตล์
# Caption หรือย่อหน้าที่มีแต่ข้อความนั้นล้วน ๆ) แล้วเพิ่มเคสร้อยแก้วเหล่านี้ลง
# scripts/test_false_positives.py ก่อนเปิดใช้


def check_chapter_layout(doc, pages, out):
    """แต่ละบทต้องขึ้นหน้าใหม่ และเลขบทต้องต่อเนื่อง"""
    THAI_D = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")
    ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10}
    chaps = []
    for i, p in enumerate(doc.paragraphs):
        raw = p.text or ""
        if "\t" in raw or re.search(r"\.{3,}", raw):
            continue                                        # ข้ามบรรทัดสารบัญ
        m = CHAPTER.match(raw.strip())
        if m:
            num = m.group(2) or m.group(3) or ""
            num = ROMAN.get(num.upper(), None) if not num.translate(THAI_D).isdigit() else int(num.translate(THAI_D))
            chaps.append((i, pages[i] if i < len(pages) else None, raw.strip(), num))

    for n, (i, pg, txt, num) in enumerate(chaps):
        if n and num is not None and chaps[n-1][3] is not None and num != chaps[n-1][3] + 1:
            out.append(F("major", "โครงสร้างบท",
                         f"เลขบทไม่ต่อเนื่อง: '{chaps[n-1][2]}' แล้วข้ามไป '{txt}'",
                         "inspection-checklist B5: เลขบทต้องเรียง 1..n ไม่ข้ามไม่ซ้ำ",
                         loc_of(i, pg), f"บทถัดจาก {chaps[n-1][2]} ควรเป็นบทที่ {chaps[n-1][3] + 1}"))
        if n and pg and chaps[n-1][1] and pg == chaps[n-1][1]:
            out.append(F("major", "โครงสร้างบท",
                         f"'{txt}' ขึ้นหน้าเดียวกับบทก่อนหน้า (หน้า {pg})",
                         "inspection-checklist B5: แต่ละบทต้องขึ้นหน้าใหม่",
                         loc_of(i, pg), "แทรก Page Break ก่อนหัวบท"))


def check_keywords_bold(doc, pages, out):
    """คำสำคัญ:/Keywords: ต้องเป็นตัวหนา"""
    for i, p in enumerate(doc.paragraphs):
        if not KEYWORDS.match(p.text or ""):
            continue
        pg = pages[i] if i < len(pages) else None
        match = KEYWORDS.match(p.text or "")
        label = match.group(1) if match else (p.text or "").split(":")[0].strip()
        label_end = match.end() if match else len(label)
        cursor = 0
        label_runs = []
        for run in p.runs:
            start, end = cursor, cursor + len(run.text or "")
            if start < label_end and (run.text or "").strip():
                label_runs.append(run)
            cursor = end
        if label_runs and not all(run_is_bold(run, p, doc) for run in label_runs):
            out.append(F("minor", "บทคัดย่อ",
                         f"'{label}:' ไม่ได้เป็นตัวหนา",
                         "inspection-checklist B3: คำว่า คำสำคัญ:/Keywords: ต้องเป็นตัวหนา",
                         loc_of(i, pg), f"**{label}:** <คำสำคัญ>"))
        if re.match(r"^\s*Keyword\s*[:：]", p.text or ""):
            out.append(F("minor", "บทคัดย่อ",
                         "ใช้ 'Keyword' เอกพจน์",
                         "เทมเพลตกำหนดให้ใช้รูปพหูพจน์", loc_of(i, pg), "Keywords:"))


# ---------------------------------------------------------------------------
# APA 7 — เชิงกลไก (เฉพาะที่ผิดแน่นอน ไม่ต้องตีความ)
# ---------------------------------------------------------------------------
REF_HEAD = re.compile(r"^\s*(รายการอ้างอิง|บรรณานุกรม|เอกสารอ้างอิง|REFERENCES?|BIBLIOGRAPHY)\s*$", re.I)
APPENDIX = re.compile(r"^\s*(ภาคผนวก|APPENDIX|ประวัติผู้เขียน|BIOGRAPHY)", re.I)
INTEXT = re.compile(r"\(([^()]{0,80}?),?\s*(?:19|20|25)\d{2}[^()]{0,20}\)|[฀-๿A-Za-z\.\s]{2,40}\s\((?:19|20|25)\d{2}\)")


def reference_entries(doc, pages, paragraphs=None):
    """คืน [(para_index, page, text)] ของรายการอ้างอิง (ไม่รวมหัวข้อและภาคผนวก)"""
    paragraphs = paragraphs if paragraphs is not None else doc.paragraphs
    start = None
    seen_chapter = False
    fallback = []
    for i, p in enumerate(paragraphs):
        raw = p.text or ""
        if "\t" in raw or re.search(r"\.{3,}", raw):
            continue
        if re.match(r"^\s*(บทที่\s*1(?:\s|$)|CHAPTER\s*(?:1|I)(?:\s|$))", raw, re.I):
            seen_chapter = True
        if REF_HEAD.match(raw.strip()):
            fallback.append(i)
            if seen_chapter:
                start = i
                break
    if start is None and fallback:
        start = fallback[-1]
    if start is None:
        return []
    out = []
    for i in range(start + 1, len(paragraphs)):
        t = (paragraphs[i].text or "").strip()
        if not t:
            continue
        if APPENDIX.match(t):
            break
        if len(t) < 15:                     # กันหัวข้อย่อย/บรรทัดสั้นที่ไม่ใช่รายการ
            continue
        out.append((i, pages[i] if i < len(pages) else None, t))
    return out


def entry_is_thai(text):
    """ตัดสินภาษาจากส่วนชื่อผู้แต่งก่อนปี ไม่ใช่จากอักษรไทยที่อาจอยู่ในชื่อแปล."""
    author = re.split(r"\((?:19|20|25)\d{2}", text, maxsplit=1)[0][:120]
    return any(is_thai(c) for c in author)


def _effective_para_value(paragraph, attr):
    value = getattr(paragraph.paragraph_format, attr)
    if value is not None:
        return value
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        value = getattr(style.paragraph_format, attr)
        if value is not None:
            return value
        style = style.base_style
    return None


_URL_RE = re.compile(r"(?:https?://|www\.|doi\.org/|10\.\d{4,})\S*")


def _inside_url(text, pos):
    """ตำแหน่ง pos อยู่ในช่วงที่เป็น URL/DOI หรือไม่."""
    return any(m.start() <= pos < m.end() for m in _URL_RE.finditer(text))


def _after_title_span(text):
    """คืนข้อความ 'หลังชื่อเรื่อง' — คือหลังจุดแรกที่ตามหลังวงเล็บปีพิมพ์.

    รูปแบบ APA คือ ผู้แต่ง. (ปี). ชื่อเรื่อง. สำนักพิมพ์. ดังนั้นทุกอย่างที่อยู่
    ก่อนจุดถัดจาก '(ปี).' ยังเป็นชื่อเรื่องอยู่ กฎที่ว่าด้วยสำนักพิมพ์จึงไม่ควรมอง.
    ถ้าหารูปแบบไม่เจอ คืนค่าว่างเพื่อไม่ให้เดา — ตรวจไม่ได้ดีกว่าฟ้องผิด.
    """
    m = re.search(r"\((?:19|20|2[45])\d{2}[a-z]?(?:,[^)]*)?\)\s*\.?\s*", text)
    if not m:
        return ""
    rest = text[m.end():]
    m2 = re.search(r"[\.\?!]\s", rest)
    return rest[m2.end():] if m2 else ""


# คำที่ขึ้นต้นประโยคแล้วบังเอิญมาก่อนวงเล็บปี — ไม่ใช่นามสกุลผู้แต่ง.
# ถ้าไม่กัน ประโยคอย่าง "Table 3 shows GDP (2019)" จะกลายเป็น finding ระดับ
# critical ว่า "อ้าง Table (2019) แต่ไม่มีในรายการอ้างอิง" ซึ่งทำให้นักศึกษา
# ไล่หาสิ่งที่ไม่มีอยู่จริง — ต้นทุนสูงกว่าการตรวจไม่เจอมาก
_NOT_A_SURNAME = {
    "the", "this", "these", "those", "that", "table", "figure", "fig", "chart",
    "appendix", "chapter", "section", "and", "but", "for", "from", "with",
    "since", "when", "while", "where", "there", "here", "it", "its", "in", "on",
    "at", "by", "as", "an", "all", "also", "however", "moreover", "thus",
    "therefore", "according", "based", "data", "year", "years", "study",
    "research", "results", "source", "note", "using", "used", "see", "such",
    "was", "were", "has", "had", "have", "is", "are", "been", "between",
    "during", "after", "before", "january", "february", "march", "april",
    "may", "june", "july", "august", "september", "october", "november",
    "december", "thailand", "bangkok", "asean", "covid", "gdp", "who",
}

# นามสกุล + ปี ต้องอยู่ติดกันตามรูปแบบการอ้างจริง (Smith, 2019 / Smith (2019) /
# Smith et al. (2019) / Smith & Jones (2019)) เท่านั้น — การปล่อยให้มีข้อความ
# คั่นกลางได้ 30 ตัวอักษรคือที่มาของการจับคำสุ่มมาเป็นชื่อผู้แต่ง
_INTEXT_RE = re.compile(
    r"\b([A-Z][a-zA-Z\-']{2,})"
    r"(?:\s+(?:et\s+al\.|and\s+[A-Z][a-zA-Z\-']{2,}|&\s*[A-Z][a-zA-Z\-']{2,}))?"
    r"\s*,?\s*[\(,]\s*((?:19|20)\d{2})[a-z]?\s*[\),;\]]")


def check_apa_mechanical(doc, pages, out):
    paragraphs = doc.paragraphs
    entries = reference_entries(doc, pages, paragraphs)
    if not entries:
        out.append(F("info", "APA7", "ไม่พบส่วนรายการอ้างอิงในไฟล์นี้", "",
                     "รายการอ้างอิง", "", domain="apa7"))
        return

    def add(sev, i, pg, n, issue, detail, correct=""):
        out.append(F(sev, "APA7 เชิงรูปแบบ", issue, detail,
                     f"รายการที่ {n} · {loc_of(i, pg)}", correct, domain="apa7"))

    for n, (i, pg, t) in enumerate(entries, 1):
        thai = entry_is_thai(t)

        # Hanging indent ตรวจได้เมื่อย่อหน้าหรือ style ระบุค่าไว้จริง; ถ้า inherit
        # ลึกกว่านี้จน resolve ไม่ได้ ให้ manual pass ตัดสินแทน ไม่เดา.
        paragraph = paragraphs[i]
        first = _effective_para_value(paragraph, "first_line_indent")
        left = _effective_para_value(paragraph, "left_indent")
        first_in = getattr(first, "inches", None)
        left_in = getattr(left, "inches", None)
        if ((first_in is not None and abs(first_in + 0.5) > 0.03) or
                (left_in is not None and abs(left_in - 0.5) > 0.03)):
            got = f"first-line {first_in if first_in is not None else 'ไม่ระบุ'}\", left {left_in if left_in is not None else 'ไม่ระบุ'}\""
            add("major", i, pg, n, f"hanging indent ไม่ตรงเทมเพลต ({got})",
                "docx-spec §7: รายการอ้างอิงต้องใช้ hanging indent 0.5 นิ้ว",
                "Left indent 0.5\" และ Special: Hanging 0.5\"")

        m = re.search(r"\bdoi\s*:\s*10\.", t, re.I)
        if m:
            add("major", i, pg, n, f"ใช้รูปแบบเก่า '{t[m.start():m.start()+22]}…'",
                "APA7 กำหนดให้ DOI เป็น URL เต็ม ไม่ใช้คำนำหน้า 'doi:'",
                "https://doi.org/10.xxxx/yyyy")

        m = re.search(r"(https?://doi\.org/\S+?)[\.。]\s*$", t)
        if m:
            add("minor", i, pg, n, "มีจุดปิดท้ายหลัง DOI",
                "APA7: ห้ามใส่เครื่องหมายใด ๆ ต่อท้าย DOI/URL เพราะทำให้ลิงก์เสีย",
                m.group(1))

        if re.search(r"\bRetrieved from\b", t, re.I):
            add("major", i, pg, n, "ยังใช้ 'Retrieved from'",
                "APA7 ตัด 'Retrieved from' ออกแล้ว (คงไว้เฉพาะกรณีต้องระบุวันที่สืบค้น)",
                "ใส่ URL ตรง ๆ ไม่ต้องมีคำนำหน้า")

        # ที่ตั้งสำนักพิมพ์: "เมือง, รัฐ: สำนักพิมพ์" — ห้ามข้ามจุดจบประโยค ไม่งั้นจะกิน
        # ชื่อเรื่องเข้ามาด้วยแล้วข้อความที่ยกมาให้ผู้ใช้ดูจะชี้ผิดจุด
        # ที่ตั้งสำนักพิมพ์อยู่ "หลังชื่อเรื่อง" เสมอ. ถ้าไม่บังคับจุดนี้ ชื่อเรื่อง
        # ที่ขึ้นต้นด้วยชื่อเมืองแล้วตามด้วย subtitle เช่น
        # "สมชาย ก. (2562). Bangkok, Thailand: A case study of…" จะถูกฟ้องผิด
        # ทั้งที่เป็นชื่อเรื่อง ไม่ใช่ที่ตั้งสำนักพิมพ์.
        if not thai:
            after_title = _after_title_span(t)
            m = re.search(r"(?:^|[\.\?!]\s)\s*([A-Z][a-zA-Z\-' ]{2,30},\s*(?:[A-Z]{2}|[A-Z][a-z]+))\s*:\s*[A-Z]",
                          after_title)
            if m:
                add("major", i, pg, n, f"ยังระบุที่ตั้งสำนักพิมพ์ '{m.group(1)}:'",
                    "APA7 ตัดที่ตั้งสำนักพิมพ์ออกแล้ว เหลือเฉพาะชื่อสำนักพิมพ์",
                    "…. ชื่อสำนักพิมพ์.")

        # ยัติภังค์ในช่วงเลขหน้า — ต้องไม่ไปแตะเลขที่อยู่ใน URL/DOI (เช่น
        # ".../11-2019/report.pdf" หรือ "10.1037/0000168-000") เพราะนั่นเป็น
        # ส่วนหนึ่งของที่อยู่จริง แก้เป็น en-dash แล้วลิงก์พัง — ผิดหนักกว่าเดิม
        for m in re.finditer(r"(?<![\d\-–])(\d{1,4})\s*-\s*(\d{1,4})(?![\d\-–])", t):
            if _inside_url(t, m.start()):
                continue
            if int(m.group(2)) >= int(m.group(1)):
                add("minor", i, pg, n, f"ช่วงเลขหน้า '{m.group(0)}' ใช้ยัติภังค์ (-)",
                    "APA7 กำหนดให้ช่วงเลขหน้าใช้ en-dash (–)",
                    f"{m.group(1)}–{m.group(2)}")
                break

        if thai:
            m = re.search(r"\b(pp?\.)\s*(\d[\d–\-]*)", t)
            if m:
                add("major", i, pg, n, f"รายการภาษาไทยใช้ '{m.group(1)} {m.group(2)}'",
                    "TULIBS APA7: รายการภาษาไทยใช้ 'น.' แทน p./pp. "
                    "(ยกเว้นผู้แต่งต่างชาติที่อ้างในเล่มไทย ให้คงรูปเดิม)",
                    f"น. {m.group(2)}")
            m = re.search(r"\((19|20)(\d{2})\)", t)
            if m and re.match(r"^[^A-Za-z]{0,40}[฀-๿]", t):
                yr = int(m.group(0)[1:5])
                add("major", i, pg, n, f"รายการภาษาไทยลงปีเป็น ค.ศ. ({yr})",
                    "TULIBS APA7: รายการภาษาไทยลงปีพิมพ์เป็น พ.ศ.",
                    f"({yr + 543})")

    # การเรียงอังกฤษเทียบ casefold ได้แน่นอน; ภาษาไทยต้องใช้กฎ collation เฉพาะ
    # (สระนำ/วรรณยุกต์) จึงปล่อยให้ manual pass เพื่อไม่สร้าง false positive.
    for label, group in (("อังกฤษ", [e for e in entries if not entry_is_thai(e[2])]),):
        keys = [re.sub(r"^[\"'“”\s]+", "", e[2]).lower() for e in group]
        for k in range(1, len(keys)):
            if keys[k] < keys[k-1]:
                i, pg, t = group[k]
                out.append(F("minor", "APA7 เชิงรูปแบบ",
                             f"รายการ '{t[:45]}…' อยู่ผิดลำดับตัวอักษร (มาหลัง '{group[k-1][2][:30]}…')",
                             f"APA7: รายการอ้างอิงกลุ่มภาษา{label}ต้องเรียงตามลำดับตัวอักษรของผู้แต่ง",
                             loc_of(i, pg), "จัดเรียงรายการใหม่ตามลำดับตัวอักษร", domain="apa7"))
                break            # รายงานจุดแรกที่เพี้ยนพอ ไม่ต้องรัวทุกบรรทัด

    # in-text ↔ รายการท้ายเล่ม (จับเฉพาะนามสกุลอังกฤษ + ปี ซึ่งจับคู่ได้แม่น)
    first_ref = entries[0][0]
    cited = {}
    for i, p in enumerate(paragraphs[:first_ref]):
        text = p.text or ""
        # สารบัญ/แถวตารางที่จัดด้วย tab หรือจุดไข่ปลา ไม่ใช่เนื้อความที่มีการอ้างอิง
        if "\t" in text or re.search(r"\.{3,}", text):
            continue
        for m in _INTEXT_RE.finditer(text):
            name, year = m.group(1), m.group(2)
            if name.lower() in _NOT_A_SURNAME:
                continue
            # ผู้แต่งที่เป็นองค์กรมีหลายคำ ("Baymard Institute", "LINE Ads Pricing")
            # แต่ regex จับได้แค่คำสุดท้าย. เก็บคำขึ้นต้นด้วยตัวใหญ่ที่อยู่ติดกัน
            # ก่อนหน้าไว้เป็นชื่อสำรอง ไม่งั้นตอนไปหาในรายการอ้างอิงจะหาไม่เจอ
            # แล้วฟ้องว่า "ไม่มีในรายการ" ทั้งที่มีอยู่ในชื่อเต็มขององค์กร
            span = text[max(0, m.start() - 60):m.start(1) + len(name)]
            aliases = tuple(w for w in re.findall(r"\b[A-Z][a-zA-Z\-'&]{2,}", span)
                            if w.lower() not in _NOT_A_SURNAME)
            cited.setdefault((name, year),
                             (i, pages[i] if i < len(pages) else None, aliases))
    reftext = " ||| ".join(e[2] for e in entries)

    def in_reflist(candidates, year):
        # ระยะ 160 ตัวอักษรเผื่อรายการที่มีผู้แต่งหลายคนก่อนถึงวงเล็บปี — แคบไป
        # จะฟ้องว่า "ไม่มีในรายการ" ทั้งที่มีอยู่ แค่ชื่อกับปีอยู่ห่างกัน.
        # ปีก็เทียบทั้ง ค.ศ. และ พ.ศ. เพราะเล่มไทยลงปีในรายการเป็น พ.ศ. แต่ผู้เขียน
        # มักพิมพ์ในเนื้อหาเป็น ค.ศ. ตามต้นฉบับ — คนละระบบปี แต่เป็นงานชิ้นเดียวกัน
        years = {year, str(int(year) + 543), str(int(year) - 543)}
        for cand in candidates:
            for y in years:
                # ไม่สนตัวพิมพ์ใหญ่-เล็ก: เนื้อหาพิมพ์ "PCMI" แต่รายการพิมพ์ "Pcmi."
                # เป็นเรื่องการสะกดชื่อ ไม่ใช่ "ไม่มีรายการ" — คนละปัญหากันคนละความรุนแรง
                if re.search(re.escape(cand) + r"[^|]{0,160}?" + y, reftext, re.I):
                    return True
        return False

    for (name, year), rec in sorted(cited.items()):
        i, pg = rec[0], rec[1]
        aliases = rec[2] if len(rec) > 2 else ()
        if not in_reflist((name,) + tuple(aliases), year):
            out.append(F("critical", "APA7 เชิงรูปแบบ",
                         f"อ้าง '{name} ({year})' ในเนื้อหา แต่ไม่พบในรายการอ้างอิงท้ายเล่ม",
                         "APA7: ทุกการอ้างอิงในเนื้อหาต้องมีรายการท้ายเล่มตรงกัน "
                         "(ถ้าผู้แต่งเป็นองค์กร ให้ตรวจว่าชื่อในเนื้อหากับในรายการสะกดตรงกัน)",
                         loc_of(i, pg), f"เพิ่มรายการของ {name} ({year}) ในรายการอ้างอิง",
                         domain="apa7"))


# ---------------------------------------------------------------------------
# TEMPLATE — ตารางส่วนนำ (ค่าจริงจาก "ตัวอย่างการพิมพ์ส่วนประกอบของวิทยานิพนธ์")
#
# ทำไมย้ายมาเป็นสคริปต์: เดิมสามอย่างนี้อยู่ในรายการ "ต้องเปิดอ่านเอง" ซึ่งเป็นงานที่
# ทำแล้วเบื่อ ข้ามง่าย และผลไม่คงที่ระหว่างรอบ ทั้งที่จริง ๆ เป็นการเทียบว่า "มีป้าย
# ชื่อรายการครบไหม" — เทียบข้อความล้วน ไม่ต้องตีความ. ป้ายทั้งหมดด้านล่างคัดมาจาก
# ไฟล์ตัวอย่างทางการของหอสมุด ไม่ได้แต่งเอง
# ---------------------------------------------------------------------------
ABSTRACT_LABELS_TH = [
    ("หัวข้อ", ("หัวข้อวิทยานิพนธ์", "หัวข้อสารนิพนธ์", "หัวข้อการค้นคว้าอิสระ",
                "หัวข้อดุษฎีนิพนธ์", "หัวข้อวิทยานิพนธ์/สารนิพนธ์")),
    ("ชื่อผู้เขียน", ("ชื่อผู้เขียน",)),
    ("ชื่อปริญญา", ("ชื่อปริญญา",)),
    ("สาขาวิชา/คณะ/มหาวิทยาลัย", ("สาขาวิชา",)),
    ("อาจารย์ที่ปรึกษา", ("อาจารย์ที่ปรึกษา",)),
    ("ปีการศึกษา", ("ปีการศึกษา",)),
]
ABSTRACT_LABELS_EN = [
    ("Thesis/Dissertation Title", ("thesis title", "dissertation title",
                                   "independent study title", "title")),
    ("Author", ("author",)),
    ("Degree", ("degree",)),
    # ป้ายนี้เขียนได้หลายแบบจริงในเล่มที่ผ่านการตรวจ: "Department/Faculty/University",
    # "Major Field Faculty University", "Faculty University" — สาระคือมีคณะ/มหาวิทยาลัย
    ("Faculty/University", ("department", "faculty", "major field")),
    ("Advisor", ("advisor", "adviser")),
    ("Academic Year", ("academic year",)),
]
# ผลงานทางวิชาการ/รางวัลเป็นของที่ "มีก็ใส่" ไม่ใช่ของบังคับ จึงไม่อยู่ในรายการที่เรียกร้อง
BIO_LABELS = [("ชื่อ", ("ชื่อ", "name")),
              ("วุฒิการศึกษา", ("วุฒิการศึกษา", "educational", "education"))]
DEAN = re.compile(r"คณบดี|\bDean\b", re.I)
CHAIR = re.compile(r"ประธานกรรมการ|\bChairman\b|\bChairperson\b", re.I)


def _table_label_map(table):
    """คืน {ข้อความคอลัมน์แรก(lower)} ของทุกแถว — ใช้เทียบว่ามีป้ายรายการไหนบ้าง"""
    labels = []
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        labels.append(_squash(cells[0].text))
    return labels


def _squash(text):
    return re.sub(r"\s+", " ", (text or "")).strip()


def _has_heading(doc, *patterns):
    """คืนข้อความหัวข้อที่เจอ (truthy) เพื่อให้ finding ยกไปเป็นคำค้นใน Ctrl+F ได้

    ถ้าคืนแค่ True/False ผู้อ่านรายงานจะเจอ finding ที่บอกว่า "หน้าบทคัดย่อไม่มีตาราง"
    แต่ไม่มีอะไรให้ค้นเพื่อไปถึงหน้านั้น
    """
    pats = [re.compile(p, re.I) for p in patterns]
    for p in doc.paragraphs:
        t = _squash(p.text)
        if t and len(t) <= 60 and any(rx.match(t) for rx in pats):
            return t
    return ""


def _find_label_table(doc, spec, min_hits=3):
    """หาตารางที่คอลัมน์แรกเป็นป้ายรายการตาม spec — คืน (table, ป้ายที่ขาด)"""
    best = None
    for table in doc.tables:
        labels = " ||| ".join(_table_label_map(table)).lower()
        missing = [name for name, keys in spec
                   if not any(k.lower() in labels for k in keys)]
        hits = len(spec) - len(missing)
        if hits >= min_hits and (best is None or hits > best[0]):
            best = (hits, table, missing)
    return (best[1], best[2]) if best else (None, None)


# --- องค์ประกอบบรรทัดบนหน้าปกและหน้าอนุมัติ ---------------------------------
# ข้อความทุกบรรทัดคัดจากไฟล์ตัวอย่างทางการ ไม่ได้แต่งเอง. ตรวจแค่ "มีบรรทัดนี้ไหม"
# ไม่ตรวจการจัดกึ่งกลาง เพราะไฟล์ตัวอย่างเองก็จัดกึ่งกลางด้วยวิธีต่างกัน (บาง
# บรรทัดใช้ jc=center บางบรรทัดเว้นวรรคดัน) อ่านจาก OOXML แล้วแยกไม่ออกว่าอันไหนผิด
#
# ⚠ อย่าเติมบรรทัด "ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์" / "COPYRIGHT OF THAMMASAT
# UNIVERSITY" กลับเข้ามา — เทมเพลตทางการทั้งสามไฟล์ (Thai rev.2024, English rev.2024,
# English Times rev.2023) **ไม่มีบรรทัดนี้เลย** บรรทัดนี้อยู่ในเอกสารตัวอย่างรุ่นเก่า
# เท่านั้น เคยใส่เป็นกฎแล้วฟ้องผิดครบทั้ง 8 เล่มที่วัด
COVER_TH = [
    ("บรรทัด “โดย” ก่อนชื่อผู้เขียน", r"^\s*โดย\s*$"),
    ("บรรทัด “…นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร…”", r"เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร"),
    ("ชื่อมหาวิทยาลัยธรรมศาสตร์", r"มหาวิทยาลัยธรรมศาสตร์"),
    ("บรรทัด “ปีการศึกษา …”", r"ปีการศึกษา"),
]
COVER_EN = [
    ("บรรทัด “BY” ก่อนชื่อผู้เขียน", r"^\s*BY\s*$"),
    ("บรรทัด “A THESIS/DISSERTATION/AN INDEPENDENT STUDY SUBMITTED IN PARTIAL FULFILLMENT…”",
     r"SUBMITTED IN PARTIAL FULFILLMENT"),
    ("บรรทัด “THAMMASAT UNIVERSITY”", r"THAMMASAT UNIVERSITY"),
    ("บรรทัด “ACADEMIC YEAR …”", r"ACADEMIC YEAR"),
]

# --- ประเภทงาน: เทมเพลตเดียวกัน เปลี่ยนแค่คำเรียก ---------------------------
# ยืนยันจากไฟล์ตัวอย่างจริงที่หอสมุดใช้ — โครงหน้าปก ขนาดฟอนต์ และลำดับบรรทัด
# เหมือนเทมเพลตวิทยานิพนธ์ทุกประการ ต่างแค่คำว่า วิทยานิพนธ์ / สารนิพนธ์ /
# การค้นคว้าอิสระ (และฝั่งอังกฤษ THESIS / DISSERTATION / INDEPENDENT STUDY)
#
# ทำไมต้องรู้ประเภท: รายงานต้องพูดคำเดียวกับที่อยู่ในเล่ม ถ้าเล่ม IS แล้วรายงาน
# เขียนว่า "วิทยานิพนธ์นี้เป็นส่วนหนึ่ง…" ผู้อ่านจะหาไม่เจอและคิดว่ารายงานผิดเล่ม
WORK_TYPES = [
    ("is", "การค้นคว้าอิสระ", "AN INDEPENDENT STUDY",
     r"การค้นคว้าอิสระ|Independent Study", r"INDEPENDENT STUDY"),
    ("term-paper", "สารนิพนธ์", "A THEMATIC PAPER",
     r"สารนิพนธ์", r"THEMATIC PAPER"),
    ("dissertation", "ดุษฎีนิพนธ์/วิทยานิพนธ์ (ปริญญาเอก)", "A DISSERTATION",
     r"ดุษฎีนิพนธ์", r"\bDISSERTATION\b"),
    ("thesis", "วิทยานิพนธ์", "A THESIS", r"วิทยานิพนธ์", r"\bTHESIS\b"),
]


def detect_work_type(blob):
    """คืน (รหัส, คำไทย, คำอังกฤษ) ของประเภทงานจากข้อความส่วนหน้า

    ไล่จากคำที่เฉพาะเจาะจงที่สุดก่อน เพราะ "การค้นคว้าอิสระ" กับ "สารนิพนธ์"
    เป็นคำเฉพาะ ส่วน "วิทยานิพนธ์" เป็นคำกลางที่โผล่ในเอกสารอื่นได้ง่ายกว่า
    """
    for code, th, en, rx_th, rx_en in WORK_TYPES:
        if re.search(rx_th, blob) or re.search(rx_en, blob, re.I):
            return code, th, en
    return "thesis", "วิทยานิพนธ์", "A THESIS"
APPROVAL_LINE = re.compile(
    r"ได้รับการตรวจสอบและอนุมัติ|was approved as partial fulfillment|has been approved",
    re.I)

# บรรทัดที่ "เคยมี" ในเทมเพลตรุ่นก่อน แต่ **ถูกตัดออกแล้ว** ใน rev.2024/rev.2023
#
# ตรวจสองทิศทางถึงจะครบ: ขาดบรรทัดที่ต้องมี = ปกไม่ครบ · มีบรรทัดที่เทมเพลตตัดออกแล้ว
# = ปกไม่ตรงรุ่น (มักเกิดจากก๊อปปกจากเล่มรุ่นพี่มาใช้). เดิมตรวจแค่ทิศทางแรก
# เล่มที่แปะบรรทัดเก่าไว้จึงผ่านฉลุย
#
# เพิ่มรายการใหม่ได้เมื่อยืนยันแล้วว่า **ไม่มีใน references/templates/ ทั้ง 3 ไฟล์**
# (ใช้ grep นับให้ได้ 0 ทั้งสามไฟล์ก่อนเสมอ)
COVER_DEPRECATED = [
    (r"ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์", "ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์"),
    # ⚠ ห้ามใส่ "COPYRIGHT OF THAMMASAT UNIVERSITY" กลับเข้ามา — ไฟล์ตัวอย่างภาษา
    # อังกฤษที่หอสมุดใช้ (ดุษฎีนิพนธ์) มีบรรทัดนี้อยู่จริงบนหน้าปก ต่างจากปกไทย
    # ที่ไม่มี. ฟ้องเมื่อไรคือฟ้องตัวอย่างที่ถูกต้อง
]
FRONT_STOP = re.compile(r"^\s*(บทคัดย่อ|ABSTRACT)\b", re.I)


def front_matter_lines(doc):
    """[(ดัชนีย่อหน้า, ข้อความหนึ่งบรรทัด)] ของส่วนหน้า (ก่อนบทคัดย่อ) รวมข้อความในตาราง

    แยกทีละ **บรรทัด** ไม่ใช่ทีละย่อหน้า เพราะบล็อกบนปกหลายอันเป็นย่อหน้าเดียวที่ขึ้น
    บรรทัดใหม่ด้วย Shift+Enter — ถ้าคืนทั้งย่อหน้า คำค้นที่ยกไปให้ผู้อ่านจะยาวเกินและ
    ค้นใน Word ไม่เจอเพราะคร่อมบรรทัด
    """
    out = []
    for i, para in enumerate(doc.paragraphs):
        text = _squash(para.text)
        if FRONT_STOP.match(text):
            break
        for line in re.split(r"[\r\n\v]", para.text or ""):
            line = _squash(line)
            if line:
                out.append((i, line))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in re.split(r"[\r\n\v]", cell.text or ""):
                    line = _squash(line)
                    if line:
                        out.append((None, line))
    return out


def front_matter_text(doc):
    """ข้อความส่วนหน้าทั้งก้อน — ใช้เทียบว่ามีบรรทัดที่ต้องมีครบไหม"""
    return "\n".join(line for _, line in front_matter_lines(doc))


def check_cover_elements(doc, pages, out):
    blob = front_matter_text(doc)
    if not blob:
        return
    # "มีปกภาษานี้ไหม" ตัดสินจากจำนวนบรรทัดเฉพาะของปกนั้นที่เจอ ไม่ใช่จากภาษาของ
    # ตัวอักษรที่โผล่มา — ไฟล์ปกอังกฤษล้วนก็มีอักษรไทยได้ (เช่นหัว "ภาคผนวก ค")
    # ถ้าใช้แค่ "มีอักษรไทย" จะไปเรียกร้องปกไทยจากไฟล์ที่ไม่มีปกไทยตั้งแต่แรก
    def found(spec):
        return [name for name, rx in spec if re.search(rx, blob, re.M | re.I)]

    th_found, en_found = found(COVER_TH), found(COVER_EN)
    has_th, has_en = len(th_found) >= 2, len(en_found) >= 2
    if not (has_th or has_en):
        return          # ไม่มีหน้าปกในไฟล์นี้ (เช่นไฟล์ที่ตัดมาเฉพาะบางส่วน) — ไม่เรียกร้อง

    code, word_th, word_en = detect_work_type(blob)

    for label, spec, active, seen in (("ไทย", COVER_TH, has_th, th_found),
                                      ("อังกฤษ", COVER_EN, has_en, en_found)):
        if not active:
            continue
        missing = [name for name, _ in spec if name not in seen]
        for name in missing:
            # แทนคำกลางในข้อความด้วยคำจริงของเล่มนี้ ผู้อ่านจะได้ค้นเจอ
            shown = name.replace("…นี้เป็น", f"{word_th}นี้เป็น").replace(
                "A THESIS/DISSERTATION/AN INDEPENDENT STUDY", word_en)
            out.append(F("major", "หน้าปก",
                         f"หน้าปก{label}ยังไม่มี{shown}",
                         f"เทมเพลตหน้าปกของ{word_th}ใช้โครงเดียวกับวิทยานิพนธ์ "
                         "ต่างแค่คำเรียกประเภทงาน — ขาดบรรทัดใดบรรทัดหนึ่งถือว่าปกไม่ครบ",
                         "หน้าปกใน", f"เพิ่ม{shown}"))

    # --- บรรทัดที่เทมเพลตรุ่นปัจจุบันตัดออกแล้ว แต่ยังค้างอยู่ในเล่ม -----------
    seen_dep = set()
    for para_index, line in front_matter_lines(doc):
        for pattern, canonical in COVER_DEPRECATED:
            if canonical in seen_dep or not re.search(pattern, line, re.I):
                continue
            # ต้องเป็น "บรรทัดนั้นทั้งบรรทัด" ไม่ใช่ข้อความที่พูดถึงบรรทัดนี้ในประโยคยาว
            # (เอกสารคู่มือ/รายงานที่อ้างถึงกฎข้อนี้ต้องไม่ถูกนับว่าเป็นหน้าปกที่ผิด)
            if len(line) > len(canonical) + 15:
                continue
            seen_dep.add(canonical)
            out.append(F("major", "หน้าปก",
                         f"หน้าปกมีบรรทัด “{line[:60]}” ซึ่งเทมเพลตรุ่นปัจจุบันตัดออกแล้ว",
                         "เทมเพลตทางการ rev.2024 (ไทย/อังกฤษ) และ rev.2023 (English Times) "
                         f"ไม่มีบรรทัด “{canonical}” บนหน้าปก — บรรทัดนี้เป็นของเทมเพลตรุ่นเก่า "
                         "มักติดมาจากการก๊อปหน้าปกจากเล่มรุ่นก่อน",
                         f"para {para_index}" if para_index is not None else "หน้าปกใน",
                         f"ลบบรรทัด “{canonical}” ออกจากหน้าปก "
                         "(ถ้าคณะ/หลักสูตรกำหนดให้มี ให้ยืนยันกับหอสมุดก่อน)"))

    if not APPROVAL_LINE.search(blob):
        out.append(F("major", "หน้าอนุมัติ",
                     "ไม่พบข้อความรับรองบนหน้าอนุมัติ",
                     "หน้าอนุมัติต้องมีประโยครับรอง เช่น “ได้รับการตรวจสอบและอนุมัติ "
                     "ให้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร…” หรือฉบับอังกฤษ "
                     "“…was approved as partial fulfillment of the requirements for…”",
                     "หน้าอนุมัติ", "เพิ่มประโยครับรองตามเทมเพลต"))

    # ชื่อประเภทงานต้องใช้คำเดียวกันทั้งเล่ม — เล่มที่ดัดแปลงจากเทมเพลตวิทยานิพนธ์
    # มักแก้คำบนปกแล้วลืมแก้บนหน้าอนุมัติหรือในตารางบทคัดย่อ กลายเป็นคนละคำในเล่มเดียว
    if code != "thesis":
        others = [w for c, w, _, _, _ in WORK_TYPES if c != code and w in blob]
        if "วิทยานิพนธ์" in blob and code in ("is", "term-paper"):
            others.append("วิทยานิพนธ์")
        others = sorted({w for w in others if w != word_th})
        if others:
            out.append(F("major", "หน้าปก",
                         f"เล่มนี้เป็น{word_th} แต่ส่วนหน้ายังมีคำว่า “{others[0]}” ปนอยู่",
                         f"เทมเพลตของ{word_th}ดัดแปลงจากเทมเพลตวิทยานิพนธ์ "
                         "ถ้าแก้คำไม่ครบทุกจุด (ปกใน · ปกภาษาที่สอง · หน้าอนุมัติ · "
                         "ตารางข้อมูลบทคัดย่อ) จะกลายเป็นคนละคำในเล่มเดียวกัน",
                         "ส่วนหน้าของเล่ม",
                         f"แก้คำว่า “{others[0]}” ให้เป็น “{word_th}” ให้ครบทุกจุดในส่วนหน้า"))


# --- บรรทัดว่างคั่นกลางย่อหน้าเนื้อความ --------------------------------------
CAPTION_LABEL = re.compile(r"^\s*(ตารางที่|ภาพที่|รูปที่|Table|Figure)\s*[\d๐-๙]+([.\-][\d๐-๙]+)?\s*$", re.I)
_PROSE_MIN = 120
_NOT_PROSE = re.compile(
    r"^\s*(\d+[\.\)]|\(\d+\)|บทที่|CHAPTER|ตารางที่|ภาพที่|รูปที่|Table|Figure|"
    r"คำสำคัญ|คำสําคัญ|Keywords?|ที่มา|Source|หมายเหตุ|Note)\b", re.I)


def _is_prose(text):
    if len(text) < _PROSE_MIN or _NOT_PROSE.match(text):
        return False
    # บรรทัดตัวพิมพ์ใหญ่ล้วนคือปก/หัวข้อ ไม่ใช่ร้อยแก้ว — ใช้ได้เฉพาะข้อความละติน
    # (ภาษาไทยไม่มีตัวพิมพ์ใหญ่-เล็ก ถ้าไม่กันไว้ ย่อหน้าไทยล้วนจะถูกตัดทิ้งทั้งหมด
    # แล้วกฎนี้จะจับได้เฉพาะย่อหน้าที่บังเอิญมีอักษรอังกฤษปน)
    latin = [c for c in text if c.isalpha() and c.isascii()]
    if len(latin) >= 8 and all(c.isupper() for c in latin):
        return False
    return True


def check_blank_lines_in_prose(doc, pages, out):
    """บรรทัดว่างคั่นระหว่างย่อหน้าเนื้อความสองย่อหน้า

    เทมเพลตกำหนดให้ย่อหน้าเนื้อความติดกันโดยไม่เว้นบรรทัด (spacing before/after = 0)
    ระยะห่างที่เห็นมาจากการเยื้องบรรทัดแรก ไม่ใช่การเคาะ Enter. การเคาะ Enter คั่น
    ทำให้ระยะห่างไม่เท่ากันทั้งเล่มและบรรทัดว่างจะไปโผล่หัวหน้าถัดไปเมื่อเนื้อหาเลื่อน.

    ตรวจเฉพาะช่วงเนื้อหา (หลัง "บทที่ 1") และเฉพาะจุดที่ **ทั้งบนและล่างเป็นร้อยแก้ว**
    — บรรทัดว่างก่อนหัวข้อ ก่อน "คำสำคัญ:" ก่อนตาราง/ภาพ หรือบนหน้าปก เป็นการจัดหน้า
    ปกติที่ไฟล์ตัวอย่างทางการก็ทำ จึงต้องไม่ถูกจับ
    """
    texts = [_squash(p.text) for p in doc.paragraphs]
    start = next((i for i, t in enumerate(texts)
                  if re.match(r"^(บทที่\s*1(?:\s|$)|CHAPTER\s*(?:1|I)(?:\s|$))", t, re.I)), None)
    if start is None:
        return
    hits = []
    i = start + 1
    while i < len(texts) - 1:
        if texts[i]:
            i += 1
            continue
        j = i
        while j < len(texts) and not texts[j]:
            j += 1
        # caption ในเล่มจริงมักแยกสองย่อหน้า: "ภาพที่ 4.13" แล้วชื่อภาพยาว ๆ ต่อบรรทัดถัดมา
        # บรรทัดที่สองยาวเกิน 120 ตัวอักษรจึงถูกนับเป็นร้อยแก้วทั้งที่เป็นคำบรรยายภาพ
        # ทำให้บรรทัดว่างหลัง caption (ซึ่งเป็นการจัดหน้าปกติ) ถูกฟ้อง — ต้องมองย้อนขึ้นไป
        after_caption = i >= 2 and CAPTION_LABEL.match(texts[i - 2])
        if (j < len(texts) and not after_caption
                and _is_prose(texts[i - 1]) and _is_prose(texts[j])):
            hits.append((i, texts[i - 1][-40:], texts[j][:40]))
        i = j
    if not hits:
        return
    first_i, before, after = hits[0]
    extra = f" (พบทั้งหมด {len(hits)} จุดในเล่ม)" if len(hits) > 1 else ""
    out.append(F("minor", "การเว้นบรรทัด",
                 f"มีบรรทัดว่างคั่นกลางระหว่างย่อหน้าเนื้อความ — ระหว่าง “…{before}” "
                 f"กับ “{after}…”{extra}",
                 "เทมเพลต TULIBS: ย่อหน้าเนื้อความพิมพ์ต่อกันโดยไม่เคาะ Enter เว้นบรรทัด "
                 "(ระยะห่างก่อน/หลังย่อหน้า = 0) ระยะที่เห็นมาจากการเยื้องบรรทัดแรก 0.80 นิ้ว",
                 f"para {first_i}",
                 "ลบบรรทัดว่างระหว่างย่อหน้าออก แล้วให้ระยะห่างมาจากการเยื้องบรรทัดแรกแทน"))


def check_front_matter_tables(doc, pages, out):
    def add(sev, issue, detail, loc, correct):
        out.append(F(sev, "ส่วนนำ", issue, detail, loc, correct))

    # --- ตารางข้อมูลบทคัดย่อ ---------------------------------------------
    for lang, spec, head in (
            ("ไทย", ABSTRACT_LABELS_TH, (r"บทคัดย่อ",)),
            ("อังกฤษ", ABSTRACT_LABELS_EN, (r"ABSTRACT\b",))):
        heading = _has_heading(doc, *head)
        if not heading:
            continue          # ไม่มีบทคัดย่อภาษานี้ ก็ไม่ต้องเรียกร้องตารางของมัน
        at = f"หน้าบทคัดย่อภาษา{lang} · ค้นด้วย “{heading}”"
        table, missing = _find_label_table(doc, spec)
        if table is None:
            add("major", f"ไม่พบตารางข้อมูลหน้าบทคัดย่อภาษา{lang}",
                "เทมเพลต TULIBS กำหนดให้หน้าบทคัดย่อเริ่มด้วยตารางข้อมูล "
                + " · ".join(n for n, _ in spec),
                at,
                "เพิ่มตาราง 2 คอลัมน์ ป้ายรายการอยู่คอลัมน์ซ้าย ข้อมูลอยู่คอลัมน์ขวา")
        elif missing:
            add("major", f"ตารางข้อมูลบทคัดย่อภาษา{lang}ขาดรายการ: {', '.join(missing)}",
                "เทมเพลต TULIBS กำหนดให้มีครบ " + " · ".join(n for n, _ in spec),
                at, "เพิ่มแถวที่ขาดในตารางข้อมูลบทคัดย่อ")

    # --- หน้าอนุมัติ -------------------------------------------------------
    approval = _has_heading(doc, r"(ได้รับการตรวจสอบและอนุมัติ|has been approved)")
    if approval:
        at_approval = f"หน้าอนุมัติ · ค้นด้วย “{approval}”"
        committee = None
        for table in doc.tables:
            labels = " ".join(_table_label_map(table))
            if CHAIR.search(labels) or DEAN.search(labels):
                committee = table
                break
        if committee is None:
            add("major", "หน้าอนุมัติไม่มีตารางรายชื่อกรรมการ",
                "เทมเพลต TULIBS: หน้าอนุมัติมีตารางตำแหน่งกรรมการคู่กับชื่อในวงเล็บ "
                "โดยมีประธานกรรมการสอบและคณบดีเสมอ",
                at_approval, "เพิ่มตารางรายชื่อกรรมการตามเทมเพลต")
        else:
            text = " ".join(c.text for r in committee.rows for c in r.cells)
            if not CHAIR.search(text):
                add("major", "หน้าอนุมัติไม่มีบรรทัดประธานกรรมการสอบ",
                    "เทมเพลต TULIBS: ต้องมีประธานกรรมการสอบเสมอ",
                    at_approval, "เพิ่มบรรทัด 'ประธานกรรมการสอบวิทยานิพนธ์'")
            if not DEAN.search(text):
                add("major", "หน้าอนุมัติไม่มีบรรทัดคณบดี",
                    "เทมเพลต TULIBS: บรรทัดสุดท้ายของหน้าอนุมัติคือคณบดี",
                    at_approval, "เพิ่มบรรทัด 'คณบดี' พร้อมชื่อในวงเล็บ")
            # ชื่อกรรมการต้องอยู่ในวงเล็บ — ตรวจเฉพาะแถวที่มีข้อความในคอลัมน์ขวาจริง
            unbracketed = []
            for row in committee.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                name = _squash(cells[-1].text)
                if name and not re.search(r"\(.+\)", name):
                    unbracketed.append(name[:40])
            if unbracketed:
                add("minor", f"ชื่อกรรมการไม่ได้อยู่ในวงเล็บ: {', '.join(unbracketed[:3])}",
                    "เทมเพลต TULIBS: ชื่อใต้เส้นลายเซ็นพิมพ์ในวงเล็บ เช่น (ศาสตราจารย์ ดร. ก ข)",
                    at_approval, "ใส่วงเล็บครอบชื่อกรรมการทุกคน")

    # --- ประวัติผู้เขียน ---------------------------------------------------
    # ไฟล์ตัวอย่างทางการเองก็มีทั้งแบบตารางและแบบย่อหน้าธรรมดา จึง**ไม่**ฟ้องเรื่อง
    # "ต้องเป็นตาราง" — ฟ้องเฉพาะกรณีที่จัดเป็นตารางแล้วแต่ขาดข้อมูลที่ต้องมี
    if _has_heading(doc, r"ประวัติผู้เขียน", r"BIOGRAPHY"):
        table, missing = _find_label_table(doc, BIO_LABELS, min_hits=1)
        if table is not None and missing:
            add("minor", f"ประวัติผู้เขียนขาดรายการ: {', '.join(missing)}",
                "เทมเพลต TULIBS: ประวัติผู้เขียนต้องมีชื่อและวุฒิการศึกษาเป็นอย่างน้อย "
                "(ผลงานทางวิชาการ/รางวัล ใส่เท่าที่มี)",
                "ประวัติผู้เขียน", "เพิ่มแถวที่ขาด")


# ---------------------------------------------------------------------------
# TEMPLATE — บันไดการเยื้องของหัวข้อย่อย
#
# ไฟล์ "ตัวอย่างการพิมพ์และจัดวางเนื้อหาวิทยานิพนธ์" วางบันไดไว้ชัดเจน:
# หัวข้อใหญ่ X.X ชิดขอบ · X.X.X เยื้อง 0.80" · X.X.X.X เยื้อง 1.10" · (1) เยื้อง 1.40"
#
# แต่**ไม่**เทียบกับค่าสัมบูรณ์เหล่านี้ตรง ๆ เพราะเทมเพลตแต่ละรุ่นใช้ค่าต่างกันเล็กน้อย
# (1.10" กับ 1.19" มีอยู่จริงทั้งคู่ในตระกูลสไตล์ TU_*) การบังคับค่าเดียวจะกลายเป็นการ
# ไล่ให้นักศึกษาแก้สิ่งที่เทมเพลตให้มาเอง. สิ่งที่ผิดแน่นอนคือ **ความไม่สม่ำเสมอภายใน
# เล่มเดียวกัน** — หัวข้อระดับเดียวกันเยื้องไม่เท่ากัน หรือหัวข้อย่อยเยื้องน้อยกว่า
# หัวข้อแม่ ซึ่งคนอ่านเห็นทันทีว่าลำดับชั้นเพี้ยน
# ---------------------------------------------------------------------------
NUMBERED_HEAD = re.compile(r"^(\d{1,2})((?:\.\d{1,2}){1,3})\s+\S")
_INDENT_TOL = 0.06          # นิ้ว — เผื่อการปัดเศษ twips ของ Word


def _indent_of(paragraph):
    total = 0.0
    for attr in ("first_line_indent", "left_indent"):
        value = _effective_para_value(paragraph, attr)
        if value is not None:
            total += value.inches
    return round(total, 2)


# แม่แบบหัวข้อของเทมเพลตทางการ: เลขข้อแบบไหน ต้องใช้สไตล์ไหน
# (ดูส่วน "เนื้อหาทั้งหมดของเทมเพลต" ใน references/templates/*.md)
#   1.1        → TU_Main Heading _ChapterN   (16pt หนา · เยื้อง 0" ไทย / 0.25" อังกฤษ)
#   1.1.1      → TU_Sub-heading 1            (16pt หนา · เยื้องบรรทัดแรก 0.80")
#   1.1.1.1    → TU_Sub-heading 2            (16pt หนา · เยื้องบรรทัดแรก 1.10")
#   (1)        → TU_Sub-heading 3            (16pt หนา · เยื้องบรรทัดแรก 1.40")
HEADING_STYLE_BY_LEVEL = {2: "TU_Main Heading", 3: "TU_Sub-heading 1",
                          4: "TU_Sub-heading 2"}


def check_heading_styles(doc, pages, out):
    """หัวข้อที่มีเลขข้อ ต้องใช้สไตล์หัวข้อของเทมเพลตให้ตรงระดับ

    ทำไมสำคัญกว่าการไล่ตั้งขนาด/เยื้องเอง: สไตล์คือสิ่งที่ทำให้ทั้งเล่มเหมือนกัน
    และเป็นตัวสร้างสารบัญอัตโนมัติ. หัวข้อที่พิมพ์ด้วย Normal แล้วจัดขนาดเอง จะไม่
    ขึ้นในสารบัญและจะเพี้ยนทันทีที่แก้ที่อื่น — ตรวจจากชื่อสไตล์จึงตรงต้นเหตุกว่า
    ตรวจค่าที่มองเห็น.

    ตรวจเฉพาะเล่มที่ **มีสไตล์ TU_ อยู่แล้ว** เพราะเล่มที่ไม่มีเลยถูกรายงานไปแล้วว่า
    ไม่ได้สร้างจากเทมเพลต การฟ้องซ้ำทีละหัวข้อไม่ได้เพิ่มข้อมูลอะไร
    """
    names = {s.name for s in doc.styles}
    if "TU_Sub-heading 1" not in names:
        return
    rows, totals = [], 0
    for i, para in enumerate(doc.paragraphs):
        raw = para.text or ""
        text = raw.strip()
        if not text or "\t" in raw or re.search(r"\.{3,}", text) or len(text) > 120:
            continue
        m = NUMBERED_HEAD.match(text)
        if not m:
            continue
        level = 1 + m.group(2).count(".")
        want = HEADING_STYLE_BY_LEVEL.get(level)
        if not want:
            continue
        totals += 1
        style = para.style.name if para.style else ""
        ok = style.startswith(want) if want == "TU_Main Heading" else style == want
        if not ok:
            rows.append((i, level, text, style or "(ไม่ระบุ)", want))
    if not rows:
        return

    label = {2: "หัวข้อใหญ่ (X.X)", 3: "หัวข้อย่อยระดับ 1 (X.X.X)",
             4: "หัวข้อย่อยระดับ 2 (X.X.X.X)"}
    if len(rows) / max(totals, 1) >= 0.6:
        # ทั้งเล่มไม่ได้ใช้สไตล์หัวข้อของเทมเพลต — เป็นปัญหาเดียว ไม่ใช่ 88 ปัญหา
        used = sorted({r[3] for r in rows})[:4]
        i, level, text, style, want = rows[0]
        out.append(F("major", "สไตล์หัวข้อ",
                     f"หัวข้อที่มีเลขข้อในเล่มนี้ {len(rows)} จาก {totals} จุด ไม่ได้ใช้สไตล์หัวข้อ "
                     f"ของเทมเพลต (ที่ใช้อยู่: {', '.join(used)}) เช่น “{text[:45]}”",
                     "เทมเพลต TULIBS กำหนดสไตล์ตามระดับหัวข้อ: X.X = TU_Main Heading_ChapterN · "
                     "X.X.X = TU_Sub-heading 1 · X.X.X.X = TU_Sub-heading 2 · (1) = TU_Sub-heading 3 "
                     "— หัวข้อที่ไม่ได้ใช้สไตล์เหล่านี้จะไม่ขึ้นในสารบัญอัตโนมัติ",
                     f"para {i}",
                     "ไล่คลิกที่หัวข้อแต่ละอันแล้วเลือกสไตล์ให้ตรงระดับจากกล่อง Styles"))
        return
    for i, level, text, style, want in rows[:8]:
        out.append(F("major", "สไตล์หัวข้อ",
                     f"{label[level]} “{text[:45]}” ใช้สไตล์ {style} ซึ่งไม่ใช่สไตล์ของระดับนี้",
                     f"เทมเพลต TULIBS: หัวข้อระดับนี้ต้องใช้สไตล์ {want} "
                     "— ใช้ผิดสไตล์ทำให้ขนาด/การเยื้องไม่ตรงกับหัวข้ออื่นระดับเดียวกัน "
                     "และสารบัญอัตโนมัติจะเก็บระดับผิด",
                     f"para {i}", f"เปลี่ยนสไตล์ของย่อหน้านี้เป็น {want}"))
    if len(rows) > 8:
        out.append(F("major", "สไตล์หัวข้อ",
                     f"ยังมีหัวข้อที่ใช้สไตล์ไม่ตรงระดับอีก {len(rows) - 8} จุด",
                     "รายงานเฉพาะ 8 จุดแรกเพื่อไม่ให้ตารางยาวเกินจำเป็น",
                     f"para {rows[8][0]}", "ตรวจสไตล์ของหัวข้อที่เหลือให้ตรงระดับ"))


def check_heading_indent_ladder(doc, pages, out):
    # เล่มที่ใช้สไตล์ TU_ อยู่แล้ว ให้ check_heading_styles ตรวจแทน — ตรงต้นเหตุกว่า
    # และไม่ต้องรายงานสองแถวสำหรับหัวข้อเดียวกัน
    if "TU_Sub-heading 1" in {s.name for s in doc.styles}:
        return
    levels = {}
    for i, para in enumerate(doc.paragraphs):
        raw = para.text or ""
        text = raw.strip()
        # ข้ามสารบัญ (tab/จุดไข่ปลา) และย่อหน้ายาวที่ขึ้นต้นด้วยเลขข้อแต่เป็นเนื้อความ
        if not text or "\t" in raw or re.search(r"\.{3,}", text) or len(text) > 120:
            continue
        m = NUMBERED_HEAD.match(text)
        if not m:
            continue
        levels.setdefault(1 + m.group(2).count("."), []).append(
            (i, pages[i] if i < len(pages) else None, _indent_of(para), text))

    modes = {}
    for level, items in sorted(levels.items()):
        if len(items) < 3:
            continue        # น้อยเกินกว่าจะรู้ว่าอันไหนคือรูปแบบหลักของเล่ม
        # จัดกลุ่มค่าที่ต่างกันแค่เศษการปัด twips ให้เป็นค่าเดียวกันก่อนนับ — ไม่งั้น
        # 0.45"/0.49"/0.5" ซึ่งตาคนมองว่าเท่ากันจะถูกนับเป็นสามค่าแล้วรายงานว่าไม่สม่ำเสมอ
        counter = {}
        for _, _, value, _ in items:
            bucket = next((b for b in counter if abs(b - value) <= _INDENT_TOL), value)
            counter[bucket] = counter.get(bucket, 0) + 1
        top, hits = max(counter.items(), key=lambda kv: (kv[1], -kv[0]))
        share = hits / len(items)
        if share < 0.6:
            values = ", ".join(f'{v}"' for v in sorted(counter))
            out.append(F("minor", "การเยื้อง",
                         f"หัวข้อระดับ {level} ({'.'.join(['X'] * level)}) เยื้องไม่สม่ำเสมอ "
                         f"— ใช้อยู่ {len(counter)} ค่า: {values}",
                         "เทมเพลต TULIBS ให้หัวข้อระดับเดียวกันเยื้องเท่ากันทุกจุด "
                         "(X.X ชิดขอบ · X.X.X 0.80\" · X.X.X.X 1.10\" · (1) 1.40\")",
                         f"para {items[0][0]}",
                         "เลือกค่าเดียวแล้วใช้ให้เหมือนกันทั้งเล่ม (แก้ที่สไตล์ได้ทีเดียว)"))
            continue
        modes[level] = top
        outliers = [it for it in items if abs(it[2] - top) > _INDENT_TOL]
        for i, pg, value, text in outliers[:8]:
            out.append(F("minor", "การเยื้อง",
                         f"หัวข้อ “{text[:50]}” เยื้อง {value}\" ต่างจากหัวข้อระดับเดียวกัน "
                         f"ที่ใช้ {top}\" ({hits} จาก {len(items)} จุด)",
                         "หัวข้อระดับเดียวกันต้องเยื้องเท่ากันทั้งเล่ม ไม่งั้นลำดับชั้นดูเพี้ยน",
                         f"para {i}", f"ตั้งเยื้องบรรทัดแรกเป็น {top}\""))
        if len(outliers) > 8:
            out.append(F("minor", "การเยื้อง",
                         f"หัวข้อระดับ {level} ยังมีจุดที่เยื้องไม่ตรงอีก {len(outliers) - 8} จุด",
                         "รายงานเฉพาะ 8 จุดแรกเพื่อไม่ให้ตารางยาวเกินจำเป็น",
                         f"para {outliers[8][0]}",
                         f"ตั้งเยื้องบรรทัดแรกเป็น {top}\" ให้ทุกหัวข้อระดับนี้"))

    for level in sorted(modes):
        if level + 1 in modes and modes[level + 1] < modes[level] - _INDENT_TOL:
            out.append(F("major", "การเยื้อง",
                         f"หัวข้อระดับ {level + 1} เยื้อง {modes[level + 1]}\" น้อยกว่าหัวข้อแม่ "
                         f"ระดับ {level} ที่เยื้อง {modes[level]}\"",
                         "หัวข้อย่อยต้องเยื้องมากกว่าหัวข้อที่ครอบมันเสมอ ไม่งั้นผู้อ่านแยกไม่ออก"
                         "ว่าหัวข้อไหนอยู่ใต้หัวข้อไหน",
                         "ทั้งเอกสาร",
                         f"ตั้งหัวข้อระดับ {level + 1} ให้เยื้องมากกว่า {modes[level]}\""))


# ---------------------------------------------------------------------------
def inspect_document(doc, path, profile=None, only=None, started_at=None, page_data=None):
    """ตรวจ deep checks บน Document ที่โหลดแล้วและคืน (payload, findings)."""
    t0 = started_at if started_at is not None else time.time()
    profile = profile or detect_profile(doc)
    resolved_pages = page_data if page_data is not None else resolve_page_data(path, doc)
    pages, has_markers = unpack_page_data(resolved_pages, doc)

    out = []
    if only != "apa7":
        check_toc_pages(doc, pages, out)
        # check_captions ถูกถอดออก — ดูเหตุผลในหมายเหตุด้านบน
        check_chapter_layout(doc, pages, out)
        check_keywords_bold(doc, pages, out)
        check_front_matter_tables(doc, pages, out)
        check_cover_elements(doc, pages, out)
        check_blank_lines_in_prose(doc, pages, out)
        check_heading_styles(doc, pages, out)
        check_heading_indent_ladder(doc, pages, out)
    if only != "template":
        check_apa_mechanical(doc, pages, out)

    enrich_dict_locations_with_pages(out, resolved_pages)

    real = [f for f in out if f["severity"] != "info"]
    counts = {s: sum(1 for f in real if f["severity"] == s) for s in ("critical", "major", "minor")}
    payload = {"file": os.path.basename(path), "profile": profile,
               "profile_label": PROFILES[profile]["label"],
               "counts": counts, "findings": real,
               "elapsed_sec": round(time.time() - t0, 2),
               "has_page_numbers": has_markers}
    if isinstance(resolved_pages, dict):
        payload.update({
            "page_source": resolved_pages.get("source", "unavailable"),
            "total_pages": resolved_pages.get("total_pages", 0),
            "page_match_ratio": resolved_pages.get("match_ratio", 0.0),
            "page_map_warning": resolved_pages.get("warning", ""),
        })
    return payload, out


def print_summary(path, payload, out):
    profile = payload["profile"]
    counts = payload["counts"]
    has_markers = payload.get("has_page_numbers", False)

    print("=" * 70)
    print(f"ตรวจชั้นลึก (deterministic) — {os.path.basename(path)}   profile: {profile}")
    print(f"พบ {SEV_ICON['critical']} {counts['critical']}  {SEV_ICON['major']} {counts['major']}  {SEV_ICON['minor']} {counts['minor']}"
          + (f"   แผ่นงาน: {payload.get('total_pages', 0)} ({payload.get('page_source', 'unavailable')})"
             if has_markers else "   ⚠ ไม่สามารถระบุแผ่นงานได้"))
    print("=" * 70)
    for f in out:
        print(f"{SEV_ICON.get(f['severity'], '·')} [{f['domain']}·{f['category']}] @ {f['location']} {f['issue']}")
        if f["correct"]:
            print(f"      → {f['correct']}")
    print("=" * 70)
    print("ขอบเขต: เฉพาะข้อผิดที่ตัดสินได้โดยไม่ต้องตีความ. สิ่งที่ยังต้องให้คนหรือ LLM อ่าน: "
          "องค์ประกอบหน้าปก/หน้าอนุมัติ, ความถูกต้องของเนื้อความรายการอ้างอิงแต่ละประเภท, "
          "Title Case/Sentence case, การเลือกประเภทแหล่งอ้างอิง")


def run(path, json_out=None, only=None, profile=None):
    t0 = time.time()
    doc = Document(path)
    payload, out = inspect_document(doc, path, profile, only, started_at=t0)
    print_summary(path, payload, out)
    if json_out:
        with open(json_out, "w", encoding="utf-8") as fh:
            json.dump(payload, fh, ensure_ascii=False, indent=2)
        print(f"JSON written: {json_out}")
    return payload, out


def main():
    ap = argparse.ArgumentParser(description="ตรวจโครงสร้างเอกสาร + APA7 เชิงกลไก (deterministic)")
    ap.add_argument("docx")
    ap.add_argument("--json", default=None, help="เขียนผลเป็น JSON (ป้อนเข้า make_report.py ได้เลย)")
    ap.add_argument("--only", choices=["template", "apa7"], default=None)
    ap.add_argument("--profile", choices=list(PROFILES.keys()), default=None)
    a = ap.parse_args()
    if not os.path.exists(a.docx):
        sys.exit(f"File not found: {a.docx}")
    run(a.docx, a.json, a.only, a.profile)


if __name__ == "__main__":
    main()
