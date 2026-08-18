#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_report.py — สร้างรายงานการแก้ไขเป็นไฟล์ .docx (ตารางแก้ทีละจุด)

ใช้กับผลตรวจจาก:
  • check_docx.py --json  (ไฟล์ Word)
  • check_pdf.py  --json  (ไฟล์ PDF จากสกิล tulibs-thesis-format-checker)
  • ไฟล์ JSON ที่เขียนเองจากการตรวจด้วยตา (visual pass)

รูปแบบรายงาน: หัวรายงาน + ตารางแก้ทีละจุด (ลำดับ / ตำแหน่ง / ระดับ / ปัญหา /
เกณฑ์ที่ถูกต้อง / วิธีแก้) + สรุปจำนวน + ประเมินความพร้อม + checklist ตรวจด้วยตา

การใช้งาน:
    # ทางลัด: ให้ check_docx.py สร้างรายงานให้เลย
    python3 check_docx.py student.docx --report report.docx

    # หรือรวมหลาย JSON (เช่น docx + visual, หรือ PDF + visual) เป็นรายงานเดียว
    python3 make_report.py auto.json visual.json -o report.docx --title "ชื่อวิทยานิพนธ์"

รูปแบบ finding ที่รองรับ (ยืดหยุ่น — map คีย์ให้อัตโนมัติ):
    severity/sev/level · location/loc/page/หน้า · issue/message/problem/msg ·
    criterion/requirement/rule/เกณฑ์ · fix/recommendation/วิธีแก้ ·
    detail (legacy fallback) · category/cat/หมวด
"""
import sys, os, json, argparse, datetime, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = "F2F2F2"

SEV_ORDER = {"critical": 0, "major": 1, "minor": 2, "info": 3, "ok": 4}
SEV_TH = {
    "critical": ("🔴 ต้องแก้ก่อนส่ง", "C00000"),
    "major":    ("🟠 ควรแก้",         "C55A11"),
    "minor":    ("🟡 แก้เล็กน้อย",    "BF9000"),
    "info":     ("ℹ️ ข้อมูล",          "2E74B5"),
    "ok":       ("🟢 ผ่าน",           "548235"),
}

def pick(d, *keys, default=""):
    for k in keys:
        if k in d and d[k] not in (None, ""):
            return d[k]
    return default

# ---- ตำแหน่งที่แสดงในรายงาน -------------------------------------------------
# รายงานสำหรับนักศึกษาแสดงเฉพาะสิ่งที่หาเจอใน Word: แผ่นที่ + ตาราง/แถว/คอลัมน์
# **ไม่แสดงเลขย่อหน้า (¶N)** เพราะ Word ไม่มีเลขย่อหน้าให้ดู นักศึกษาเอาไปใช้ไม่ได้
# ¶N ยังรับเข้ามาได้ตามเดิม (สคริปต์ตรวจอัตโนมัติใส่มาเอง) และ annotate_docx.py ยังใช้มัน
# ปักคอมเมนต์ — แค่ถูกตัดออกตอนแสดงผลเท่านั้น
_TBL_INNER = re.compile(r"tbl(\d+):r(\d+):c(\d+)(?::¶\d+)?", re.I)
# รูปแบบไทยจาก check_docx/check_deep: "ตาราง 27 แถว 1 คอลัมน์ 2 ย่อหน้า 1"
_TBL_THAI = re.compile(r"ตาราง\s+(\d+(?:\.\d+)*)\s+แถว\s+(\d+)\s+คอลัมน์\s+(\d+)"
                       r"(?:\s+ย่อหน้า\s+\d+)?")
_PARA_TOKEN = re.compile(r"(?:\bpara\b\s*|¶\s*|ย่อหน้า\s*)\d+", re.I)


def _strip_para(text):
    """ตัด ¶N / para N ออก แล้วเก็บกวาดตัวคั่นที่ค้าง"""
    text = _TBL_INNER.sub(
        lambda m: f"ตารางที่ {m.group(1)} แถว {m.group(2)} คอลัมน์ {m.group(3)}", text)
    text = _TBL_THAI.sub(
        lambda m: f"ตารางที่ {m.group(1)} แถว {m.group(2)} คอลัมน์ {m.group(3)}", text)
    text = _PARA_TOKEN.sub("", text)
    text = re.sub(r"[(\[]\s*[)\]]", "", text)          # วงเล็บที่ว่างเพราะตัด ¶ ออก
    text = re.sub(r"(?:\s*·\s*)+", " · ", text)
    text = re.sub(r"\s*:\s*(?=[\"“'‘])", " · ", text)   # "para 12: “ข้อความ”" → "“ข้อความ”"
    text = re.sub(r"^\s*[:·|,]\s*", "", text)
    return re.sub(r"\s{2,}", " ", text).strip(" ·|,:")


def display_location(raw):
    """คืนข้อความตำแหน่งที่นักศึกษาใช้ได้จริง — แผ่นที่ (+ ตาราง/แถว/คอลัมน์) ไม่มีเลขย่อหน้า"""
    location = str(raw or "")

    def _tag(m, tilde=""):
        inner = _strip_para(m.group(2))
        return f"แผ่นที่ {tilde}{m.group(1)}" + (f" · {inner}" if inner else "")

    location = re.sub(r"\[sheet\s*(\d+)\|([^\]]*)\]", _tag, location, flags=re.I)
    location = re.sub(r"\[p\s*(\d+)\|([^\]]*)\]",
                      lambda m: _tag(m, "~"), location, flags=re.I)
    location = re.sub(r"(?<!ย่อ)(?<!แผ่น)หน้า\s*(~?)\s*(\d+)",
                      lambda m: f"แผ่นที่ {m.group(1)}{m.group(2)}", location)
    return _strip_para(location)


# ---- คำค้นสำหรับ Ctrl+F ------------------------------------------------------
# เลขแผ่นเป็นตัวเลขที่คนเอาไปใช้ยาก (เลื่อนได้ทุกครั้งที่แก้ไฟล์ และ Word ไม่มี
# "แผ่นที่" ให้กระโดดไปตรง ๆ) สิ่งที่ใช้ได้จริงคือ **ข้อความในเล่ม** — ก๊อปไปวางใน
# ช่องค้นหาแล้วเคอร์เซอร์ไปถึงจุดนั้นทันที รายงานจึงยกข้อความจริงมาเป็นหลัก
# แล้วให้เลขแผ่นเป็นข้อมูลประกอบ
# ห้ามให้ 1.5" (นิ้ว) กลายเป็นเครื่องหมายคำพูด — negative lookbehind กันตัวเลขนำหน้า
_QUOTED_RE = re.compile(r"(?<![\d])[\"“'‘]([^\"”'’\n]{6,120})[\"”'’]")
_PARA_KEY_RE = re.compile(r"(?:\bpara\b\s*|¶)\s*(\d+)", re.I)
_TBL_KEY_RE = re.compile(r"tbl(\d+):r(\d+):c(\d+):¶(\d+)", re.I)
_TBL_KEY_TH_RE = re.compile(
    r"ตาราง\s+(\d+(?:\.\d+)*)\s+แถว\s+(\d+)\s+คอลัมน์\s+(\d+)\s+ย่อหน้า\s+(\d+)")
SEARCH_MAX = 58


def corpus_from_index(index):
    """ยุบข้อความจากดัชนีเป็นก้อนเดียว — ใช้ยืนยันว่าคำค้นค้นแล้วเจอจริง

    สร้างจากดัชนีที่มีอยู่แล้วแทนการเปิดไฟล์ .docx ซ้ำรอบสอง (เล่มจริง ~9 MB
    การเปิดสองรอบกินเวลาราวครึ่งหนึ่งของขั้นออกรายงานทั้งขั้น)
    """
    parts = [v for k, v in index.items() if not k.startswith("style:") and k != ORDER_KEY]
    return re.sub(r"\s+", " ", "\n".join(parts).replace("\u00a0", " "))


ORDER_KEY = "__order__"        # คีย์พิเศษในดัชนี: ลำดับ (key, text) ตามเอกสารจริง


def context_key(finding):
    """คืนคีย์ดัชนีของย่อหน้าที่ finding ชี้ (para N / tblX:...) หรือ None"""
    location = str(pick(finding, "location", "loc", "page", "หน้า", "ตำแหน่ง"))
    m = _TBL_KEY_RE.search(location) or _TBL_KEY_TH_RE.search(location)
    if m:
        return "tbl{}:r{}:c{}:¶{}".format(*m.groups())
    m = _PARA_KEY_RE.search(location)
    return f"para {m.group(1)}" if m else None


def build_search_index(docx_path):
    """คืน dict คีย์ตำแหน่ง → ข้อความจริงในเล่ม (ไว้เติมคำค้นให้ finding ที่ไม่ได้ยกข้อความมา)"""
    index = {}
    if not docx_path or not os.path.exists(docx_path):
        return index
    try:
        from docx import Document as _Doc
        from docx.table import Table as _Table
        from docx.oxml.ns import qn as _qn
    except ImportError:
        return index
    try:
        doc = _Doc(docx_path)
    except Exception:                                          # noqa: BLE001
        return index
    def remember_style(name, text):
        if not name or not text:
            return
        key = f"style:{str(name).strip().casefold()}"
        if key not in index and len(text) >= 4:
            index[key] = text

    # เดินตาม body ตามลำดับจริง (ย่อหน้ากับตารางสลับกัน) — เลข para/tbl ตรงกับ
    # python-docx เพราะทั้ง doc.paragraphs และ doc.tables ก็เรียงตาม body เหมือนกัน
    # ลำดับนี้ทำให้บอกได้ว่าคำค้นที่ซ้ำหลายจุด "จุดที่ต้องแก้คือจุดที่เท่าไร"
    from docx.table import Table as _Table
    order = []
    para_i, tbl_i = 0, 0
    body_paragraphs = doc.paragraphs
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if para_i < len(body_paragraphs):
                paragraph = body_paragraphs[para_i]
                text = paragraph.text or ""      # ดิบ — ช่องว่างหัว/ท้ายย่อหน้าคือสิ่งที่ finding ชี้
                if text.strip():
                    index[f"para {para_i}"] = text
                    order.append((f"para {para_i}", text))
                    try:
                        remember_style(paragraph.style.name if paragraph.style else "", text)
                        for run in paragraph.runs:      # character style เช่น TU_..._Char
                            if run.style is not None and (run.text or "").strip():
                                remember_style(run.style.name, run.text.strip())
                    except Exception:                          # noqa: BLE001
                        pass
            para_i += 1
        elif child.tag == qn("w:tbl"):
            tbl_i += 1
            seen = set()
            table = _Table(child, doc._body)
            for ri, row in enumerate(table.rows, 1):
                for ci, cell in enumerate(row.cells, 1):
                    if cell._tc in seen:
                        continue
                    seen.add(cell._tc)
                    for pi, paragraph in enumerate(cell.paragraphs, 1):
                        text = paragraph.text or ""
                        if text.strip():
                            key = f"tbl{tbl_i}:r{ri}:c{ci}:¶{pi}"
                            index[key] = text
                            order.append((key, text))
    index[ORDER_KEY] = order
    return index


# อักขระไทยที่ขึ้นต้นคำไม่ได้ (สระหลัง/วรรณยุกต์/ตัวสะกดพิเศษ) — คำค้นที่ขึ้นต้นด้วย
# พวกนี้คือหลักฐานว่าถูกตัดกลางคำ อ่านก็ไม่รู้เรื่อง (สระลอยเป็นวงกลมประ) และทำให้
# ผู้ใช้ไม่กล้าเชื่อรายงาน ต้องเล็มออกจนขึ้นต้นด้วยอักขระที่เริ่มคำได้จริง
_THAI_DEPENDENT = set("ะัาำิีึืุู็่้๊๋์ํ๎ๆฯ")


def clean_word_start(text):
    """เล็มหัวข้อความให้ขึ้นต้นที่อักขระที่เริ่มคำได้

    ข้ามไปเริ่มหลังช่องว่าง **เฉพาะเมื่อหัวข้อความขาดจริง** (ขึ้นต้นด้วยสระตาม/วรรณยุกต์)
    — เวอร์ชันก่อนข้ามทุกกรณี ทำให้ "ตารางที่ 2.2" โดนตัดเหลือ "2.2" ซึ่งค้นเจอเป็นสิบจุด
    """
    text = str(text or "").lstrip("… ")
    if text and text[0] in _THAI_DEPENDENT:
        space = text.find(" ", 0, 14)
        if space != -1:                   # หัวขาดและมีขอบคำจริงใกล้ ๆ — เริ่มหลังช่องว่าง
            text = text[space + 1:]
        while text and text[0] in _THAI_DEPENDENT:
            text = text[1:]
    return text.strip()


def _expand_word(text, start, end):
    """ขยายช่วง [start,end) ออกไปจนสุดคำ (ชนช่องว่าง/วรรคตอน) — ใช้ดึง 'คำผิดทั้งคำ' จากบริบท"""
    boundary = " \t\n\r,.;:()[]{}\"'“”‘’!?/\\"
    while start > 0 and text[start - 1] not in boundary:
        start -= 1
    while end < len(text) and text[end] not in boundary:
        end += 1
    return text[start:end]


def _edit_distance(a, b, cap=4):
    """Levenshtein แบบตัดจบเร็ว — สองคำสะกดต่างกันไม่กี่อักขระเท่านั้นที่เราสนใจ"""
    if abs(len(a) - len(b)) > cap:
        return cap + 1
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        cur = [i]
        for j, cb in enumerate(b, 1):
            cur.append(min(prev[j] + 1, cur[-1] + 1, prev[j - 1] + (ca != cb)))
        if min(cur) > cap:
            return cap + 1
        prev = cur
    return prev[-1]


def sentence_window(ctx, start, end, pad=26):
    """คืน (ก่อนช่วง, หลังช่วง) รอบ ๆ [start,end) โดยตัดที่ขอบคำ — ไว้ยกประโยคจริงจากเล่ม

    ห้ามมี "…" ปนเด็ดขาด เพราะข้อความนี้ผู้ใช้ต้องก๊อปไปวางในช่อง Find/Replace ทั้งก้อน
    """
    a, b = max(0, start - pad), min(len(ctx), end + pad)
    if a > 0:
        space = ctx.rfind(" ", a, start)
        if space != -1 and space - a < 12:
            a = space + 1        # เศษหน้าช่องว่างสั้น ๆ — เริ่มหลังช่องว่างให้ขอบคำสวย
    space = ctx.find(" ", end, b)
    if space != -1 and b - space < 12:
        b = space                # ตัดเศษท้ายที่สั้นทิ้ง แต่ไม่หั่นหน้าต่างจนเหลือนิดเดียว
    left = clean_word_start(ctx[a:start]) if a > 0 else ctx[a:start].lstrip()
    return left, ctx[end:b].rstrip()


def find_wrong_word(context, correct):
    """หา 'คำที่สะกดผิด' ในบริบท โดยเทียบกับคำที่ถูก

    ตัวตรวจบอกมาแค่คำที่ถูก (เช่น "แอปพลิเคชัน") แต่ผู้ใช้ต้องเอา **คำที่ผิด**
    ("แอปพลิเคชั่น") ไปวางในช่อง Find — เลือกช่วงในบริบทที่ "แก้น้อยครั้งที่สุดแล้วได้คำถูก"
    ห้ามใช้การขยายให้สุดคำ เพราะภาษาไทยไม่มีวรรคคั่น จะกินคำข้างเคียงเข้ามา
    (เช่นได้ "หรืออีเมล์" แทน "อีเมล์" — ผู้ใช้ทำตามแล้วคำว่า "หรือ" หายไปทั้งเล่ม)
    """
    if not context or not correct:
        return ""
    # จุดเริ่มค้น: ตำแหน่งของส่วนร่วมยาวสุดระหว่างคำถูกกับบริบท
    anchor = -1
    for size in range(len(correct), 3, -1):
        for offset in range(0, len(correct) - size + 1):
            pos = context.find(correct[offset:offset + size])
            if pos != -1:
                anchor = pos - offset      # ประมาณจุดที่คำผิดเริ่ม
                break
        if anchor != -1:
            break
    if anchor == -1:
        return ""
    best, best_score = "", 30              # คะแนน = ระยะแก้ไข×10 + ค่าปรับขอบคำ
    for start in range(max(0, anchor - 2), anchor + 3):
        for length in range(max(2, len(correct) - 2), len(correct) + 4):
            end = start + length
            cand = context[start:end]
            if not cand or cand == correct:
                continue
            if cand in correct:
                # คำผิดที่เป็นแค่ "เศษ" ของคำถูก (เช่น "อีเม" ⊂ "อีเมล") ใช้ไม่ได้ —
                # Replace แล้วอักษรที่เหลือในเอกสารจะไปต่อท้ายคำถูกกลายเป็นตัวซ้ำ
                continue
            d = _edit_distance(cand, correct, cap=2)
            if d > 2:
                continue
            score = d * 10
            # ตัดจบก่อนสระ/วรรณยุกต์ท้ายคำ = คำขาด (เช่นได้ "ออีเมล" ทิ้ง ์ ไว้ข้างหลัง
            # พอ Replace แล้วเหลือ "หรือีเมล์" — คำข้างเคียงพังทั้งเล่ม) ปรับหนัก
            if end < len(context) and context[end] in _THAI_DEPENDENT:
                score += 5
            # เริ่มหลังสระตาม = เริ่มกลางคำของคำข้างเคียง ปรับหนักเช่นกัน
            if start > 0 and context[start - 1] in _THAI_DEPENDENT:
                score += 5
            if cand[0] in _THAI_DEPENDENT:
                score += 5
            # เสมอกันเอาตัวยาวกว่า — ตัวสั้นคือตัวยาวที่โดนตัดท้าย ซึ่ง Replace แล้วจะ
            # เหลืออักขระตกค้าง (เช่น "แอปพลิเคชั่" ทิ้ง "น" ไว้ → ได้ "แอปพลิเคชันน")
            if score < best_score or (score == best_score and len(cand) > len(best)):
                best, best_score = cand, score
    if not best or best[0] in _THAI_DEPENDENT:
        return ""
    return best.strip()


def _searchable(text):
    """เลือกช่วงที่ "ค้นแล้วเจอจริง" — เลี่ยงช่วงที่มีเว้นวรรคซ้อน/ขึ้นบรรทัดใหม่

    Ctrl+F ของ Word เทียบตัวอักษรตรง ๆ ถ้าคำค้นคร่อมจุดที่เคาะ space 2 ครั้ง
    (ซึ่งคือ "ตัวปัญหา" ของ finding ประเภทเว้นวรรคซ้อนพอดี) แล้วเราส่งไปเป็น
    space เดียว ผู้ใช้จะค้นไม่เจอ — ตัดเอาเฉพาะฝั่งที่ยาวกว่าแทน
    """
    text = str(text or "").replace("\n", " ").replace("\t", " ").strip()
    # \u00a0 (non-breaking space) พิมพ์ในช่องค้นหาไม่ได้ ตัดฝั่งที่มีมันทิ้งเหมือนเว้นวรรคซ้อน
    parts = [p.strip() for p in re.split(r"\s{2,}|\u00a0+", text) if p.strip()]
    if len(parts) > 1:
        return max(parts, key=len)
    return text


def _trim(text, limit=SEARCH_MAX):
    text = re.sub(r"\s+", " ", clean_word_start(_searchable(text))).strip()
    if len(text) <= limit:
        return text
    cut = text[:limit]
    space = cut.rfind(" ")
    if space > limit * 0.6:
        cut = cut[:space]
    return cut.rstrip(" ·,;:") + "…"


def full_context(finding, index=None):
    """คืนข้อความเต็มของย่อหน้าที่ finding ชี้ (จากดัชนี) — เก็บช่องว่างซ้อน/อักขระพิเศษไว้ครบ

    ใช้เป็นวัตถุดิบให้ humanize ระบุ "คำผิดคือคำไหน อยู่ระหว่างคำอะไรกับคำอะไร"
    ถ้าไม่มีดัชนี ใช้ข้อความที่ยกมาใน location แทน (สั้นกว่าแต่ยังมีช่องว่างซ้อนของจริง)
    """
    index = index or {}
    location = str(pick(finding, "location", "loc", "page", "หน้า", "ตำแหน่ง"))
    m = _TBL_KEY_RE.search(location) or _TBL_KEY_TH_RE.search(location)
    if m and index.get("tbl{}:r{}:c{}:¶{}".format(*m.groups())):
        return index["tbl{}:r{}:c{}:¶{}".format(*m.groups())]
    m = _PARA_KEY_RE.search(location)
    if m and index.get(f"para {m.group(1)}"):
        return index[f"para {m.group(1)}"]
    quoted = _QUOTED_RE.findall(location)
    return quoted[0] if quoted else ""


def search_text(finding, index=None):
    """ข้อความที่ผู้อ่านรายงานเอาไป Ctrl+F ได้จริง — คืน "" ถ้าเป็นปัญหาระดับทั้งเล่ม"""
    index = index or {}
    direct = str(pick(finding, "snippet", "quote", "search", default="")).strip()
    if len(direct) >= 4:
        return _trim(direct)
    location = str(pick(finding, "location", "loc", "page", "หน้า", "ตำแหน่ง"))
    # ตัวตรวจอัตโนมัติใส่ "ข้อความรอบจุดที่ผิด" ไว้ใน location อยู่แล้ว (เช่น
    # para 78: “…ะของงานวิจัยนี้  ศาสตราจารย…”) — อันนี้ตรงจุดที่สุด ใช้ก่อนเสมอ
    # ส่วนข้อความใน issue มักเป็น "ค่าที่ถูกต้อง" ซึ่งยังไม่มีในเล่ม ค้นแล้วไม่เจอ
    for quoted in _QUOTED_RE.findall(location):
        cleaned = quoted.strip().strip("…").strip()
        if len(cleaned) >= 5 and not re.search(r"[XxN]\.[XxN]", cleaned):
            return _trim(cleaned)
    for field in ("issue", "correct", "detail", "fix"):
        for quoted in _QUOTED_RE.findall(str(finding.get(field) or "")):
            cleaned = quoted.strip().strip("…").strip()
            if len(cleaned) >= 6 and not re.search(r"[XxN]\.[XxN]", cleaned):
                return _trim(cleaned)
    m = _TBL_KEY_RE.search(location) or _TBL_KEY_TH_RE.search(location)
    if m:
        text = index.get("tbl{}:r{}:c{}:¶{}".format(*m.groups()))
        if text:
            return _trim(text)
    m = _PARA_KEY_RE.search(location)
    if m:
        text = index.get(f"para {m.group(1)}")
        if text:
            return _trim(text)
    return ""



# ---- แปลผลตรวจให้เป็นภาษาที่นักศึกษาอ่านแล้วลงมือแก้ได้ทันที --------------------
# ตัวตรวจอัตโนมัติพูดภาษาเครื่อง ("left margin 1.5" — required 1.0"") และมักไม่ได้
# เขียนช่อง "เกณฑ์" กับ "วิธีแก้" มาให้เลย รายงานดิบจึงอ่านแล้วไม่รู้ว่าต้องกดอะไร
# ชั้นนี้เติมสามอย่างให้ครบทุกแถว: พบอะไร (ภาษาคน) · เกณฑ์ของ TULIBS · กดตรงไหนใน Word
_SIDE_TH = {"left": "ซ้าย", "right": "ขวา", "top": "บน", "bottom": "ล่าง",
            "header": "หัวกระดาษ", "footer": "ท้ายกระดาษ", "gutter": "สันเล่ม"}
_PROP_TH = {"ตัวหนา": "ตัวหนา", "ตัวเอียง": "ตัวเอียง", "ขนาด": "ขนาดฟอนต์",
            "จัดกึ่งกลาง": "การจัดวาง", "ระยะห่างบรรทัด": "ระยะห่างบรรทัด"}

_RE_MARGIN = re.compile(r"(left|right|top|bottom|header|footer|gutter)\s+margin\s+"
                        r"([\d.]+)\"?\s*[—-]\s*required\s+([\d.]+)", re.I)
# ชื่อสไตล์จริงมีทั้งขีดกลางและตัวเลข (TU_Sub-heading 1, TU_Main Heading_Chapter2)
# ถ้า class ตัวอักษรแคบไป กฎนี้จะไม่จับ แล้ว finding ตกไปใช้ข้อความดิบภาษาเครื่อง
_RE_STYLE_DRIFT = re.compile(r"^(TU_[\w \-]+?):\s*(.+?)\s*=\s*(.+?)\s*\(เทมเพลต\s*=\s*(.+?)\)")
_RE_SIZE = re.compile(r"([\w_ ]+)/(thai|latin|ไทย|ละติน):\s*([\d.]+)pt\s*→\s*([\d.]+)pt")
_RE_FOREIGN_FONT = re.compile(r"Foreign fonts found[^:]*:\s*(.+)", re.I)
_RE_COLOR = re.compile(r"(\d+)\s*run\(s\) use non-black text colour", re.I)
_RE_HYPHEN = re.compile(r"ช่วงเลขหน้า\s*'([^']+)'\s*ใช้ยัติภังค์")
_RE_NO_ENTRY = re.compile(r"อ้าง\s*'([^']+)'\s*ในเนื้อหา แต่ไม่พบในรายการอ้างอิง")
_RE_ORDER = re.compile(r"รายการ\s*'([^']+)'\s*อยู่ผิดลำดับตัวอักษร")
_RE_SPELL = re.compile(r"คำสะกดผิด:\s*แก้เป็น\s*[“\"']([^”\"']+)")


def _bool_th(value):
    v = str(value).strip().lower()
    return {"true": "ใช่", "false": "ไม่ใช่", "none": "ไม่ได้กำหนด"}.get(v, str(value))


def humanize(f, ctx=""):
    """เติม/เขียนใหม่ช่อง issue · criterion · fix · correct ให้เป็นภาษาที่ลงมือทำได้

    รับ dict ที่ normalize แล้ว แก้ในตัว (in place) และคืนค่ากลับ
    """
    category = f.get("category") or ""
    issue = f.get("issue") or ""
    fix = f.get("fix") or ""
    correct = f.get("correct") or ""
    generic = ("แก้ค่า/ข้อความให้ตรงเกณฑ์ที่ระบุ", "")
    if fix in generic or fix == correct:
        fix = ""

    # --- กฎที่มาจากไฟล์ตัวอย่างทางการ: ต้องบอก "กดตรงไหนใน Word" ให้ครบเหมือนกฎอื่น --
    if category == "การเยื้อง":
        m = re.search(r"เยื้อง ([\d.]+)\" ต่างจาก", issue)
        target = re.search(r"ที่ใช้ ([\d.]+)\"", issue)
        if m and target:
            want = target.group(1)
            # เยื้อง 0" ใน Word คือ Special: (none) ไม่ใช่ First line ที่ใส่เลข 0 —
            # บอกผิดตรงนี้ผู้ใช้จะกดตามแล้วไม่เห็นอะไรเปลี่ยน แล้วเลิกเชื่อรายงาน
            step = ("ช่อง Special เลือก (none)" if float(want) < 0.01
                    else f"ช่อง Special เลือก First line → ช่อง By ใส่ {want}\"")
            f["fix"] = (f"คลิกในย่อหน้าหัวข้อนี้ → แท็บ Home → กดลูกศรมุมขวาล่างของกลุ่ม Paragraph "
                        f"→ {step} → OK "
                        f"(ถ้าหัวข้อระดับนี้ใช้สไตล์เดียวกันทั้งเล่ม แก้ที่สไตล์ทีเดียวจบ: "
                        f"คลิกขวาที่ชื่อสไตล์ในกล่อง Styles → Modify)")
        else:
            f["fix"] = ("เลือกค่าเยื้องของหัวข้อระดับนี้ให้เหลือค่าเดียว แล้วตั้งที่สไตล์: "
                        "แท็บ Home → กล่อง Styles → คลิกขวาที่สไตล์ของหัวข้อ → Modify → "
                        "Format → Paragraph → Special: First line")
        return f

    if category in ("หน้าปก", "หน้าอนุมัติ"):
        if "ลิขสิทธิ์" in issue or "COPYRIGHT" in issue:
            f["fix"] = ("พิมพ์บรรทัดนี้เป็นบรรทัดสุดท้ายของหน้าปก จัดกึ่งกลาง ขนาดเท่าข้อความอื่นบนปก: "
                        "ปกไทยใช้ “ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์” · "
                        "ปกอังกฤษใช้ “COPYRIGHT OF THAMMASAT UNIVERSITY”")
        elif "รับรอง" in issue:
            f["fix"] = ("เพิ่มประโยครับรองบนหน้าอนุมัติ เหนือตารางรายชื่อกรรมการ: "
                        "“ได้รับการตรวจสอบและอนุมัติ ให้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร…” "
                        "แล้วต่อด้วยชื่อปริญญาและบรรทัด “เมื่อ วันที่ …”")
        else:
            f["fix"] = ("เพิ่มบรรทัดนี้ลงหน้าปก โดยเทียบลำดับกับไฟล์ "
                        "`ตัวอย่างการพิมพ์ส่วนประกอบของวิทยานิพนธ์` ของหอสมุด "
                        "(ชื่อเรื่อง → โดย/BY → ชื่อผู้เขียน → ข้อความหลักสูตร → "
                        "คณะ/มหาวิทยาลัย → ปีการศึกษา → ลิขสิทธิ์)")
        return f

    if category == "สไตล์หัวข้อ":
        want = re.search(r"ต้องใช้สไตล์ (TU_[\w\- ]*[\w\-])", str(f.get("criterion") or ""))
        if want:
            style = want.group(1).strip()
            # ชื่อจริงของสไตล์หัวข้อใหญ่มีเลขบทต่อท้าย (TU_Main Heading_Chapter3)
            # ถ้าบอกแค่ "TU_Main Heading" ผู้อ่านจะหาไม่เจอในกล่อง Styles
            hint = ""
            if style == "TU_Main Heading":
                num = re.search(r"“(\d+)\.", str(f.get("issue") or ""))
                style = f"TU_Main Heading_Chapter{num.group(1)}" if num else "TU_Main Heading_ChapterN"
                hint = " (เลขท้ายชื่อสไตล์คือเลขบท — บทที่ 1 ใช้ TU_Main Heading _Chapter1 มีเว้นวรรคหน้า _Chapter1)"
            f["fix"] = (f"คลิกที่บรรทัดหัวข้อนี้ (คลิกที่ไหนก็ได้ในบรรทัด ไม่ต้องลากคลุม) "
                        f"→ แท็บ Home → ในกล่อง Styles เลื่อนหา “{style}” แล้วคลิก{hint} "
                        f"— ขนาด ความหนา และการเยื้องจะเปลี่ยนตามเทมเพลตให้เอง ไม่ต้องตั้งเอง")
            f["correct"] = f"สไตล์ {style}"
        else:
            f["fix"] = ("ไล่คลิกหัวข้อทีละอันแล้วเลือกสไตล์จากกล่อง Styles ให้ตรงระดับ: "
                        "X.X = TU_Main Heading_ChapterN (N คือเลขบท) · X.X.X = TU_Sub-heading 1 · "
                        "X.X.X.X = TU_Sub-heading 2 · (1) = TU_Sub-heading 3\n"
                        "ทำเสร็จแล้วอย่าลืมอัปเดตสารบัญ: คลิกที่สารบัญ → References → Update Table "
                        "→ Update entire table")
            f["correct"] = "สไตล์หัวข้อตามระดับที่เทมเพลตกำหนด"
        return f

    if category == "การเว้นบรรทัด":
        f["fix"] = ("กด Ctrl+H → กดปุ่ม More → ติ๊ก Use wildcards ออกก่อน → "
                    "ช่อง Find พิมพ์ ^p^p → ช่อง Replace พิมพ์ ^p → Replace All "
                    "(ทำซ้ำจนขึ้นว่าแทนที่ 0 ครั้ง) — วิธีนี้ลบบรรทัดว่างซ้อนทั้งเล่มรวดเดียว. "
                    "จากนั้นตรวจว่าหัวข้อและหน้าปกที่ตั้งใจเว้นบรรทัดไว้ยังอยู่ครบ")
        f["correct"] = "ย่อหน้าเนื้อความติดกัน ไม่มีบรรทัดว่างคั่น"
        return f

    if category == "ส่วนนำ":
        if "ไม่พบตารางข้อมูล" in issue:
            f["fix"] = ("คัดลอกตารางข้อมูลจากไฟล์เทมเพลต/ตัวอย่างของหอสมุดมาวางที่หัวหน้าบทคัดย่อ "
                        "แล้วพิมพ์ข้อมูลของตัวเองแทน — พิมพ์เองเป็นบรรทัดเปล่าจะจัดคอลัมน์ไม่ตรง")
        elif "ขาดรายการ" in issue:
            f["fix"] = ("คลิกในแถวสุดท้ายของตาราง → กด Tab เพื่อเพิ่มแถว (หรือแท็บ Layout → "
                        "Insert Below) → พิมพ์ป้ายรายการที่ขาดในคอลัมน์ซ้าย ข้อมูลในคอลัมน์ขวา")
        elif "วงเล็บ" in issue:
            f["fix"] = "พิมพ์วงเล็บครอบชื่อกรรมการทุกคน เช่น (ศาสตราจารย์ ดร. ก ข)"
        return f

    m = _RE_MARGIN.search(issue)
    if m:
        side = _SIDE_TH.get(m.group(1).lower(), m.group(1))
        found, need = m.group(2), m.group(3)
        # "ส่วนนี้" เพียว ๆ ไม่พอ — Word ไม่แสดงเลข section ที่ไหนเลย ผู้อ่านต้องรู้ว่า
        # หน้าไหน. ตัวตรวจแนบข้อความบรรทัดแรกของ section มาให้ใน location แล้ว
        anchor = re.search(r"ส่วนที่ขึ้นต้นด้วย “([^”]+)”", str(f.get("location") or ""))
        land = "landscape" in str(f.get("location") or "")
        spot = f"หน้าที่ขึ้นต้นด้วย “{anchor.group(1)}”" if anchor else "ส่วนนี้"
        kind = " (หน้าแนวนอน)" if land else ""
        f["issue"] = (f"ระยะขอบ{side}ของ{spot}{kind} ตั้งไว้ {found} นิ้ว "
                      f"(เกณฑ์คือ {need} นิ้ว)")
        f["criterion"] = ("TULIBS: ขอบบน 1.5 · ขอบล่าง 1.0 · ขอบซ้าย 1.5 · ขอบขวา 1.0 นิ้ว "
                          "(หน้าแนวนอนสลับเป็นซ้าย 1.0 · ขวา 1.5)")
        f["fix"] = (
            (f"ก๊อป “{anchor.group(1)}” ไปค้นใน Ctrl+F เพื่อไปให้ถึงหน้านั้นก่อน → "
             if anchor else "คลิกในหน้าที่มีปัญหา → ")
            + f"แท็บ Layout → Margins → Custom Margins… "
              f"→ ช่อง{side} ใส่ {need} → ช่อง Apply to เลือก This section → OK "
              f"(ถ้าเลือก Whole document จะไปเปลี่ยนหน้าอื่นที่ตั้งถูกอยู่แล้วด้วย)")
        f["correct"] = f"{need} นิ้ว"
        return f

    # --- ข้อความที่ตัวตรวจยังเขียนเป็นภาษาอังกฤษเชิงเครื่อง: แปลให้ลงมือได้ ----------
    m = re.search(r"Core TULIBS styles are missing:\s*(.+)", issue)
    if m:
        f["issue"] = ("ไฟล์นี้ไม่มีชุดสไตล์ของเทมเพลต TULIBS อยู่เลย "
                      f"(ที่ขาด: {m.group(1).strip()}) — แปลว่าเล่มนี้ไม่ได้พิมพ์ทับบนไฟล์เทมเพลต "
                      "จึงต้องตั้งฟอนต์/ขนาด/การเยื้องเองทุกจุด และมีโอกาสพลาดสูง")
        f["criterion"] = "เทมเพลต TULIBS มาพร้อมสไตล์ TU_Chapter · TU_Paragraph_Normal · TU_Main Heading · TU_Sub-heading"
        f["fix"] = ("เปิดไฟล์เทมเพลตเปล่าของหอสมุด → เลือกเนื้อหาทั้งหมดในเล่มปัจจุบัน (Ctrl+A) "
                    "→ Ctrl+C → ไปวางในไฟล์เทมเพลตด้วย Paste Special แบบ Merge Formatting "
                    "→ แล้วไล่ตั้งสไตล์ให้หัวข้อแต่ละระดับ. "
                    "ถ้าไม่อยากย้ายไฟล์ ต้องยอมรับว่ารายการรูปแบบด้านล่างจะต้องแก้เองทีละจุด")
        f["correct"] = "ไฟล์ที่สร้างจากเทมเพลต TULIBS"
        return f

    m = re.search(r"page ([\d.]+)x([\d.]+)\" is not A4 \(([\d.]+)x([\d.]+)", issue)
    if m:
        gw, gh, ew, eh = m.groups()
        f["issue"] = (f"ขนาดกระดาษของส่วนนี้ตั้งไว้ {gw}×{gh} นิ้ว ซึ่งเป็นขนาด Letter "
                      f"ไม่ใช่ A4 ({ew}×{eh} นิ้ว) — พิมพ์ออกมาแล้วขอบกระดาษจะไม่ตรงเกณฑ์ทั้งเล่ม")
        f["criterion"] = "TULIBS: กระดาษ A4 เท่านั้น (8.27 × 11.69 นิ้ว / 210 × 297 มม.)"
        f["fix"] = ("แท็บ Layout → Size → เลือก A4 · ถ้าเล่มมีหลายส่วน ให้กด Ctrl+A ก่อนเลือก "
                    "แล้วไล่ตรวจซ้ำว่าหน้าแนวนอน (ถ้ามี) ยังเป็น A4 แนวนอนอยู่")
        f["correct"] = f"A4 = {ew} × {eh} นิ้ว"
        return f

    if "Could not confirm automatic page-number fields" in issue:
        f["issue"] = ("ตรวจไม่พบเลขหน้าแบบอัตโนมัติในหัวกระดาษ — ถ้าพิมพ์เลขหน้าด้วยมือ "
                      "เลขจะไม่ขยับตามเมื่อเนื้อหาเลื่อน และรูปแบบ (1)/1 จะเพี้ยนเมื่อแก้ไฟล์")
        f["criterion"] = "TULIBS: เลขหน้าอยู่มุมขวาบน ห่างขอบบน 1.0 นิ้ว ใส่เป็น field อัตโนมัติ"
        f["fix"] = ("ดับเบิลคลิกที่หัวกระดาษเพื่อเข้าโหมด Header → ลบเลขที่พิมพ์เองออก → "
                    "แท็บ Insert → Page Number → Top of Page → Plain Number 3 (ชิดขวา) "
                    "→ ตั้งรูปแบบส่วนนำเป็น (1) ที่ Page Number → Format Page Numbers")
        f["correct"] = "เลขหน้าเป็น field อัตโนมัติ มุมขวาบน"
        return f

    if "ยังไม่มีข้อความไหนในเล่มใช้สไตล์เหล่านี้" in issue:
        # แถวสรุปสไตล์ที่ไม่ได้ใช้ — เขียนให้ชัดว่า "ไม่ต้องรีบแก้" ไม่ใช่งานค้าง
        f["criterion"] = ("สไตล์ที่ไม่มีข้อความใช้ ไม่มีผลกับสิ่งที่พิมพ์ออกมา "
                          "จึงไม่ใช่ข้อผิดพลาดของรูปเล่ม")
        f["correct"] = "ปล่อยไว้ได้ (หรือแก้ให้ตรงเทมเพลตเพื่อความเรียบร้อย)"
        return f

    m = _RE_STYLE_DRIFT.search(issue)
    if m:
        style, prop, now, want = m.groups()
        prop_th = _PROP_TH.get(prop, prop)
        hits = re.search(r"มีผลกับข้อความ (\d+) จุด", issue)
        seen = re.search(r"เห็นได้ที่ “([^”]+)”", issue)
        # เขียนจาก "สิ่งที่ตาเห็น" ก่อน แล้วค่อยบอกว่าต้นเหตุคือสไตล์ — ผู้อ่านที่ไม่รู้จัก
        # ระบบสไตล์ของ Word ต้องเข้าใจได้ว่ามันคือข้อความตรงไหนในเล่มของตัวเอง
        where = ""
        if seen:
            where = f"เช่นบรรทัด “{seen.group(1)}”"
        scope = f" (มีทั้งหมด {hits.group(1)} จุดในเล่ม)" if hits else ""
        f["issue"] = (f"หัวข้อ/ข้อความที่จัดรูปแบบด้วยสไตล์ {style} มี{prop_th} = "
                      f"{_bool_th(now)} แต่เทมเพลตกำหนดไว้ที่ {_bool_th(want)} {where}{scope}")
        f["criterion"] = f"เทมเพลต TULIBS กำหนดให้สไตล์ {style} มี{prop_th} = {_bool_th(want)}"
        f["fix"] = (
            f"วิธีที่เร็วที่สุด (แก้ครั้งเดียวได้ทุกจุด): คลิกในบรรทัดที่ยกมาในช่องซ้าย "
            f"→ แท็บ Home → ในกล่อง Styles ชื่อสไตล์ที่ถูกไฮไลต์คือ {style} "
            f"→ คลิกขวาที่ชื่อนั้น → Modify… → ตั้ง{prop_th}เป็น {_bool_th(want)} → OK\n"
            f"ถ้าหากล่อง Styles ไม่เจอ: เลือกข้อความบรรทัดนั้น แล้วตั้ง{prop_th}เป็น "
            f"{_bool_th(want)} จากแถบเครื่องมือ Home ตรง ๆ แต่ต้องทำซ้ำทุกจุด")
        f["correct"] = f"{prop_th} = {_bool_th(want)}"
        return f

    m = _RE_SIZE.search(issue)
    if m:
        style, script, want, found = m.groups()
        script_th = "ไทย" if script.lower() in ("thai", "ไทย") else "อังกฤษ/ละติน"
        f["issue"] = (f"ข้อความนี้ถูกตั้งขนาดฟอนต์ {script_th} เป็น {found} pt ทับสไตล์ {style.strip()} "
                      f"ซึ่งกำหนดไว้ {want} pt")
        f["criterion"] = f"สไตล์ {style.strip()} ของเทมเพลต TULIBS = {want} pt ({script_th})"
        f["fix"] = (f"ลากคลุมข้อความนี้ → แท็บ Home → ช่องขนาดฟอนต์ ใส่ {want} "
                    f"หรือกด Ctrl+Spacebar เพื่อล้างการตั้งค่าที่ทับสไตล์ทิ้ง")
        f["correct"] = f"{want} pt"
        return f

    m = _RE_FOREIGN_FONT.search(issue)
    if m:
        f["issue"] = f"พบฟอนต์นอกเทมเพลตปนอยู่: {m.group(1)[:110]}"
        f["criterion"] = correct or "ทั้งเล่มต้องใช้ TH Sarabun New 16 pt (โปรไฟล์ Times = Times New Roman 12 pt)"
        f["fix"] = ("ลากคลุมข้อความที่ระบุ → แท็บ Home → ช่องชื่อฟอนต์ พิมพ์ชื่อฟอนต์ของเทมเพลต "
                    "(ถ้าเป็นตารางให้คลุมทั้งตาราง) — ถ้าเจอหลายจุด ใช้ Ctrl+A แล้วตั้งทีเดียว "
                    "แต่ต้องตั้งค่าสไตล์ให้ถูกก่อน")
        f["correct"] = correct or "TH Sarabun New 16 pt"
        return f

    m = _RE_COLOR.search(issue)
    if m:
        f["replace_all"] = True            # แก้ด้วย Ctrl+A ทีเดียวทั้งเล่ม ไม่ต้องชี้จุด
        f["issue"] = f"มีข้อความ {m.group(1)} จุดที่ไม่ใช่สีดำ (เช่น ตัวอักษรสีแดง/น้ำเงินที่ค้างมาจากการแก้งาน)"
        f["criterion"] = "ตัวอักษรในเล่มต้องเป็นสีดำล้วนทั้งหมด"
        f["fix"] = ("ลากคลุมข้อความ (หรือ Ctrl+A ทั้งเล่ม) → แท็บ Home → ปุ่ม Font Color (ตัว A มีแถบสี) "
                    "→ เลือก Automatic")
        f["correct"] = "สีดำ (Automatic)"
        return f

    m = _RE_HYPHEN.search(issue)
    if m:
        rng = m.group(1)
        f["issue"] = f"ช่วงเลขหน้า {rng} ใช้ขีดสั้น (-) ซึ่งไม่ใช่เครื่องหมายของ APA 7"
        f["criterion"] = "APA 7: ช่วงเลขหน้าใช้ en dash (–) ไม่ใช่ hyphen (-)"
        f["fix"] = ("แทนที่ขีดสั้นด้วย en dash: กด Ctrl+H → ช่อง Find พิมพ์เลขช่วงนี้ "
                    "→ ช่อง Replace พิมพ์ค่าที่ถูกต้องในคอลัมน์ขวา (en dash พิมพ์ได้ด้วย Ctrl+ลบบนแป้นตัวเลข)")
        f["correct"] = correct or rng.replace("-", "–")
        return f

    if "ไม่มีคำบรรยาย" in issue:
        anchor = _trim(ctx) if (ctx or "").strip() else ""
        is_table = "ตาราง" in issue or "Table" in issue
        kind = "ตาราง" if is_table else "ภาพ"
        if anchor:
            f["search"], f["search_kind"] = anchor, "exact"
            f["issue"] = (f"{kind}ที่อยู่ถัดจากข้อความ “{anchor}” ยังไม่มีบรรทัดคำบรรยาย "
                          f"({kind}ที่ บท.ลำดับ ชื่อ{kind})")
        f["criterion"] = ("คำบรรยายตารางอยู่เหนือตาราง · คำบรรยายภาพอยู่ใต้ภาพ "
                          "รูปแบบ: '<ตาราง/ภาพ>ที่ <บท>.<ลำดับที่ในบท> <ชื่อ>'")
        f["fix"] = (f"ค้นด้วยคำค้นในช่องซ้าย → {kind}อยู่ถัดลงไป → พิมพ์บรรทัดคำบรรยาย"
                    + ("เหนือตาราง" if is_table else "ใต้ภาพ")
                    + f" เช่น '{kind}ที่ 2.5 ชื่อ{kind}' แล้วไล่เลขให้ต่อเนื่องกับ{kind}ก่อนหน้า")
        f["correct"] = f"{kind}ที่ <บท>.<ลำดับ> <ชื่อ>"
        return f

    m = _RE_NO_ENTRY.search(issue)
    if m:
        cite = m.group(1)
        cm = re.match(r"(.+?)\s*\(\s*(\d{4})\s*\)", cite)
        if cm:
            f["_cite"] = (cm.group(1).strip(), cm.group(2))
        f["issue"] = f"ในเนื้อหาอ้าง {cite} แต่ไม่มีรายการนี้ในบรรณานุกรมท้ายเล่ม"
        f["criterion"] = "APA 7: ทุกการอ้างอิงในเนื้อหาต้องมีรายการเต็มท้ายเล่ม และทุกรายการท้ายเล่มต้องถูกอ้างในเนื้อหา"
        f["fix"] = ("ตรวจสองทาง: ถ้าเป็นงานที่อ้างจริง ให้เพิ่มรายการเต็มในบรรณานุกรม; "
                    "ถ้าเป็นชื่อเครื่องมือ/หัวข้อที่ไม่ใช่การอ้างอิง ให้แก้ข้อความไม่ให้อยู่ในรูป (ชื่อ, ปี) "
                    "เพราะจะถูกอ่านว่าเป็นการอ้างอิง")
        f["correct"] = correct or "เพิ่มรายการเต็มท้ายเล่ม หรือแก้ข้อความให้ไม่อยู่ในรูปการอ้างอิง"
        return f

    m = _RE_ORDER.search(issue)
    if m:
        f["criterion"] = "APA 7: เรียงบรรณานุกรมตามตัวอักษรของชื่อผู้แต่งคนแรก"
        f["fix"] = "ย้ายรายการนี้ขึ้นไปไว้ก่อนรายการที่ระบุ ให้ลำดับตัวอักษรถูกต้อง"
        return f

    m = _RE_SPELL.search(issue)
    if m:
        right = m.group(1)
        wrong = find_wrong_word(ctx, right)
        if wrong:
            # รู้ทั้งคำผิดและคำถูก → บอกครบในประโยคเดียว และให้คำผิดเป็นคำค้นตรง ๆ
            pos = ctx.find(wrong)
            example = ""
            if pos != -1:
                before, after = sentence_window(ctx, pos, pos + len(wrong), pad=20)
                example = f" — ตัวอย่างในเล่ม: “{before}{wrong}{after}”"
            f["issue"] = f"สะกดผิด: “{wrong}” — ที่ถูกคือ “{right}”{example}"
            f["fix"] = f"กด Ctrl+H → ช่อง Find พิมพ์ “{wrong}” → ช่อง Replace พิมพ์ “{right}” → Replace All"
            f["search"], f["search_kind"], f["replace_all"] = wrong, "exact", True
        else:
            f["issue"] = f"สะกดผิด — ที่ถูกคือ “{right}”"
            f["fix"] = (f"ค้นด้วยคำค้นในช่องซ้าย แล้วแก้คำที่สะกดผิดตรงนั้นเป็น “{right}” "
                        f"(หรือ Ctrl+H ถ้าคำนี้ผิดซ้ำหลายจุด)")
        f["criterion"] = "ใช้คำสะกดตามพจนานุกรมฉบับราชบัณฑิตยสถาน"
        f["correct"] = right
        return f

    if "นิคหิต" in issue or "ํา" in issue:
        m = re.search(r"[\u0E01-\u0E2E][\u0E31-\u0E4E]*ํา[\u0E01-\u0E4E]*", ctx or "")
        if m:
            wrong = _expand_word(ctx, m.start(), m.end())
            right = wrong.replace("ํา", "ำ")
            f["issue"] = (f"คำว่า “{wrong}” พิมพ์สระอำด้วยอักขระสองตัว (นิคหิต ํ + สระอา า) "
                          f"ตาเห็นเหมือน “{right}” แต่เครื่องมองเป็นคนละคำ — ค้นหาไม่เจอ ตัดคำผิด และบางฟอนต์แสดงเพี้ยน")
            f["criterion"] = "สระอำต้องเป็นอักขระเดียว (ำ U+0E33)"
            f["fix"] = (f"กด Ctrl+H → ช่อง Find วางคำจากช่องคำค้น (“{wrong}”) → "
                        f"ช่อง Replace พิมพ์ “{right}” ใหม่ด้วยแป้นพิมพ์ปกติ → Replace All")
            f["search"], f["search_kind"], f["replace_all"] = wrong, "exact", True
            f["correct"] = right
        else:
            f["criterion"] = "สระอำต้องเป็นอักขระเดียว (ำ U+0E33)"
            f["fix"] = "ค้นด้วยคำค้นในช่องซ้าย แล้วพิมพ์สระอำใหม่ด้วยแป้นพิมพ์ปกติ"
        return f

    if "เว้นวรรคซ้อน" in issue:
        stripped = (ctx or "").rstrip()
        lead = re.match(r"[ \t\u00a0]{2,}", ctx or "")
        trail = re.search(r"\S([ \t\u00a0]{2,})$", ctx or "")
        m = re.search(r"\S[ \t\u00a0]{2,}\S", ctx or "")
        if not m and trail:
            # เคาะซ้อนอยู่ท้ายย่อหน้า — ไม่มีคำตามหลังให้ยกมา
            tail_start = trail.start(1)
            before, _ = sentence_window(ctx, tail_start, len(ctx))
            f["issue"] = (f"มีช่องว่างเกินค้างอยู่ **ท้ายย่อหน้า** หลังคำว่า “{before}” "
                          f"(มองไม่เห็นด้วยตา แต่ทำให้การจัดชิดขอบ/ตัดบรรทัดเพี้ยน)")
            f["fix"] = ("ค้นด้วยคำค้นในช่องซ้าย → คลิกท้ายบรรทัดนั้น → กด End แล้วลบช่องว่างจนชิดตัวอักษรสุดท้าย "
                        "(หรือเปิด Home → ¶ เพื่อมองเห็นช่องว่างก่อนลบ)")
            f["search"], f["search_kind"] = before, "exact"
            f["_expand"] = "left"
            f["correct"] = "ไม่มีช่องว่างท้ายย่อหน้า"
            f["criterion"] = "เว้นวรรคระหว่างคำ/ประโยคเคาะครั้งเดียว"
            return f
        if not m and lead and stripped:
            # เคาะซ้อนนำหน้าย่อหน้า — นักศึกษาใช้ space จัดย่อหน้าแทนการตั้งเยื้อง
            head = stripped.strip()
            first_words = " ".join(head.split()[:3])[:36] or head[:36]
            f["issue"] = (f"ย่อหน้านี้ใช้การเคาะ space/Tab {len(lead.group(0))} ครั้งจัดย่อหน้า "
                          f"หน้าคำว่า “{first_words}” แทนการตั้งเยื้องบรรทัดแรกของสไตล์")
            f["fix"] = ("ค้นด้วยคำค้นในช่องซ้าย → คลิกหน้าตัวอักษรแรกของบรรทัด → ลบช่องว่างจนชิดขอบ "
                        "→ ถ้าต้องการย่อหน้า ให้ตั้งที่สไตล์: Format Paragraph → Special: First line = 0.8\"")
            f["search"], f["search_kind"] = first_words, "exact"
            f["correct"] = "เยื้องบรรทัดแรกด้วยการตั้งค่าย่อหน้า ไม่ใช่การเคาะ space"
            f["criterion"] = "การเยื้องทำที่สไตล์/การตั้งค่าย่อหน้า ไม่ใช้ space หรือ Tab"
            return f
        if m:
            gap = re.search(r"[ \t\u00a0]{2,}", ctx[m.start():]).span()
            gap_start, gap_end = m.start() + gap[0], m.start() + gap[1]
            gap_text = ctx[gap_start:gap_end]
            before, after = sentence_window(ctx, gap_start, gap_end)
            if "\t" in gap_text:
                # ช่องว่างมี Tab ปน — วาง Tab ในช่อง Find ของ Word ไม่ได้ ให้ชี้จุดด้วยคำซ้ายแทน
                f["issue"] = (f"มีช่องว่าง/แท็บเกิน ({len(gap_text)} อักขระ) ระหว่าง “{before}” กับ “{after}”")
                f["fix"] = ("ค้นด้วยคำค้นในช่องซ้าย → เปิด Home → ¶ (Show formatting) จะเห็นช่องว่างและ "
                            "ลูกศรแท็บหลังคำนั้น → ลบให้เหลือเคาะเดียว")
                f["search"], f["search_kind"] = before, "exact"
                f["_expand"] = "left"
                f["correct"] = f"{before} {after}"
                f["criterion"] = "เว้นวรรคระหว่างคำ/ประโยคเคาะครั้งเดียว ไม่ใช้ Tab คั่นกลางประโยค"
                return f
            wrong_phrase = f"{before}{gap_text}{after}"
            fixed_phrase = f"{before} {after}"
            # ยกประโยคผิดมาทั้งท่อน **โดยเก็บเคาะซ้อนไว้ข้างใน** — ก๊อปไปค้นแล้วเจอจุดเดียว
            # แน่นอน (เคาะซ้อนทำให้ไม่ซ้ำกับข้อความปกติ) และใช้เป็นช่อง Find ของ Ctrl+H ได้เลย
            f["issue"] = (f"เคาะเว้นวรรคเกิน 1 ครั้ง ระหว่าง “{before.split()[-1] if before.split() else before}” "
                          f"กับ “{after.split()[0] if after.split() else after}” — ข้อความที่ผิด: “{wrong_phrase}”")
            f["fix"] = ("กด Ctrl+H → ช่อง Find วางข้อความจากช่องคำค้น (มีเคาะซ้อนอยู่ข้างใน "
                        "จึงเจอจุดนี้จุดเดียว) → ช่อง Replace วางข้อความจากช่อง 'ค่าที่ถูกต้อง' → Replace")
            f["search"], f["search_kind"] = wrong_phrase, "exact"
            f["_expand"] = "none"          # คำค้นคือประโยคผิดตรงตัว ห้ามให้ระบบยืด/แก้ทับ
            f["correct"] = fixed_phrase
        else:
            f["issue"] = "เคาะเว้นวรรคเกิน 1 ครั้งติดกัน"
            f["fix"] = ("กด Ctrl+H → ช่อง Find เคาะ space 2 ครั้ง → ช่อง Replace เคาะ 1 ครั้ง "
                        "→ กด Replace All ซ้ำจนขึ้นว่าแทนที่ 0 จุด")
            f["correct"] = "เว้นวรรคครั้งเดียว"
        f["criterion"] = "เว้นวรรคระหว่างคำ/ประโยคเคาะครั้งเดียว"
        return f

    if "non-breaking space" in issue or "U+00A0" in issue:
        m = re.search(r"\u00a0+", ctx or "")
        f["issue"] = "มีอักขระเว้นวรรคชนิดพิเศษ (non-breaking space) ปนอยู่ มักติดมาจากการก๊อปข้อความจากเว็บ"
        if m:
            before, after = sentence_window(ctx, m.start(), m.end())
            wrong_phrase = f"{before}{ctx[m.start():m.end()]}{after}"
            f["issue"] += f" — ข้อความที่ผิด: “{wrong_phrase}”"
            f["search"], f["search_kind"] = wrong_phrase, "exact"
            f["_expand"] = "none"
            f["correct"] = f"{before} {after}"
            f["fix"] = ("กด Ctrl+H → ช่อง Find วางข้อความจากช่องคำค้น (อักขระพิเศษติดมาด้วย) → "
                        "ช่อง Replace วางจากช่อง 'ค่าที่ถูกต้อง' → Replace · "
                        "หรือกวาดทั้งเล่ม: Find พิมพ์ ^s → Replace เคาะ space → Replace All")
        else:
            f["fix"] = "กด Ctrl+H → ช่อง Find พิมพ์ ^s → ช่อง Replace เคาะ space 1 ครั้ง → Replace All"
            f["correct"] = "เว้นวรรคปกติ"
        f["criterion"] = "ใช้เว้นวรรคปกติเท่านั้น"
        return f

    if category.startswith("caption"):
        is_figure = "ภาพ" in category
        f["criterion"] = ("TULIBS: คำบรรยายภาพอยู่ **ใต้** ภาพ · คำบรรยายตารางอยู่ **เหนือ** ตาราง")
        f["fix"] = ("คลิกที่บรรทัดคำบรรยาย → กด Ctrl+X (ตัด) → คลิกตำแหน่งที่ถูกต้อง "
                    + ("ใต้ภาพ" if is_figure else "เหนือตาราง") + " → กด Ctrl+V (วาง)")
        f["correct"] = correct or ("ย้ายคำบรรยายลงไปใต้ภาพ" if is_figure else "ย้ายคำบรรยายขึ้นไปเหนือตาราง")
        return f

    # ไม่เข้ากฎไหน — อย่างน้อยต้องไม่ปล่อยช่อง "วิธีแก้" ว่างหรือซ้ำกับค่าที่ถูกต้อง
    if not f.get("fix") or f["fix"] in generic:
        f["fix"] = (f"แก้ให้เป็น: {correct}" if correct
                    else "ดูรายละเอียดในช่อง 'ปัญหาที่พบ' แล้วแก้ให้ตรงเกณฑ์")
    return f


_RE_STYLE_IN_LOC = re.compile(r"styles?\s+([\w][\w .\-_]*?)\s*(?:$|·|\||,)", re.I)


def style_example(finding, index=None):
    """ข้อความตัวอย่างของสไตล์ที่มีปัญหา — ให้ผู้อ่านค้นไปดูได้ว่าหน้าตาเป็นยังไง

    ปัญหาระดับสไตล์ (เช่นขนาดฟอนต์ของ TU_Chapter เพี้ยน) แก้ที่ตัวสไตล์ครั้งเดียวจบ
    แต่ถ้าไม่มีอะไรให้ค้นเลย ผู้อ่านจะไม่รู้ว่ามันคือหัวข้อไหนในเล่ม
    """
    index = index or {}
    if not index:
        return ""

    def lookup(name):
        name = re.sub(r"\s+", " ", name).strip(" .:·|,")
        if not name:
            return ""
        # "TU_X Char" = character style ที่ผูกกับ paragraph style "TU_X"
        variants = [name, re.sub(r"\s*Char$", "", name, flags=re.I)]
        for variant in variants:
            words = variant.split(" ")
            # ตัดคำท้ายทีละคำ เผื่อจับชื่อมาเกิน (ชื่อสไตล์มีเว้นวรรคในตัวเอง
            # เลยแยกจากข้อความที่ตามมาด้วย regex อย่างเดียวไม่ได้)
            for cut in range(len(words), 0, -1):
                text = index.get("style:" + " ".join(words[:cut]).casefold())
                if text:
                    return text
        return ""

    for field in ("location", "issue", "correct", "criterion"):
        value = str(finding.get(field) or "")
        for m in re.finditer(r"TU_[\w .\-]+", value):
            text = lookup(m.group(0))
            if text:
                return _trim(text)
        m = _RE_STYLE_IN_LOC.search(value)
        if m:
            text = lookup(m.group(1))
            if text:
                return _trim(text)
    return ""


def normalize(finding, index=None):
    domain = str(pick(finding, "domain", "type", "kind", "ประเภทตรวจ", default="template")).lower()
    category = str(pick(finding, "category", "cat", "หมวด"))
    if domain == "spelling" or category.casefold() == "spelling" or "คำผิด" in category:
        domain = "spelling"
    if domain not in ("template", "spelling", "apa7"):
        domain = "template"
    legacy_detail = str(pick(finding, "detail", default=""))
    criterion = str(pick(finding, "criterion", "requirement", "rule", "เกณฑ์", default=""))
    fix = str(pick(finding, "fix", "recommendation", "วิธีแก้", default=""))
    correct = str(pick(finding, "correct", "expected", "ค่าที่ถูกต้อง", "target"))
    if not criterion and legacy_detail:
        if any(token in legacy_detail.lower() for token in
               ("apa7", "apa 7", "tulibs", "docx-spec", "inspection-checklist", "เทมเพลต", "กำหนด")):
            criterion = legacy_detail
        else:
            fix = fix or legacy_detail
    if not criterion:
        criterion = ("คู่มือ TULIBS APA7 / APA Style 7 ตามภาษาของรายการ"
                     if domain == "apa7" else
                     "TULIBS template profile ที่ระบุ (ดู references/docx-spec.md)")
    if not fix:
        fix = correct or legacy_detail or "แก้ค่า/ข้อความให้ตรงเกณฑ์ที่ระบุ"
    location = display_location(pick(finding, "location", "loc", "page", "หน้า", "ตำแหน่ง"))
    search = search_text(finding, index)
    search_kind = "exact" if search else ""
    if not search:
        search = style_example(finding, index)
        search_kind = "example" if search else ""
    ctx = full_context(finding, index)
    key = context_key(finding)
    return humanize({
        "_ctx": ctx,
        "_key": key,
        "search": search,
        "search_kind": search_kind,
        "domain":   domain,
        "severity": str(pick(finding, "severity", "sev", "level", default="minor")).lower(),
        "location": location,
        "category": category,
        "issue":    str(pick(finding, "issue", "message", "problem", "msg", "ปัญหา")),
        "criterion": criterion,
        "fix":       fix,
        "correct":   correct,
    }, ctx)

# ---- ตารางการตั้งค่า layout / หน้ากระดาษ (อ้างอิงตามโปรไฟล์) ----------------
LAYOUT = {
    "thai": {"font": "TH Sarabun New", "body": "16 pt", "ls": "เดี่ยว (1.0)", "xx": "0.0 นิ้ว (ชิดขอบ)"},
    "english": {"font": "TH Sarabun New", "body": "16 pt", "ls": "เดี่ยว (1.0)", "xx": "0.25 นิ้ว"},
    "english-times": {"font": "Times New Roman (อังกฤษ) / TH Sarabun New (ไทย)",
                      "body": "12 pt (อังกฤษ) / 16 pt (ไทย)", "ls": "1.5", "xx": "0.25 นิ้ว"},
}
def layout_rows(profile):
    L = LAYOUT.get(profile, LAYOUT["thai"])
    return [
        ("ขนาดกระดาษ", "A4 (8.27 × 11.69 นิ้ว / 210 × 297 มม.)"),
        ("แนวหน้ากระดาษ", "แนวตั้ง (Portrait) เป็นหลัก; landscape เฉพาะตาราง/ภาพใหญ่"),
        ("ขอบบน (Top)", "1.5 นิ้ว"),
        ("ขอบล่าง (Bottom)", "1.0 นิ้ว"),
        ("ขอบซ้าย (Left) — สันเล่ม", "1.5 นิ้ว  (landscape = 1.0 นิ้ว)"),
        ("ขอบขวา (Right)", "1.0 นิ้ว  (landscape = 1.5 นิ้ว)"),
        ("Header (เลขหน้าจากขอบบน)", "1.0 นิ้ว"),
        ("Footer", "0.49 นิ้ว"),
        ("ฟอนต์เนื้อหา", L["font"]),
        ("ขนาดเนื้อหา", L["body"]),
        ("ระยะห่างบรรทัด", L["ls"]),
        ("เยื้องย่อหน้า (บรรทัดแรก)", "0.80 นิ้ว"),
        ("หัวข้อ X.X เยื้องซ้าย", L["xx"]),
        ("ชื่อบท / หัวข้อส่วนนำ", "18 pt หนา กึ่งกลาง"),
        ("ชื่อเรื่องปก / หน้าคั่นภาคผนวก", "20 pt หนา"),
        ("เลขหน้าส่วนนำ", "(1)(2)(3)… เริ่มที่บทคัดย่อ, มุมขวาบน"),
        ("เลขหน้าเนื้อหา", "1 2 3 … เริ่มที่บทที่ 1, มุมขวาบน"),
    ]

def _fill_location(cell, finding):
    """ช่อง 'จุดที่ต้องแก้' — บรรทัดแรกคือข้อความที่ก๊อปไปค้นใน Word ได้ทันที

    บรรทัดที่สองคือแผ่น/ส่วนของเล่ม ไว้ช่วยกะตำแหน่งคร่าว ๆ เท่านั้น ไม่ใช่ตัวหลัก
    เพราะเลขแผ่นเลื่อนได้ทุกครั้งที่นักศึกษาแก้ไฟล์
    """
    cell.text = ""
    paragraph = cell.paragraphs[0]
    search = finding.get("search") or ""
    # **ไม่แสดงเลขแผ่นหรือชื่อบทเลย** — เลขแผ่นเลื่อนทุกครั้งที่แก้ไฟล์ ใส่ไปมีแต่ทำให้
    # ผู้อ่านสับสนว่าจะเชื่ออันไหน ช่องนี้จึงมีอย่างเดียวคือ "ข้อความที่ก๊อปไปค้นได้"
    if search:
        run = paragraph.add_run(f"🔎 “{search}”")
        run.bold = True
        run.font.size = Pt(11)
        hint = cell.add_paragraph()
        hint_run = hint.add_run(
            "ตัวอย่างจุดที่ใช้สไตล์นี้ — ก๊อปไปค้นใน Ctrl+F เพื่อดูของจริง (แก้ที่สไตล์ทีเดียวได้ทั้งเล่ม)"
            if finding.get("search_kind") == "example" else "ก๊อปข้อความนี้ไปวางใน Ctrl+F")
        hint_run.italic = True
        hint_run.font.size = Pt(9)
        hint_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        note = finding.get("search_note")
        if note:
            note_p = cell.add_paragraph()
            _add_markup_runs(note_p, note, base_size=9, color="1F3864")
            for run in note_p.runs:
                run.italic = True
        if finding.get("search_ok") is False:
            warn = cell.add_paragraph()
            # finding ประเภท "ยังไม่มีบรรทัดนี้ในเล่ม" ค้นไม่เจอเป็นเรื่องปกติ — นั่นคือ
            # ตัวปัญหาเอง. ถ้าติดธงเตือนเหมือนกรณีอื่น ผู้อ่านจะเข้าใจว่ารายงานพลาด
            # แล้วข้ามแถวที่ถูกต้องไป
            if finding.get("category") in ("หน้าปก", "หน้าอนุมัติ"):
                text, color = ("นี่คือข้อความที่ยังขาด — ก๊อปไปวางในเล่มได้เลย",
                               RGBColor(0x1F, 0x38, 0x64))
            else:
                text, color = ("⚠ ค้นข้อความนี้ในเล่มไม่พบ — ตรวจรายการนี้ซ้ำก่อนส่ง",
                               RGBColor(0xC0, 0x00, 0x00))
            warn_run = warn.add_run(text)
            warn_run.italic = True
            warn_run.font.size = Pt(10)
            warn_run.font.color.rgb = color
        return
    # ไม่มีคำค้น = แก้ที่การตั้งค่า **หรือ** เป็นเรื่อง "สิ่งที่ยังไม่มีในเล่ม"
    # (เช่น ภาพที่ยังไม่มี caption) ซึ่งคนละเรื่องกันคนละวิธีหา ถ้าเขียนรวมเป็น
    # "แก้ที่การตั้งค่าหน้ากระดาษ/สไตล์" เหมือนกันหมด ผู้อ่านจะไปงมผิดที่
    note = finding.get("search_note")
    run = paragraph.add_run("ยังไม่มีข้อความนี้ในเล่ม" if note else "ทั้งเอกสาร")
    run.bold = True
    run.font.size = Pt(11)
    hint = cell.add_paragraph()
    hint_run = hint.add_run(
        note or "ไม่ใช่ข้อความจุดใดจุดหนึ่ง — ทำตามขั้นตอนในช่อง “วิธีแก้” ได้เลย")
    hint_run.italic = True
    hint_run.font.size = Pt(9)
    hint_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    tcpr.append(sh)

def _set(cell, text, bold=False, color=None, size=11, align=None):
    cell.text = ""
    par = cell.paragraphs[0]
    if align is not None:
        par.alignment = align
    run = par.add_run(str(text))
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def readiness(counts):
    c, m, mi = counts.get("critical", 0), counts.get("major", 0), counts.get("minor", 0)
    if c == 0 and m == 0 and mi == 0:
        return "ยังไม่พบปัญหาในผลที่ป้อน — ยืนยัน checklist ก่อนส่ง"
    if c == 0 and m == 0:
        return "พบเฉพาะจุดเล็กน้อย"
    if c == 0 and m <= 3:
        return "ควรแก้จุดสำคัญก่อนส่ง"
    return "ต้องแก้ก่อนส่ง"

# ---- ตัวรายการ checklist ตรวจด้วยตา (สคริปต์ทำแทนไม่ได้) --------------------
VISUAL_ITEMS = [
    "รูปแบบเลขหน้า (1)(2) ในส่วนนำ และ 1 2 3 ในเนื้อหา แสดงถูกส่วน",
    "ข้อความบนหน้าปก/ปกภาษาที่สองครบและถูกต้อง; ปีการศึกษา พ.ศ./ค.ศ. ตรงภาษา",
    "หน้าอนุมัติใช้แบบเดียว (มี/ไม่มีที่ปรึกษาร่วม) กรรมการครบ มีเส้นเซ็น+วงเล็บชื่อ",
    "บทคัดย่อ: ตารางข้อมูลครบ 6 รายการและเนื้อหาไม่เกิน 1 หน้า",
    "ชื่อบทภาษาอังกฤษเป็นตัวพิมพ์ใหญ่; ไม่มีหัวข้อลอยท้ายหน้า",
    "ตารางตกหน้า: ทวนหัวตาราง + (ต่อ)/(Continued)",
]

APA7_VISUAL_ITEMS = [
    "แต่ละรายการจัดกลุ่มภาษาถูก: ไทย → TULIBS APA7, อังกฤษ/ละติน → APA Style ทางการ",
    "ชื่อผู้แต่งไทยไม่กลับคำ (ชื่อ+สกุล), ผู้แต่งอังกฤษกลับคำแบบ Last, F. M.",
    "ปีพิมพ์: ไทยใช้ พ.ศ., อังกฤษใช้ ค.ศ. — ไม่แปลงข้ามระบบในรายการเดียวกัน",
    "\"และ\"/\"และคณะ\" ใช้กับรายการไทย; \"&\"/\"et al.\" ใช้กับรายการอังกฤษ",
    "ตัวเอียงลงถูกที่: ชื่อวารสาร/หนังสือ/รายงาน เอียง, ชื่อบทความ/บทไม่เอียง",
    "Title Case สำหรับชื่อวารสาร, Sentence case สำหรับชื่อบทความ/เว็บเพจ",
    "เลือกรูปแบบให้ตรงประเภทแหล่ง (วารสาร หนังสือ เว็บ วิทยานิพนธ์ ฯลฯ)",
    "รายการท้ายเล่มทุกชิ้นถูกอ้างในเนื้อหา และการสะกดชื่อ/ปีตรงกัน",
]

def _add_markup_runs(par, text, base_size=11, color=None, bold=False):
    """เติม run ลงย่อหน้า/เซลล์ โดยตีความ *...* เป็นตัวเอียง (ชื่อวารสาร/หนังสือ/เลขเล่ม)
    ใช้กับช่อง 'ค่าที่ถูกต้อง' ของตาราง APA7 เพื่อให้ตัวอย่างที่แก้แล้วแสดงตัวเอียงตรงจุด."""
    segs = str(text).split("*")
    if len(segs) == 1:            # ไม่มี markup — คืน run เดียว
        run = par.add_run(str(text))
        run.font.name = "TH Sarabun New"; run.font.size = Pt(base_size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return par
    for i, seg in enumerate(segs):
        if seg == "":
            continue
        run = par.add_run(seg)
        run.font.name = "TH Sarabun New"
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = (i % 2 == 1)   # ส่วนที่อยู่ระหว่างเครื่องหมาย * = เอียง
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return par

def _counts_of(rows):
    return {s: sum(1 for f in rows if f["severity"] == s)
            for s in ("critical", "major", "minor")}

def _fmt_duration(sec):
    try:
        sec = float(sec)
    except (TypeError, ValueError):
        return None
    if sec < 90:
        return f"{sec:.0f} วินาที" if sec >= 1 else f"{sec:.1f} วินาที"
    m, s = divmod(int(round(sec)), 60)
    return f"{m} นาที {s} วินาที" if s else f"{m} นาที"


_SHEET_RE = re.compile(r"แผ่น(?:งาน)?ที่\s*~?\s*(\d+)")
_PARA_RE = re.compile(r"para\s*(\d+)")
_ENTRY_RE = re.compile(r"รายการที่\s*(\d+)")


def _sheet_sort_key(finding):
    """เรียงตามลำดับที่คนจะเดินไล่แก้ในไฟล์จริง: แผ่น → ย่อหน้า → ความรุนแรง.

    เลขย่อหน้าจำเป็นต้องอยู่ในคีย์แม้จะไม่แสดงในรายงาน เพราะเลขแผ่นที่นับจาก
    marker ของ Word มักยุบรวมกัน (เล่มหนาหลายเล่มได้ 12 แผ่นทั้งเล่ม) ถ้าเรียงด้วย
    แผ่นอย่างเดียว finding ท้ายเล่มจะไปโผล่ก่อน finding ต้นเล่มเพราะแค่ severity
    สูงกว่า คนอ่านต้องเลื่อนกลับไปกลับมาในเอกสาร ซึ่งเป็นสิ่งที่รายงานนี้ควรกำจัด
    """
    loc = finding.get("location", "") or ""
    pages = [int(v) for v in _SHEET_RE.findall(loc)]
    paras = [int(v) for v in _PARA_RE.findall(loc)]
    entries = [int(v) for v in _ENTRY_RE.findall(loc)]
    return (min(pages) if pages else 10**9,
            min(paras) if paras else (min(entries) if entries else 10**9),
            SEV_ORDER.get(finding.get("severity"), 9))


def collapse_repeats(rows):
    """ยุบแถวคำผิดที่ใช้ 'คำค้นเดียวกัน + แก้เป็นค่าเดียวกัน' ให้เหลือแถวเดียว.

    คำสะกดผิดคำหนึ่งที่โผล่ 14 จุด แก้จบด้วย Replace All ครั้งเดียว การพิมพ์เป็น
    14 แถวไม่ได้ทำให้แก้ครบขึ้น แต่ทำให้รายงานดูยาวเกินจริงจนคนอ่านท้อ และกลบ
    finding อื่นที่ต้องลงมือแยกทีละจุด. เก็บแถวแรก (ตำแหน่งต้นเล่มที่สุด) แล้วบอก
    จำนวนที่พบไว้ในแถวนั้นแทน.
    """
    seen = {}
    out = []
    for f in rows:
        key = (_collapse(f.get("search", "")), f.get("correct", ""),
               f.get("category", ""))
        if not key[0]:
            # ไม่มีคำค้น: ยุบได้ก็ต่อเมื่อ "ปัญหาที่พบ" เขียนเหมือนกันเป๊ะ — แถวที่
            # อ่านแล้วแยกไม่ออกว่าคนละจุดกัน การพิมพ์ซ้ำ 29 ครั้งไม่ได้ช่วยให้แก้ครบ
            key = ("", _collapse(f.get("issue", "")), f.get("category", ""))
            if not key[1]:
                out.append(f)
                continue
        elif not key[1]:
            out.append(f)
            continue
        if key in seen:
            seen[key]["_repeats"] = seen[key].get("_repeats", 1) + 1
            continue
        seen[key] = f
        out.append(f)
    for f in out:
        n = f.get("_repeats", 1)
        if n > 1:
            # Replace All ใช้ได้เฉพาะกับปัญหาที่เป็น "ข้อความ" — ปัญหาการตั้งค่า
            # (ขนาดกระดาษ/ระยะขอบ) ต้องไปแก้ทีละ section บอกผิดแล้วผู้อ่านหาปุ่มไม่เจอ
            how = ("แก้ครั้งเดียวด้วย Replace All" if f.get("search")
                   else "ต้องไล่แก้ให้ครบทุกส่วน")
            f["issue"] = f"{f['issue']} (รูปแบบเดียวกันนี้พบ {n} จุดในเล่ม — {how})"
    return out


def verify_report_structure(path):
    """คืนรายการปัญหาโครงสร้างรายงาน; ไม่ต้อง render ภาพ/PDF."""
    doc = Document(path)
    issues = []
    if len(doc.tables) < 4:
        issues.append(f"คาดอย่างน้อย 4 ตาราง (Template, คำผิด, APA7, layout) แต่พบ {len(doc.tables)}")
        return issues
    expected = ["ลำดับ", "ระดับ", "ปัญหาที่พบ", "เกณฑ์/ที่มา", "วิธีแก้", "ค่าที่ถูกต้อง"]
    for index, table in enumerate(doc.tables[:3], 1):
        if len(table.columns) != 7:
            issues.append(f"ตารางที่ {index} มี {len(table.columns)} คอลัมน์ (ต้องเป็น 7)")
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for header in expected:
            if header not in headers:
                issues.append(f"ตารางที่ {index} ไม่มีหัวคอลัมน์ '{header}'")
        body_text = "\n".join(cell.text for row in table.rows[1:] for cell in row.cells)
        if "🟢 ผ่าน" in body_text:
            issues.append(f"ตารางที่ {index} อ้างว่า 'ผ่าน' ทั้งที่ coverage อาจยังไม่ครบ")
        # ช่อง "จุดที่ต้องแก้" ต้องไม่มีเลขแผ่น/เลขหน้าหลงเหลือ — ตั้งใจตัดออกเพื่อ
        # ไม่ให้ผู้อ่านไขว้เขวระหว่าง "ค้นด้วยข้อความ" กับ "ไล่ตามเลขหน้า"
        for row in table.rows[1:]:
            if _SHEET_RE.search(row.cells[1].text):
                issues.append(f"ตารางที่ {index} ยังมีเลขแผ่นค้างอยู่ในช่องจุดที่ต้องแก้")
                break
    headings = "\n".join(p.text for p in doc.paragraphs)
    if any(f"ตารางที่ {index}" not in headings for index in (1, 2, 3)):
        issues.append("ไม่พบหัวข้อแยกตาราง Template/คำผิด/APA7 ครบ")
    # รายงานต้องไม่มีบล็อกเวลา/ความยาวเล่ม/คำอธิบายวิธีใช้ — ตัดออกโดยตั้งใจ
    for banned in ("เวลาที่ใช้ในการตรวจสอบ", "วิธีใช้รายงานนี้", "ความยาวเล่ม"):
        if any(banned in h for h in headings):
            issues.append(f"รายงานยังมีบล็อก “{banned}” ซึ่งถูกตัดออกจากรูปแบบแล้ว")
    if "รายการที่ต้องตรวจสอบเพิ่มเติม (สคริปต์ตรวจแทนไม่ได้)" in headings:
        issues.append("ยังพบข้อความ '(สคริปต์ตรวจแทนไม่ได้)' ที่ต้องลบ")
    return issues

REPORT_FONT = "TH Sarabun New"


def _force_thai_fonts(doc, family=REPORT_FONT):
    """ตั้งฟอนต์ให้ครบทั้ง 4 ช่อง — ascii / hAnsi / **cs** / eastAsia

    python-docx ตั้งให้แค่ ascii+hAnsi เวลาสั่ง run.font.name แต่ Word ถือว่า
    **ข้อความไทยเป็น complex script** จึงไปหยิบฟอนต์จากช่อง w:cs แทน ผลคือรายงาน
    ที่ตั้งฟอนต์ไว้อย่างดีกลับแสดงไทยด้วยฟอนต์อื่น (บาง renderer ถึงขั้นไม่มีสระ/
    วรรณยุกต์ให้เห็น) ต้องเซ็ต w:cs ด้วยเสมอ ไม่งั้นบั๊กนี้จะโผล่ทุกครั้งที่เปิดคนละเครื่อง
    """
    def apply(rpr):
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        name = rf.get(qn("w:ascii")) or family
        for slot in ("ascii", "hAnsi", "cs", "eastAsia"):
            rf.set(qn("w:" + slot), name)
            # ถ้ามี *Theme ค้างอยู่ Word จะยึด theme แทนค่าที่เพิ่งตั้ง ต้องลบทิ้ง
            theme = qn("w:" + slot + "Theme") if slot != "cs" else qn("w:cstheme")
            if rf.get(theme) is not None:
                del rf.attrib[theme]
        # บอกภาษาไทยไว้ด้วย เพื่อให้ตัดคำ/จัดบรรทัดถูกต้อง
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:bidi"), "th-TH")

    roots = [doc.element, doc.styles.element]
    for section in doc.sections:
        roots.extend([section.header._element, section.footer._element])
    for root in roots:
        for rpr in root.iter(qn("w:rPr")):
            apply(rpr)
    # ค่าเริ่มต้นของเอกสาร เผื่อ run ที่ไม่มี rPr เลย
    defaults = doc.styles.element.find(qn("w:docDefaults"))
    if defaults is not None:
        rpr_default = defaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                rpr_default.insert(0, rpr)
            apply(rpr)


def _collapse(text):
    return re.sub(r"\s+", " ", str(text or "").replace("\u00a0", " ")).strip()


def uniquify_search(search, ctx, corpus, mode="both", limit=80):
    """ขยายคำค้นด้วยบริบทรอบ ๆ จนกว่าจะเจอแค่จุดเดียวในเล่ม

    คำค้นสั้น ๆ อย่าง "toward" หรือ "2.2" เจอเป็นสิบจุด ผู้ใช้ไม่รู้จะแก้จุดไหน —
    ยืดออกทีละนิดจากย่อหน้าจริงจนเหลือจุดเดียว. mode="left" = ยืดไปทางซ้ายเท่านั้น
    (ใช้กับ finding เว้นวรรคซ้อน ที่ตำแหน่งท้ายคำค้นต้องตรึงไว้ตรงช่องว่างที่เกิน)
    """
    needle = _collapse(search)
    if not needle or corpus.count(needle) <= 1:
        return search
    cctx = _collapse(ctx)
    pos = cctx.find(needle)
    if pos == -1:
        return search
    start, end = pos, pos + len(needle)
    while corpus.count(cctx[start:end]) > 1 and end - start < limit:
        moved = False
        if start > 0:
            start = max(0, start - 6); moved = True
        if mode == "both" and end < len(cctx):
            end = min(len(cctx), end + 6); moved = True
        if not moved:
            break
    candidate = cctx[start:end]
    cleaned = clean_word_start(candidate) if start > 0 else candidate
    if cleaned and corpus.count(_collapse(cleaned)) == 1:
        return cleaned.strip()
    return candidate.strip() if corpus.count(candidate) == 1 else search


def shrink_to_found(needle, corpus, floor=8):
    """ตัดท้ายคำค้นทีละคำจนกว่าจะเป็นข้อความที่มีอยู่จริงในเล่ม.

    ทำไมต้องมี: Word เก็บ caption/หัวข้อเป็นหลาย run และบางทีมี tab หรือ field
    คั่นกลาง ข้อความที่อ่านออกมาจึงเป็น "ภาพที่ 2.1แผนที่ตำแหน่ง…" ซึ่งไม่มีใน
    เล่มแบบตรงตัว. ส่วนหน้าของคำค้นเป็นส่วนที่ระบุจุดได้อยู่แล้ว ("ภาพที่ 2.1")
    การคืนส่วนที่ค้นเจอจริงจึงมีประโยชน์กว่าการติดธง ⚠ แล้วปล่อยให้ผู้อ่านหาเอง.
    """
    words = _collapse(needle).split(" ")
    while len(words) > 1:
        words.pop()
        cand = " ".join(words)
        if len(cand) >= floor and corpus.count(cand) >= 1:
            return cand
    # ไม่มีช่องว่างให้ตัด (ข้อความไทยยาวติดกัน) — ไล่ตัดทีละอักขระ
    text = _collapse(needle)
    for end in range(len(text) - 1, floor - 1, -1):
        cand = text[:end].strip()
        if corpus.count(cand) >= 1:
            return cand
    return ""


def occurrence_ordinal(order, key, needle):
    """คำค้นซ้ำหลายจุด — จุดของ finding นี้คือจุดที่เท่าไรเมื่อกด Find Next จากต้นเล่ม"""
    if not key or not needle:
        return None
    before = 0
    for part_key, text in order:
        if part_key == key:
            return before + 1
        before += _collapse(text).count(needle)
    return None


def refine_searches(findings, index, corpus):
    """หลังรวม finding แล้ว: ทำให้ทุกคำค้นชี้จุดเดียว หรือบอกว่าเป็นจุดที่เท่าไร"""
    order = index.get(ORDER_KEY, [])
    for f in findings:
        search = f.get("search")
        if not search:
            continue
        needle = _collapse(search)
        count = corpus.count(needle) if corpus else 0
        if count == 0 and f.get("_cite") and corpus:
            # ตัวตรวจเขียนการอ้างอิงในรูป "Author (Year)" แต่ในเล่มจริงอาจเป็น
            # "(Author, Year)" หรือ "Author and B (Year)" — ไล่หารูปที่มีจริงมาเป็นคำค้น
            author, year = f["_cite"]
            for variant in (f"{author} ({year})", f"({author}, {year})",
                            f"{author}, {year}", author):
                c = corpus.count(_collapse(variant))
                if c:
                    f["search"] = variant
                    needle, count = _collapse(variant), c
                    if variant == author:
                        f["search_note"] = (f"ในเล่มไม่พบรูป “{search}” ตรง ๆ — ค้นด้วยชื่อผู้แต่ง "
                                            f"“{author}” แล้วดูจุดที่อ้างปี {year}")
                    break
        if count == 0 and corpus:
            shorter = shrink_to_found(needle, corpus)
            if shorter:
                # caption/หัวข้อมักถูกประกอบจากหลาย run ("ภาพที่ 2.1" + tab + ชื่อภาพ)
                # ข้อความที่ตัวตรวจอ่านได้จึงไม่มีในเล่มแบบตรงตัว ค้นแล้วไม่เจอ.
                # ตัดท้ายทิ้งจนเหลือส่วนที่มีจริง ดีกว่าให้ผู้อ่านค้นข้อความที่ไม่มีอยู่
                f["search"] = shorter
                f["search_note"] = (f"ข้อความเต็มในเล่มถูกแบ่งเป็นหลายช่วง — ค้นด้วย “{shorter}” "
                                    "แล้วดูบรรทัดนั้นทั้งบรรทัด")
                needle, count = _collapse(shorter), corpus.count(_collapse(shorter))
            elif re.search(r"[<>]", needle):
                # finding แบบ "ไม่มีสิ่งนี้" (เช่น ภาพไม่มี caption) — ไม่มีข้อความให้ค้น
                # อยู่แล้ว การใส่แม่แบบ "ภาพที่ <บท>.<ลำดับ>" ลงช่องคำค้นทำให้ผู้อ่าน
                # ไปค้นสิ่งที่ไม่มีวันเจอ บอกตรง ๆ ว่าต้องหาอย่างไรแทน
                f["search"] = ""
                f["search_note"] = ("ยังไม่มีข้อความนี้ในเล่ม (นี่คือสิ่งที่ต้องเพิ่ม) — "
                                    "เลื่อนหาภาพ/ตารางที่ยังไม่มีคำบรรยายตามลำดับในบท")
                continue
        if count <= 1:
            continue
        if f.get("replace_all"):
            # การแก้เป็นแบบกวาดทั้งเล่ม (Replace All) — ซ้ำหลายจุดคือเรื่องดี บอกไว้เฉย ๆ
            f["search_note"] = f"คำนี้ปรากฏ {count} จุด — Replace All จัดการครบในครั้งเดียว"
            continue
        if f.get("search_kind") == "example":
            # ตัวอย่างจุดที่ใช้สไตล์ — เปิดดูจุดไหนก็ได้เพราะแก้ที่ตัวสไตล์ครั้งเดียว
            f["search_note"] = f"ข้อความนี้ปรากฏ {count} จุด — เปิดดูจุดไหนก็ได้ ทุกจุดใช้สไตล์เดียวกัน"
            continue
        mode = f.get("_expand", "both")
        if mode == "none":
            # คำค้นเป็นข้อความผิดตรงตัว (มีเคาะซ้อน/อักขระพิเศษข้างใน) — ระบบยืด/สร้างใหม่
            # จะทำอักขระพิเศษหาย กลายเป็นค้นไม่เจอ ปล่อยไว้ตามเดิม
            continue
        ctx = f.get("_ctx") or ""
        better = uniquify_search(search, ctx, corpus, mode=mode)
        if _collapse(better) != needle and corpus.count(_collapse(better)) == 1:
            f["search"] = better
            continue
        ordinal = occurrence_ordinal(order, f.get("_key"), needle)
        if ordinal:
            f["search_note"] = (f"คำนี้ปรากฏ {count} จุดในเล่ม — จุดที่ต้องแก้คือ**จุดที่ {ordinal}** "
                                f"(กด Find Next ไล่ไปจากต้นเล่ม)")
        else:
            f["search_note"] = f"คำนี้ปรากฏ {count} จุดในเล่ม — เทียบกับข้อความในช่อง 'ปัญหาที่พบ' ประกอบ"


def _guess_source(payload, out_path):
    """หาไฟล์ต้นฉบับ .docx จาก payload — ไว้ดึงข้อความมาทำคำค้นให้ finding ที่ไม่ได้ยกข้อความมา"""
    name = str(payload.get("file") or "").strip()
    if not name:
        return None
    for candidate in (name, os.path.join(os.path.dirname(os.path.abspath(out_path)), os.path.basename(name)),
                      os.path.join(os.getcwd(), os.path.basename(name))):
        if candidate and os.path.exists(candidate):
            return candidate
    return None


def build_report(payload, out_path, title=None, time_note=None, source_docx=None,
                 font=REPORT_FONT):
    """payload = {file, profile, profile_label, counts, findings[]}  (หรือรวมหลายชุด)
    finding แต่ละรายการมี domain = template, spelling หรือ apa7 —
    ใช้แยกเป็น 3 ตารางรายงานคนละส่วน ไม่ปนกัน"""
    source = source_docx or _guess_source(payload, out_path)
    index = build_search_index(source)
    corpus = corpus_from_index(index)
    findings = [normalize(f, index) for f in payload.get("findings", [])]
    refine_searches(findings, index, corpus)
    # ยืนยันว่าคำค้นแต่ละอันมีอยู่จริงในเล่ม — ที่ยืนยันไม่ได้ต้องติดธงไว้ ไม่ใช่ปล่อยให้
    # ผู้อ่านเสียเวลาค้นข้อความที่ไม่มีอยู่ (มักแปลว่า finding นั้นเองน่าสงสัย)
    unverified = 0
    for f in findings:
        if not f.get("search"):
            continue
        needle = re.sub(r"\s+", " ", f["search"].rstrip("…")).strip()
        if corpus and needle and needle not in corpus:
            f["search_ok"] = False
            unverified += 1
        else:
            f["search_ok"] = True
    # เอาเฉพาะที่ต้องแก้จริง (ไม่รวม info/ok) และเรียงตามแผ่นงานก่อนระดับความรุนแรง
    actionable = [f for f in findings if f["severity"] in ("critical", "major", "minor")]
    template_rows = collapse_repeats(
        sorted((f for f in actionable if f["domain"] == "template"),
               key=_sheet_sort_key))
    spelling_rows = collapse_repeats(
        sorted((f for f in actionable if f["domain"] == "spelling"),
               key=_sheet_sort_key))
    apa7_rows = sorted((f for f in actionable if f["domain"] == "apa7"),
                       key=_sheet_sort_key)

    counts = payload.get("counts") or _counts_of(actionable)
    counts_template = _counts_of(template_rows)
    counts_spelling = _counts_of(spelling_rows)
    counts_apa7 = _counts_of(apa7_rows)

    doc = Document()
    doc.styles["Normal"].font.name = "TH Sarabun New"
    doc.styles["Normal"].font.size = Pt(14)
    for s in doc.sections:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Inches(11.69), Inches(8.27)
        s.left_margin = s.right_margin = Inches(0.6)
        s.top_margin = s.bottom_margin = Inches(0.7)

    # ---- หัวรายงาน ----
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("รายงานการตรวจสอบวิทยานิพนธ์ (TULIBS Thesis .docx Checker)")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ตรวจรูปแบบเทมเพลต คำผิด และการอ้างอิง APA 7 แยก 3 ตาราง")
    sr.italic = True; sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x59,0x59,0x59)
    if title:
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t.add_run(title); tr.italic = True; tr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.add_run("ไฟล์ที่ตรวจ: ").bold = True
    meta.add_run(str(payload.get("file", "-")))
    meta.add_run("    |  เทมเพลต: ").bold = True
    meta.add_run(str(payload.get("profile_label") or payload.get("profile", "-")))
    meta.add_run("    |  วันที่ตรวจ: ").bold = True
    meta.add_run(datetime.date.today().isoformat())
    # ไม่แสดงความยาวเล่ม เวลาที่ใช้ตรวจ หรือคำอธิบายวิธีใช้รายงาน — สามอย่างนี้เป็น
    # ข้อมูลของ "ผู้ตรวจ" ไม่ใช่ของผู้ที่ต้องลงมือแก้ไฟล์ การมีอยู่ทำให้ผู้อ่านต้อง
    # เลื่อนผ่านก่อนถึงตารางที่ใช้งานจริง และหัวตารางบอกวิธีใช้อยู่แล้วในตัว

    summ = doc.add_paragraph()
    summ.add_run("สรุปรวมทั้งฉบับ: ").bold = True
    summ.add_run(f"🔴 {counts.get('critical',0)} ต้องแก้ก่อนส่ง   "
                 f"🟠 {counts.get('major',0)} ควรแก้   "
                 f"🟡 {counts.get('minor',0)} เล็กน้อย")
    s1 = doc.add_paragraph()
    s1.add_run("• รูปแบบเทมเพลต (Template): ").bold = True
    s1.add_run(f"🔴 {counts_template.get('critical',0)}  🟠 {counts_template.get('major',0)}  "
               f"🟡 {counts_template.get('minor',0)}   —   ").italic = False
    s1.add_run(readiness(counts_template)).bold = True
    s2 = doc.add_paragraph()
    s2.add_run("• คำผิด/อักขระผิดปกติ: ").bold = True
    s2.add_run(f"🔴 {counts_spelling.get('critical',0)}  🟠 {counts_spelling.get('major',0)}  "
               f"🟡 {counts_spelling.get('minor',0)}   —   ")
    s2.add_run(readiness(counts_spelling)).bold = True
    s3 = doc.add_paragraph()
    s3.add_run("• การอ้างอิง APA 7: ").bold = True
    s3.add_run(f"🔴 {counts_apa7.get('critical',0)}  🟠 {counts_apa7.get('major',0)}  "
               f"🟡 {counts_apa7.get('minor',0)}   —   ")
    s3.add_run(readiness(counts_apa7)).bold = True

    # ---- helper: ตารางแก้ไขทีละจุด ----
    def correction_table(heading, rows, loc_header="ตำแหน่งที่ต้องแก้",
                         empty_msg="ไม่พบจุดที่ต้องแก้จากการตรวจอัตโนมัติ",
                         heading_color=NAVY):
        hh = doc.add_paragraph()
        hr = hh.add_run(heading); hr.bold = True; hr.font.size = Pt(16)
        hr.font.color.rgb = heading_color
        headers = ["ลำดับ", loc_header, "ระดับ", "ปัญหาที่พบ", "เกณฑ์/ที่มา", "วิธีแก้", "ค่าที่ถูกต้อง"]
        widths = [0.4, 1.95, 0.95, 1.95, 1.8, 1.75, 1.6]
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, htext in enumerate(headers):
            _set(table.rows[0].cells[i], htext, bold=True, color="FFFFFF", size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade(table.rows[0].cells[i], "1F3864")
            table.rows[0].cells[i].width = Inches(widths[i])
        if not rows:
            row = table.add_row().cells
            _set(row[0], "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set(row[1], "—"); _set(row[2], "ℹ️ ยังไม่พบ", color="2E74B5")
            _set(row[3], empty_msg); _set(row[4], "—")
            _set(row[5], "ยืนยัน checklist ก่อนส่ง"); _set(row[6], "—")
        for idx, f in enumerate(rows, 1):
            row = table.add_row().cells
            sev_label, sev_color = SEV_TH.get(f["severity"], (f["severity"], "000000"))
            _set(row[0], idx, align=WD_ALIGN_PARAGRAPH.CENTER)
            _fill_location(row[1], f)
            _set(row[2], sev_label, color=sev_color, bold=True)
            cat = f"[{f['category']}] " if f["category"] else ""
            _set(row[3], cat + f["issue"])
            _set(row[4], f["criterion"] or "-")
            _set(row[5], f["fix"] or "-")
            # ช่อง "ค่าที่ถูกต้อง": render ตัวอย่างที่แก้แล้ว โดยตีความ *...* เป็นตัวเอียง
            # (ชื่อวารสาร/หนังสือ/เลขเล่ม) — สำหรับ APA7 จะได้เห็นรูปที่ถูกพร้อมตัวเอียงในช่องนี้เลย
            row[6].text = ""
            _add_markup_runs(row[6].paragraphs[0],
                             f["correct"] or (f["fix"] or "-"),
                             base_size=11, color="548235", bold=True)
            for i in range(len(headers)):
                row[i].width = Inches(widths[i])
            if idx % 2 == 0:
                for c in row:
                    _shade(c, GREY)

    # ---- ตารางที่ 1: แก้ไขรูปแบบ/หน้ากระดาษ (Template) ----
    correction_table("ตารางที่ 1 — แก้ไขรูปแบบเทมเพลต/หน้ากระดาษทีละจุด (Template Format)",
                      template_rows,
                      loc_header="จุดที่ต้องแก้ — ก๊อปไปค้นใน Word (Ctrl+F)",
                      empty_msg="ไม่พบจุดที่ต้องแก้ด้านรูปแบบเทมเพลต")

    doc.add_paragraph()

    # ---- ตารางที่ 2: คำผิด ----
    correction_table("ตารางที่ 2 — แก้ไขคำผิดและอักขระผิดปกติทีละจุด (Spelling)",
                      spelling_rows,
                      loc_header="จุดที่ต้องแก้ — ก๊อปไปค้นใน Word (Ctrl+F)",
                      empty_msg="ไม่พบคำผิดหรืออักขระผิดปกติจากกฎที่ตรวจอัตโนมัติ",
                      heading_color=RGBColor(0x70, 0x30, 0xA0))

    doc.add_paragraph()

    # ---- ตารางที่ 3: การอ้างอิง APA 7 ----
    correction_table("ตารางที่ 3 — แก้ไขการอ้างอิง APA 7 ทีละรายการ (APA7 Reference Check)",
                      apa7_rows,
                      loc_header="รายการที่ต้องแก้ — ก๊อปไปค้นใน Word (Ctrl+F)",
                      empty_msg="ไม่พบจุดผิดพลาด APA 7 จากรายการอ้างอิงที่ตรวจ",
                      heading_color=RGBColor(0x1F, 0x64, 0x3A))

    # ---- ตารางการตั้งค่า layout / หน้ากระดาษที่ถูกต้อง ----
    doc.add_paragraph()
    lp = doc.add_paragraph()
    lr = lp.add_run("รูปแบบการตั้งค่าหน้ากระดาษที่กำหนดไว้ (เทมเพลต "
                    + str(payload.get("profile", "-")) + ")")
    lr.bold = True; lr.font.size = Pt(16); lr.font.color.rgb = NAVY
    ltab = doc.add_table(rows=1, cols=2)
    ltab.style = "Table Grid"
    ltab.alignment = WD_TABLE_ALIGNMENT.CENTER
    lw = [3.2, 6.0]
    for i, htext in enumerate(["รายการตั้งค่า", "ค่าที่ถูกต้อง"]):
        _set(ltab.rows[0].cells[i], htext, bold=True, color="FFFFFF", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(ltab.rows[0].cells[i], "1F3864")
        ltab.rows[0].cells[i].width = Inches(lw[i])
    for ri, (k, v) in enumerate(layout_rows(payload.get("profile", "thai")), 1):
        row = ltab.add_row().cells
        _set(row[0], k, bold=True)
        _set(row[1], v)
        row[0].width = Inches(lw[0]); row[1].width = Inches(lw[1])
        if ri % 2 == 0:
            for c in row:
                _shade(c, GREY)

    # ---- checklist ตรวจด้วยตา ----
    doc.add_paragraph()
    cp = doc.add_paragraph()
    cr = cp.add_run("รายการที่ต้องตรวจสอบเพิ่มเติม")
    cr.bold = True; cr.font.size = Pt(16); cr.font.color.rgb = NAVY
    for item in VISUAL_ITEMS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ " + item)

    doc.add_paragraph()
    cp2 = doc.add_paragraph()
    cr2 = cp2.add_run("รายการตรวจสอบ APA 7 เพิ่มเติม (ต้องอ่านเนื้อหา)")
    cr2.bold = True; cr2.font.size = Pt(16); cr2.font.color.rgb = RGBColor(0x1F, 0x64, 0x3A)
    for item in APA7_VISUAL_ITEMS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ " + item)

    note = doc.add_paragraph()
    nr = note.add_run("หมายเหตุ: ผลตรวจอัตโนมัติอ่านจากค่าจริงในไฟล์ .docx (styles.xml/sectPr/run) "
                      "ไม่ใช่จากการดูภาพ; การตรวจไฟล์ .docx แม่นกว่า PDF จึงควรตรวจก่อนแปลงเป็น PDF. "
                      "ผลเชิงความหมายของ APA 7 ต้องมาจาก reading pass ที่เทียบ TULIBS APA7 (ไทย) / APA Style (อังกฤษ); "
                      "เกณฑ์ template เป็น pinned ruleset ตาม profile/revision ที่แสดงด้านบน ไม่ใช่คำยืนยันว่าเป็น revision ล่าสุด; "
                      "ตาราง Template, คำผิด และ APA7 แยกกันโดยเจตนา ไม่ปนหมวดกัน")
    nr.italic = True; nr.font.size = Pt(12); nr.font.color.rgb = RGBColor(0x59,0x59,0x59)

    _force_thai_fonts(doc, font or REPORT_FONT)
    doc.save(out_path)
    structural_issues = verify_report_structure(out_path)
    if structural_issues:
        raise ValueError("รายงาน DOCX โครงสร้างไม่ผ่าน: " + "; ".join(structural_issues))
    actionable_all = template_rows + spelling_rows + apa7_rows
    return {
        "path": out_path,
        "searchable": sum(1 for f in actionable_all if f.get("search") and f.get("search_ok") is not False),
        "unverified": sum(1 for f in actionable_all if f.get("search") and f.get("search_ok") is False),
        "document_level": sum(1 for f in actionable_all if not f.get("search")),
    }

def load_json(path):
    with open(path, encoding="utf-8") as fh:
        data = json.load(fh)
    if isinstance(data, list):          # bare findings list
        return {"findings": data}
    return data

def merge(payloads):
    base = dict(payloads[0])
    all_f, seen = [], set()
    for p in payloads:
        for finding in p.get("findings", []):
            normalized = normalize(finding)
            key = tuple(normalized[k] for k in
                        ("domain", "severity", "location", "category", "issue"))
            if key in seen:
                continue
            seen.add(key)
            all_f.append(finding)
    base["findings"] = all_f
    elapsed = [p.get("elapsed_sec") for p in payloads]
    if all(isinstance(value, (int, float)) for value in elapsed):
        base["elapsed_sec"] = round(sum(elapsed), 2)
    # recompute counts across merged
    norm = [normalize(f) for f in all_f]
    base["counts"] = {s: sum(1 for f in norm if f["severity"] == s)
                      for s in ("critical", "major", "minor")}
    return base

def main():
    ap = argparse.ArgumentParser(description="สร้างรายงานการแก้ไข .docx จากผลตรวจ JSON")
    ap.add_argument("json", nargs="+", help="ไฟล์ JSON ผลตรวจ (docx/pdf/visual) หนึ่งไฟล์ขึ้นไป")
    ap.add_argument("-o", "--out", default="thesis_format_report.docx")
    ap.add_argument("--title", default=None)
    ap.add_argument("--docx", dest="source_docx", default=None,
                    help="ไฟล์ต้นฉบับ .docx — ใช้ดึงข้อความจริงมาเป็นคำค้นในช่อง 'จุดที่ต้องแก้' "
                         "(ไม่ใส่ก็ได้ สคริปต์จะหาจากชื่อไฟล์ใน JSON ให้เอง)")
    ap.add_argument("--font", default=REPORT_FONT,
                    help=f"ฟอนต์ของตัวรายงาน (ค่าเริ่มต้น {REPORT_FONT}) — ตั้งครบทั้งช่อง "
                         "ascii/hAnsi/cs/eastAsia เพราะ Word ถือว่าไทยเป็น complex script "
                         "ถ้าเครื่องปลายทางไม่มีฟอนต์นี้ ให้ใส่ชื่อฟอนต์ไทยที่มี เช่น Sarabun")
    ap.add_argument("--time", dest="time_note", default=None,
                    help="เวลารวมที่ใช้ตรวจทั้งหมด (รวมรอบอ่าน APA ด้วยตา) เช่น 'ประมาณ 3 นาที'")
    a = ap.parse_args()
    payloads = []
    for path in a.json:
        if not os.path.exists(path):
            print(f"WARNING: ข้ามไฟล์ผลตรวจที่ไม่มีอยู่: {path}", file=sys.stderr)
            continue
        payloads.append(load_json(path))
    if not payloads:
        sys.exit("ไม่พบไฟล์ JSON ที่อ่านได้")
    payload = payloads[0] if len(payloads) == 1 else merge(payloads)
    stats = build_report(payload, a.out, a.title, a.time_note, a.source_docx, a.font)
    print(f"เขียนรายงานแล้ว: {a.out}")
    if isinstance(stats, dict) and stats.get("searchable") is not None:
        print(f"คำค้นในช่อง 'จุดที่ต้องแก้': ยืนยันว่าค้นเจอจริง {stats['searchable']} รายการ"
              + (f" · ค้นไม่พบ {stats['unverified']} รายการ (ติดธง ⚠ ไว้ในรายงานแล้ว)"
                 if stats.get("unverified") else "")
              + (f" · เป็นปัญหาระดับเอกสาร {stats['document_level']} รายการ"
                 if stats.get("document_level") else ""))

if __name__ == "__main__":
    main()
