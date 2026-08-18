#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""split_docx.py — ตัดเล่มวิทยานิพนธ์ .docx เป็น "ชิ้นงาน" ที่ไม่ทับกัน
สำหรับแจกให้ agent หลายตัวตรวจขนานกัน

ทำไมต้องมีขั้นนี้
----------------
ถ้าให้ agent ทุกตัวเปิดไฟล์ .docx ทั้งเล่มเอง แต่ละตัวจะดูดเนื้อหาทั้ง 100+ หน้า
เข้า context ซ้ำกัน — ช้า แพง และพอ context เต็ม agent จะเริ่มลืมกฎที่สั่งไว้ตอนต้น
สคริปต์นี้อ่านไฟล์ *ครั้งเดียว* แล้วเขียนออกมาเป็นไฟล์ข้อความแยกส่วน โดยแต่ละบรรทัด
ติดป้ายตำแหน่งไว้ให้แล้ว (`[sheet12|¶345]` = แผ่นที่ 12 ย่อหน้าที่ 345) ทำให้ agent
อ่านเฉพาะส่วนของตัวเอง แต่ยังอ้างตำแหน่งใน finding ได้แม่นเท่าเดิม

ชิ้นงานที่ตัดออกมา
-----------------
  frontmatter.txt  ปก → ก่อน "บทที่ 1"/"CHAPTER 1"   (ปก/อนุมัติ/บทคัดย่อ/สารบัญ)
  body.txt         "บทที่ 1" → ก่อนรายการอ้างอิง      (เนื้อบท/caption/ตาราง/ภาพ)
  references.txt   เฉพาะรายการอ้างอิง/บรรณานุกรม       (ไม่พ่วงภาคผนวกให้ agent APA)
  backmatter.txt   ภาคผนวก/ประวัติผู้เขียน             (งานตรวจ format ของ Agent 2)
  intext.txt       ประโยคที่มีการอ้างอิงในเนื้อหา       (ดึงจาก body ให้ agent APA7
                                                       ไม่ต้องอ่าน body ทั้งก้อน)
  manifest.json    ขอบเขต/เลขหน้า/สถิติของแต่ละชิ้น + ชื่อไฟล์ที่ agent ต้องอ่าน

การหาเส้นแบ่งใช้หัวข้อจริงในเล่ม ไม่ใช่การเดาจากเลขหน้า และข้ามบรรทัดสารบัญ
(บรรทัดที่มี tab นำเลขหน้า) เพื่อไม่ให้ "รายการอ้างอิง" ในสารบัญถูกนับเป็นจุดเริ่มส่วน

ใช้:
    python3 scripts/split_docx.py thesis.docx -o work/
    python3 scripts/split_docx.py thesis.docx -o work/ --max-chars 60000
"""
import sys, os, re, json, argparse

try:
    from docx import Document
except ImportError:
    sys.exit("ต้องติดตั้ง python-docx ก่อน:  pip install python-docx")

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from check_docx import detect_profile, PROFILES, resolve_page_data, unpack_page_data
from docx.oxml.ns import qn
from docx.table import Table

# เส้นแบ่งส่วน — ยอมรับทั้งไทยและอังกฤษ และรูปแบบที่นักศึกษามักพิมพ์เพี้ยนเล็กน้อย
CHAPTER_START = re.compile(r"^\s*(บทที่\s*1(?![0-9])|CHAPTER\s*(1|I)(?![0-9A-Z]))", re.I)
REF_START     = re.compile(r"^\s*(รายการอ้างอิง|บรรณานุกรม|เอกสารอ้างอิง|REFERENCES?|BIBLIOGRAPHY)\s*$", re.I)
TAIL_START    = re.compile(r"^\s*(ภาคผนวก|APPENDI(?:X|CES)|ประวัติผู้เขียน|BIOGRAPHY)(?:\s|$)", re.I)

# การอ้างอิงในเนื้อหา: (ผู้แต่ง, ปี) / (Author, 2019) / ผู้แต่ง (2563)
INTEXT = re.compile(r"\(([^()]{0,80}?),?\s*(?:19|20|25)\d{2}[^()]{0,20}\)|[฀-๿A-Za-z\.\s]{2,40}\s\((?:19|20|25)\d{2}\)")


def para_lines(doc, page_data):
    """คืน (location, page, style, text) ตามลำดับจริง รวมข้อความใน table cells.

    ``Document.paragraphs`` ไม่รวมข้อความในตาราง แต่หน้าอนุมัติและตารางข้อมูล
    บทคัดย่อของเทมเพลตอยู่ใน table cells; ถ้าตัดเฉพาะ body paragraphs Agent 1
    จะไม่มีข้อมูลสำคัญให้ตรวจเลย.
    """
    page_by_index, _ = unpack_page_data(page_data, doc)
    location_pages = page_data.get("location_pages", {}) if isinstance(page_data, dict) else {}
    out = []
    page, body_pi, table_i = 1, 0, 0
    P, TBL = qn("w:p"), qn("w:tbl")
    body_paragraphs = doc.paragraphs  # cache: Document.paragraphs สร้าง list ใหม่ทุกครั้ง

    def count_breaks(el):
        total = 0
        for node in el.iter():
            if node.tag == qn("w:lastRenderedPageBreak"):
                total += 1
            elif node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
                total += 1
        return total

    def add_paragraph(p, loc, pg):
        t = (p.text or "").strip()
        if not t:
            return
        try:
            st = p.style.name if p.style is not None else ""
        except Exception:
            st = ""
        out.append((loc, pg, st, t))

    for child in doc.element.body.iterchildren():
        if child.tag == P:
            p = body_paragraphs[body_pi]
            pg = page_by_index[body_pi] if body_pi < len(page_by_index) else page
            add_paragraph(p, f"¶{body_pi}", pg)
            body_pi += 1
            page += count_breaks(child)
        elif child.tag == TBL:
            table_i += 1
            table = Table(child, doc._body)
            seen_cells = set()
            local_page = page
            for ri, row in enumerate(table.rows, 1):
                for ci, cell in enumerate(row.cells, 1):
                    if cell._tc in seen_cells:
                        continue
                    seen_cells.add(cell._tc)
                    for pi, p in enumerate(cell.paragraphs, 1):
                        logical = f"ตาราง {table_i} แถว {ri} คอลัมน์ {ci} ย่อหน้า {pi}"
                        pg = location_pages.get(logical, local_page)
                        add_paragraph(p, f"tbl{table_i}:r{ri}:c{ci}:¶{pi}", pg)
                        local_page += count_breaks(p._p)
            page += count_breaks(child)
    return out


def find_boundaries(lines):
    """หา index ที่บท 1, รายการอ้างอิง และภาคผนวก/ประวัติเริ่ม

    ข้ามบรรทัดสารบัญ: สารบัญพิมพ์ 'บทที่ 1 บทนำ ....... 1' ซึ่ง match pattern
    เหมือนหัวบทจริง ถ้าไม่กรองออกจะได้เส้นแบ่งผิดตั้งแต่หน้าสารบัญ วิธีกรองคือ
    ดูว่าบรรทัดนั้นลงท้ายด้วยเลขหน้าหรือมีจุดไข่ปลา/tab คั่นหรือไม่
    """
    toc_like = re.compile(r"(\.{3,}|\t)\s*\(?\d+\)?\s*$")
    chap = ref = tail = None
    for n, (_, _, _, t) in enumerate(lines):
        if toc_like.search(t):
            continue
        if chap is None and CHAPTER_START.match(t):
            chap = n
        if REF_START.match(t):
            ref = n          # เอาอันหลังสุด — เผื่อคำนี้โผล่ในสารบัญแบบไม่มีจุดไข่ปลา
            tail = None      # เริ่มหา appendix ใหม่หลัง reference heading ตัวล่าสุด
        if ref is not None and n > ref and tail is None and TAIL_START.match(t):
            tail = n
    if chap is None:
        chap = 0
    if ref is None or ref <= chap:
        ref = len(lines)
    if tail is None or tail <= ref:
        tail = len(lines)
    return chap, ref, tail


# ---- คัดเฉพาะย่อหน้าที่ "คนต้องดู" --------------------------------------------
# เนื้อหาส่วนใหญ่ของวิทยานิพนธ์คือร้อยแก้ว ซึ่งกฎที่ agent ต้องตรวจ (ลำดับหัวข้อ,
# caption, ตารางตกหน้า, ความสม่ำเสมอของหัวข้อ) ไม่ได้ใช้เลย — ส่วนคำผิด/ฟอนต์/ระยะขอบ
# สคริปต์ตรวจไปแล้วทั้งหมด. วัดกับเล่มจริง 135 หน้า: ร้อยแก้วกินพื้นที่ 81% ของ slice
# แต่ไม่ก่อให้เกิด finding เลย ตัดออกแล้ว agent อ่านน้อยลง 5 เท่าโดยไม่เสียความสามารถ
SUSPECT_TEXT = re.compile(
    r"^\s*("
    r"บทที่\s|CHAPTER\s|ภาคผนวก|APPENDI(X|CES)|ประวัติผู้เขียน|BIOGRAPHY|"
    r"ตารางที่\s|ภาพที่\s|แผนภาพที่\s|Table\s+\d|Figure\s+\d|"
    r"\d+\.\d+|\(\d+\)|[ก-ฮ]\.\d+|"
    r"หมายเหตุ|Note\.|ที่มา|Source"
    r")", re.I)
# สไตล์ที่บ่งว่าเป็นหัวข้อ/คำบรรยาย — ต้องไม่จับ TU_Paragraph_Normal ซึ่งคือร้อยแก้ว
# ทั้งเล่ม (ตอนแรกเขียน r"TU_" ไว้แล้วตัดได้แค่ 2% เพราะทุกย่อหน้าขึ้นต้นด้วย TU_)
SUSPECT_STYLE = re.compile(r"heading|caption|title|chapter|toc|สารบัญ", re.I)
PROSE_STYLE = re.compile(r"para|normal|body|เนื้อหา", re.I)
SHORT_LINE = 60          # หัวข้อ/ป้าย/เซลล์ตาราง มักสั้นกว่านี้


TBL_CELL = re.compile(r"tbl(\d+):r(\d+):c(\d+)")
CONT_MARK = re.compile(r"\(\s*(ต่อ|continued)\s*\)|ตารางที่|ภาพที่|Table\s+\d|Figure\s+\d", re.I)


def is_suspect(loc, style, text, page=None, prev_page=None):
    m = TBL_CELL.match(loc)
    if m:
        # ในตารางยาว ๆ (แบบสอบถามในภาคผนวก) agent ไม่ต้องอ่านทุกแถว — สิ่งที่ต้องตรวจคือ
        # หัวตารางกับจุดที่ตารางตกหน้าว่ามีการทวนหัวตารางและคำว่า (ต่อ) หรือไม่
        row = int(m.group(2))
        if row <= 2 or CONT_MARK.search(text):
            return True
        return page is not None and prev_page is not None and page != prev_page
    if style and SUSPECT_STYLE.search(style) and not PROSE_STYLE.search(style):
        return True
    if len(text) <= SHORT_LINE:               # บรรทัดสั้น = หัวข้อ ป้าย หรือคำโดด
        return True
    return bool(SUSPECT_TEXT.match(text))


def keep_suspects(rows):
    """คืน rows ที่คัดแล้ว โดยแทนที่ช่วงร้อยแก้วที่ตัดออกด้วยบรรทัดสรุปหนึ่งบรรทัด

    ต้องเหลือร่องรอยไว้ ไม่ใช่ลบเงียบ ๆ เพราะ agent ยังต้องรู้ว่าระหว่างหัวข้อสองอัน
    มีเนื้อหาคั่นอยู่กี่ย่อหน้าและกินแผ่นไหนบ้าง (ใช้ดูหัวข้อลอยท้ายหน้า/บทขึ้นหน้าใหม่)
    """
    out, dropped, first_page, last_page = [], 0, None, None

    def flush():
        nonlocal dropped, first_page, last_page
        if dropped:
            span = (f"แผ่น {first_page}" if first_page == last_page
                    else f"แผ่น {first_page}–{last_page}")
            out.append(("—", last_page, "",
                        f"[ตัดร้อยแก้ว {dropped} ย่อหน้า ({span}) — สคริปต์ตรวจคำผิด/ฟอนต์ให้แล้ว]"))
        dropped, first_page, last_page = 0, None, None

    prev_page = None
    for loc, pg, style, text in rows:
        suspect = is_suspect(loc, style or "", text, pg, prev_page)
        prev_page = pg
        if suspect:
            flush()
            out.append((loc, pg, style, text))
        else:
            dropped += 1
            first_page = first_page or pg
            last_page = pg
    flush()
    return out


# ---- เดาประเภทแหล่งอ้างอิง ----------------------------------------------------
# ต้นทุนของงาน APA7 คิดต่อ "รายการ" เพราะแต่ละรายการต้องเปิดคู่มือคนละไฟล์ตามประเภท
# แหล่ง ถ้าปล่อยให้ agent ไล่เดาเองทีละรายการ เล่มหนึ่งจะเปิดไฟล์คู่มือหลายสิบครั้ง
# จัดกลุ่มให้ก่อนแล้วส่งเป็นก้อนตามประเภท → เปิดคู่มือประเภทละครั้งเดียว
REF_TYPES = [
    ("วิทยานิพนธ์", re.compile(r"วิทยานิพนธ์|สารนิพนธ์|การค้นคว้าอิสระ|thesis|dissertation", re.I),
     "ref-06-วิทยานิพนธ์ปริญญามหาบัณฑิตและดุษฎีบัณฑิต.md", "Unpublished dissertation or thesis references.md"),
    ("ประชุมวิชาการ", re.compile(r"การประชุม|ประชุมวิชาการ|proceedings?|conference|symposium", re.I),
     "ref-05-เอกสารประกอบการประชุมวิชาการ.md", "Conference presentation references.md"),
    ("รายงาน/หน่วยงาน", re.compile(r"รายงานการวิจัย|รายงานประจำปี|กรม|กระทรวง|สำนักงาน|ธนาคารแห่งประเทศไทย|"
                                   r"\breport\b|ministry|agency|bureau|OECD|WHO", re.I),
     "ref-04-รายงานทางเทคนิค-และรายงานการวิจัย.md", "Report by a government agency.md"),
    ("วารสาร", re.compile(r"\d+\s*\(\s*\d+\s*\)\s*,\s*\d+|doi\.org|วารสาร|journal", re.I),
     "ref-01-วารสาร-นิตยสาร-หนังสือพิมพ์และจดหมายข่าว.md", "Journal article references.md"),
    ("เว็บ/ออนไลน์", re.compile(r"https?://|สืบค้นจาก|retrieved from|www\.", re.I),
     "ref-14-เว็บเพจและเว็บไซต์.md", "Webpage on a website references.md"),
    ("หนังสือ", re.compile(r"สำนักพิมพ์|พิมพ์ครั้งที่|press\b|publish|wiley|sage|routledge|mcgraw|"
                           r"\(\s*(พิมพ์ครั้งที่|\d+(st|nd|rd|th)\s+ed)", re.I),
     "ref-02-หนังสือ-หนังสืออ้างอิง-และบางบทจากหนังสือ.md", "Bookebook references.md"),
]
THAI_CHARS = re.compile(r"[ก-๙]")


def classify_reference(text):
    """คืน (ชื่อประเภท, ภาษา, ไฟล์คู่มือไทย, ไฟล์คู่มืออังกฤษ) — เดาแบบอนุรักษ์นิยม

    เดาไม่ออกให้ตอบ "ไม่ระบุ" ดีกว่าเดาผิด เพราะ agent จะได้เปิด INDEX แล้วเลือกเอง
    ตามจริง แทนที่จะเชื่อป้ายผิด ๆ แล้วตรวจด้วยกฎของแหล่งคนละประเภท
    """
    lang = "ไทย" if len(THAI_CHARS.findall(text)) >= 5 else "อังกฤษ"
    for name, pattern, tha_file, eng_file in REF_TYPES:
        if pattern.search(text):
            return name, lang, tha_file, eng_file
    return "ไม่ระบุ", lang, "00-INDEX.md", ""


MIN_GROUP = 5          # ก้อนเล็กกว่านี้ไม่คุ้มกับการเปิด agent ใหม่


def manual_for(lang, tha_file, eng_file):
    if lang == "ไทย":
        return f"references/apa7-tha-parts/{tha_file}"
    return (f"references/apa7-eng/{eng_file}" if eng_file
            else "references/apa7-eng/ (ดูรายชื่อไฟล์แล้วเปิดเฉพาะประเภทที่ตรง)")


def group_references(rows, max_groups=4):
    """จัดรายการอ้างอิงเป็นก้อนตามประเภทแหล่ง แล้วยุบก้อนเล็กเข้าด้วยกัน

    คืน list ของ (ชื่อก้อน, [คู่มือที่ต้องเปิด], rows)

    การจัดกลุ่มมีไว้ลดจำนวนครั้งที่ต้องเปิดคู่มือ ไม่ใช่เพื่อเปิด agent ให้มากที่สุด —
    ถ้าปล่อยให้ประเภทละ agent เล่มที่มี 33 รายการจะกลายเป็น 9 agent ที่แต่ละตัวตรวจ
    2-3 รายการ ซึ่ง overhead การอ่านสัญญา+คู่มือแพงกว่างานจริง
    """
    buckets = {}
    for row in rows:
        name, lang, tha_file, eng_file = classify_reference(row[3])
        buckets.setdefault((lang, name, tha_file, eng_file), []).append(row)
    big, small = [], []
    for (lang, name, tha_file, eng_file), items in sorted(buckets.items(), key=lambda kv: -len(kv[1])):
        entry = (f"{name} ({lang})", [manual_for(lang, tha_file, eng_file)], items)
        (big if len(items) >= MIN_GROUP else small).append(entry)
    # เกินโควตา agent → ยุบก้อนที่เล็กที่สุดลงไปกองรวม
    while len(big) > max(1, max_groups - (1 if small else 0)):
        small.append(big.pop())
    if small:
        manuals, items, names = [], [], []
        for name, mans, rows_s in small:
            names.append(f"{name} × {len(rows_s)}")
            manuals.extend(m for m in mans if m not in manuals)
            items.extend(rows_s)
        items.sort(key=lambda r: r[0])
        big.append(("คละประเภท: " + ", ".join(names), manuals, items))
    return big


def citation_windows(text, pad=45):
    """คืนเฉพาะช่วงรอบ ๆ การอ้างอิงในเนื้อหา เชื่อมด้วย " … " """
    spans = []
    for m in INTEXT.finditer(text):
        a, b = max(0, m.start() - pad), min(len(text), m.end() + pad)
        if spans and a <= spans[-1][1] + 10:
            spans[-1] = (spans[-1][0], b)
        else:
            spans.append((a, b))
    if not spans:
        return text
    parts = []
    for i, (a, b) in enumerate(spans):
        prefix = "…" if a > 0 and i == 0 else ""
        parts.append(prefix + text[a:b].strip())
    out = " … ".join(parts)
    return out + ("…" if spans[-1][1] < len(text) else "")


def write_slice(path, header, rows):
    with open(path, "w", encoding="utf-8") as f:
        f.write(header.rstrip() + "\n\n")
        for loc, pg, st, t in rows:
            tag = f"[sheet{pg}|{loc}]" if pg else f"[{loc}]"
            sty = f" <{st}>" if st and st.lower() not in ("normal", "default paragraph font") else ""
            f.write(f"{tag}{sty} {t}\n")
    return os.path.getsize(path)


def chunk(rows, max_chars):
    """แบ่ง rows เป็นก้อนละไม่เกิน max_chars ตัวอักษร ตัดที่ขอบย่อหน้าเสมอ"""
    if max_chars <= 0:
        return [rows]
    out, cur, n = [], [], 0
    for r in rows:
        ln = len(r[3]) + 20
        if cur and n + ln > max_chars:
            out.append(cur); cur, n = [], 0
        cur.append(r); n += ln
    if cur:
        out.append(cur)
    return out


def chunk_by_count(rows, max_rows):
    """แบ่ง rows เป็นก้อนละไม่เกิน max_rows แถว แล้ว**เกลี่ยให้เท่ากัน**

    ถ้าตัดตรง ๆ ทุก max_rows เศษท้ายอาจเหลือรายการเดียว ซึ่งแปลว่าเสีย agent
    ทั้งตัวไปกับงานชิ้นเดียว — เกลี่ยแล้วทุกตัวได้งานพอ ๆ กันและจบพร้อมกัน
    """
    n = len(rows)
    if max_rows <= 0 or n <= max_rows:
        return [rows]
    parts = -(-n // max_rows)                       # ceil
    size = -(-n // parts)
    return [rows[i:i + size] for i in range(0, n, size)]


def main():
    ap = argparse.ArgumentParser(description="ตัด .docx เป็นชิ้นงานสำหรับ agent ขนาน")
    ap.add_argument("docx")
    ap.add_argument("-o", "--out", default="work", help="โฟลเดอร์ปลายทาง (ค่าเริ่มต้น: work/)")
    ap.add_argument("--profile", choices=list(PROFILES.keys()), default=None)
    ap.add_argument("--body-mode", choices=("suspects", "full"), default="suspects",
                    help="suspects (ค่าเริ่มต้น) = ส่งให้ agent เฉพาะย่อหน้าที่คนต้องดู "
                         "(หัวข้อ/caption/บรรทัดสั้น/เซลล์ตาราง) แล้วสรุปช่วงร้อยแก้วที่ตัดออกเป็นบรรทัดเดียว "
                         "· full = ส่งทั้งหมดเหมือนเดิม")
    ap.add_argument("--max-chars", type=int, default=60000,
                    help="ซอย body ที่ยาวเกินเป็นหลายก้อน (body-1.txt, body-2.txt ...) "
                         "เพื่อกระจายให้ agent เพิ่ม; 0 = ไม่ซอย (ค่าเริ่มต้น 60000)")
    ap.add_argument("--max-ref-agents", type=int, default=4,
                    help="จำนวน agent สูงสุดสำหรับรายการอ้างอิง (ค่าเริ่มต้น 4) — "
                         "ก้อนที่เล็กกว่า 5 รายการจะถูกยุบรวมกันเพราะไม่คุ้ม overhead")
    ap.add_argument("--no-group-refs", dest="group_refs", action="store_false",
                    help="ไม่ต้องจัดกลุ่มรายการอ้างอิงตามประเภทแหล่ง (ค่าเริ่มต้นคือจัดให้)")
    ap.add_argument("--max-ref-entries", type=int, default=60,
                    help="จำนวนรายการอ้างอิงต่อ agent หนึ่งตัว (references-1.txt ...) "
                         "APA7 คือคอขวดเพราะต้นทุนคิดต่อ 'รายการ' ไม่ใช่ต่อความยาว — "
                         "แต่ละรายการต้องเปิดคู่มือคนละไฟล์ตามประเภทแหล่ง; 0 = ไม่ซอย (ค่าเริ่มต้น 60)")
    ap.add_argument("--page-engine", choices=("auto", "rendered", "markers"),
                    default="markers",
                    help="markers (ค่าเริ่มต้น) = ใช้ตำแหน่งแบ่งหน้าที่ Word บันทึกไว้ในไฟล์ "
                         "· rendered = ตัวเลือกเสริม ต้องเรียก LibreOffice ได้")
    a = ap.parse_args()

    doc = Document(a.docx)
    profile = a.profile or detect_profile(doc)
    try:
        page_data = resolve_page_data(a.docx, doc, a.page_engine)
    except RuntimeError as exc:
        sys.exit(f"คำนวณแผ่นงานไม่สำเร็จ: {exc}")
    page_by_index, has_markers = unpack_page_data(page_data, doc)
    lines = para_lines(doc, page_data)
    chap, ref, tail = find_boundaries(lines)

    os.makedirs(a.out, exist_ok=True)
    base = os.path.basename(a.docx)
    src = os.path.abspath(a.docx)

    def hdr(title, what):
        return (f"# {title} — {base}\n"
                f"# profile: {profile} | ไฟล์ต้นฉบับ: {src}\n"
                f"# ป้าย: [sheet<แผ่น>|¶<ย่อหน้า>] — **ใช้แค่เลขแผ่น** เช่น location = \"แผ่นที่ 12\"\n"
                f"#   ไม่ต้องลอก ¶ (Word ไม่มีเลขย่อหน้าให้ดู รายงานตัดทิ้งอยู่แล้ว)\n"
                f"#   สิ่งที่ต้องมีแทน: ยกข้อความจริงที่ผิดใส่เครื่องหมายคำพูดใน issue\n"
                f"# ขอบเขตชิ้นนี้: {what}")

    fm, bd, rf, tail_rows = lines[:chap], lines[chap:ref], lines[ref:tail], lines[tail:]
    manifest = {
        "source": src, "profile": profile, "profile_label": PROFILES[profile]["label"],
        "page_markers_available": has_markers,
        "page_source": page_data.get("source", "unavailable"),
        "total_paragraphs": len(doc.paragraphs), "total_text_paragraphs": len(lines),
        "total_tables": len(doc.tables),
        "last_page": page_data.get("total_pages") or max([p for p in page_by_index if p] or [0]),
        "slices": {},
    }

    def record(key, fname, rows, role, what):
        if not rows:
            manifest["slices"][key] = {"file": None, "paragraphs": 0, "note": "ไม่พบส่วนนี้ในเล่ม"}
            return
        path = os.path.join(a.out, fname)
        size = write_slice(path, hdr(role, what), rows)
        manifest["slices"][key] = {
            "file": path, "role": role, "paragraphs": len(rows),
            "location_range": [rows[0][0], rows[-1][0]],
            "page_range": [rows[0][1], rows[-1][1]],
            "chars": size,
        }

    record("frontmatter", "frontmatter.txt", fm,
           "ส่วนนำ (Agent 1)", "ปกใน/ปกภาษาที่สอง, หน้าอนุมัติ, บทคัดย่อ, กิตติกรรมประกาศ, สารบัญ/สารบัญตาราง/สารบัญภาพ")

    bd_full, tail_full = list(bd), list(tail_rows)   # เก็บไว้ดึงการอ้างอิงในเนื้อหา
    if a.body_mode == "suspects":
        before = sum(len(r[3]) for r in bd) + sum(len(r[3]) for r in tail_rows)
        bd, tail_rows = keep_suspects(bd), keep_suspects(tail_rows)
        after = sum(len(r[3]) for r in bd) + sum(len(r[3]) for r in tail_rows)
        trimmed = (1 - after / before) if before else 0.0
        manifest["body_mode"] = "suspects"
        manifest["body_trimmed_ratio"] = round(trimmed, 4)
    else:
        manifest["body_mode"] = "full"
        trimmed = 0.0

    body_chunks = chunk(bd, a.max_chars)
    if len(body_chunks) == 1:
        record("body", "body.txt", bd,
               "เนื้อหา (Agent 2)", "บทที่ 1 ถึงบทสุดท้าย: หัวข้อบท/หัวข้อย่อย, caption ตาราง-ภาพ, ความสม่ำเสมอของรูปแบบ")
    else:
        for n, ck in enumerate(body_chunks, 1):
            record(f"body-{n}", f"body-{n}.txt", ck,
                   f"เนื้อหา ส่วนที่ {n}/{len(body_chunks)} (Agent 2.{n})",
                   "บทเนื้อหาช่วงนี้: หัวข้อบท/หัวข้อย่อย, caption ตาราง-ภาพ, ความสม่ำเสมอของรูปแบบ")

    ref_groups = group_references(rf, a.max_ref_agents) if a.group_refs else []
    if len(ref_groups) <= 1 and len(rf) <= a.max_ref_entries:
        manuals = ref_groups[0][1] if ref_groups else []
        record("references", "references.txt", rf,
               "รายการอ้างอิง (Agent 3)",
               "เฉพาะรายการอ้างอิง/บรรณานุกรม — ตรวจ APA7 ทุกรายการ"
               + (f"; ส่วนใหญ่เป็น {ref_groups[0][0]} เปิดคู่มือ {manuals[0]} ก่อน" if ref_groups else ""))
        if manuals:
            manifest["slices"]["references"]["manual"] = manuals[0]
    elif ref_groups:
        n = 0
        for label, manuals, rows_g in ref_groups:
            for part in chunk_by_count(rows_g, a.max_ref_entries):
                n += 1
                manual_text = " · ".join(manuals)
                record(f"references-{n}", f"references-{n}.txt", part,
                       f"รายการอ้างอิง: {label} — {len(part)} รายการ (Agent 3.{n})",
                       f"รายการในชิ้นนี้ถูกจัดประเภทไว้ให้แล้วว่าเป็น **{label}** — "
                       f"เปิดคู่มือเท่าที่ระบุนี้พอ: {manual_text} "
                       f"(ถ้าเจอรายการที่ประเภทไม่ตรงกับป้าย ให้ตรวจตามประเภทจริงแล้วเขียนบอกไว้ใน issue); "
                       f"การจับคู่ in-text ↔ รายการท้ายเล่มเป็นงานของ check_deep.py และ Agent 3.1 เท่านั้น")
                manifest["slices"][f"references-{n}"]["ref_type"] = label
                manifest["slices"][f"references-{n}"]["manual"] = manuals[0]
                manifest["slices"][f"references-{n}"]["manuals"] = manuals
    else:
        for n, ck in enumerate(chunk_by_count(rf, a.max_ref_entries), 1):
            record(f"references-{n}", f"references-{n}.txt", ck,
                   f"รายการอ้างอิง ส่วนที่ {n} (Agent 3.{n})",
                   "รายการอ้างอิงช่วงนี้ — ตรวจ APA7 ทุกรายการตามภาษาของแต่ละรายการ")

    record("backmatter", "backmatter.txt", tail_rows,
           "ภาคผนวกและท้ายเล่ม (Agent 2)", "ภาคผนวกและประวัติผู้เขียน — ตรวจรูปแบบเทมเพลต ไม่ใช่ APA7")

    # in-text citations — ดึงเฉพาะย่อหน้าที่มีการอ้างอิง เพื่อให้ agent APA7
    # ตรวจ in-text ↔ รายการท้ายเล่มได้โดยไม่ต้องอ่าน body ทั้งก้อน
    # ดึงจากข้อความเต็มเสมอ — การอ้างอิงในเนื้อหาส่วนใหญ่อยู่กลางร้อยแก้วซึ่งโหมด
    # suspects ตัดทิ้งไปแล้ว ถ้าดึงจากที่ตัดแล้ว Agent 3 จะเห็น in-text แทบไม่เหลือ
    cited = [r for r in (bd_full + tail_full) if INTEXT.search(r[3])]
    if a.body_mode == "suspects":
        # Agent 3 ตรวจ "รูปแบบการอ้างอิง" ไม่ได้ตรวจเนื้อความรอบ ๆ — ส่งทั้งย่อหน้าไป
        # เปลืองเปล่า ตัดเหลือหน้าต่างรอบจุดอ้างอิงพอให้เห็นบริบทว่าเป็นการอ้างแบบไหน
        cited = [(loc, pg, st, citation_windows(t)) for loc, pg, st, t in cited]
    record("intext", "intext.txt", cited,
           "การอ้างอิงในเนื้อหา (Agent 3)", "เฉพาะย่อหน้าที่ตรวจพบรูปแบบการอ้างอิงในเนื้อหา — ใช้จับคู่กับรายการท้ายเล่ม")

    # แผนการ spawn — orchestrator อ่านจากตรงนี้ได้เลย ไม่ต้องเดาว่าต้องใช้ agent กี่ตัว
    plan = []
    fm_slice = manifest["slices"].get("frontmatter", {})
    if fm_slice.get("file"):
        plan.append({"agent": "agent1-frontmatter", "contract": "Agent 1",
                     "inputs": [fm_slice["file"]],
                     "out": os.path.join(a.out, "agent1-frontmatter.json")})
    body_keys = sorted(k for k in manifest["slices"]
                       if (k == "body" or k.startswith("body-")) and manifest["slices"][k].get("file"))
    tail_file = (manifest["slices"].get("backmatter") or {}).get("file")
    for n, key in enumerate(body_keys, 1):
        inputs = [manifest["slices"][key]["file"]]
        if key == body_keys[-1] and tail_file:
            inputs.append(tail_file)          # ภาคผนวกไปกับก้อนสุดท้ายเสมอ
        name = "agent2-body" if len(body_keys) == 1 else f"agent2-body-{n}"
        plan.append({"agent": name, "contract": f"Agent 2{'' if len(body_keys) == 1 else f'.{n}'}",
                     "inputs": inputs, "out": os.path.join(a.out, name + ".json")})
    ref_keys = sorted(k for k in manifest["slices"]
                      if (k == "references" or k.startswith("references-")) and manifest["slices"][k].get("file"))
    intext_file = (manifest["slices"].get("intext") or {}).get("file")
    for n, key in enumerate(ref_keys, 1):
        inputs = [manifest["slices"][key]["file"]]
        if n == 1 and intext_file:
            inputs.append(intext_file)        # จับคู่ in-text เป็นงานของตัวแรกตัวเดียว
        name = "agent3-apa7" if len(ref_keys) == 1 else f"agent3-apa7-{n}"
        item = {"agent": name, "contract": f"Agent 3{'' if len(ref_keys) == 1 else f'.{n}'}",
                "inputs": inputs, "out": os.path.join(a.out, name + ".json")}
        slice_info = manifest["slices"][key]
        if slice_info.get("manual"):
            item["manual"] = slice_info["manual"]
            item["ref_type"] = slice_info.get("ref_type")
        plan.append(item)
    manifest["agent_plan"] = plan

    mpath = os.path.join(a.out, "manifest.json")
    with open(mpath, "w", encoding="utf-8") as fh:
        json.dump(manifest, fh, ensure_ascii=False, indent=2)

    print("=" * 70)
    print(f"ตัดชิ้นงานเสร็จ — {base}   profile: {profile} ({PROFILES[profile]['label']})")
    print(f"ย่อหน้าที่มีข้อความ {len(lines)} · ตาราง {len(doc.tables)} · แผ่นงานทั้งหมด {manifest['last_page']}")
    if manifest.get("body_mode") == "suspects":
        print(f"โหมด suspects: ตัดร้อยแก้วออก {manifest['body_trimmed_ratio']:.0%} "
              f"ของเนื้อหา+ภาคผนวก (agent อ่านน้อยลงเท่านั้น เกณฑ์ตรวจเท่าเดิม)")
    if not has_markers:
        print("⚠ ไม่สามารถยืนยันแผ่นงานจาก Word marker ได้")
    print("=" * 70)
    for k, v in manifest["slices"].items():
        if not v.get("file"):
            print(f"  {k:12s} — {v.get('note')}")
            continue
        pr = v["page_range"]
        print(f"  {k:12s} {v['paragraphs']:5d} ย่อหน้า · แผ่นที่ {pr[0]}–{pr[1]} · {v['chars']//1024} KB · {os.path.basename(v['file'])}")
    print(f"\nmanifest: {mpath}")
    print(f"แผนการ spawn — agent {len(plan)} ตัว (สร้างพร้อมกันในเทิร์นเดียว):")
    for item in plan:
        extra = f"  [คู่มือ: {os.path.basename(item['manual'])}]" if item.get("manual") else ""
        print(f"  {item['contract']:12s} {', '.join(os.path.basename(x) for x in item['inputs']):36s}"
              f" → {os.path.basename(item['out'])}{extra}")
    print("ใช้ prompt ตาม references/agent-contracts.md")


if __name__ == "__main__":
    main()
