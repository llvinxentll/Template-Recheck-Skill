#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
annotate_docx.py — แทรก "คอมเมนต์ Word" ลงในไฟล์วิทยานิพนธ์ตรงจุดที่ต้องแก้

ทำไมต้องมีขั้นนี้
----------------
รายงาน .docx บอกได้แค่ "แผ่นที่ ~12" ซึ่งนักศึกษายังต้องไล่หาเองทั้งหน้า
(และเลขแผ่นก็คลาดได้ถ้าแก้ไฟล์ต่อ) สคริปต์นี้จึงสร้าง
**สำเนาไฟล์ต้นฉบับที่มีคอมเมนต์ปักไว้ตรงจุด** — เปิดไฟล์แล้วเห็นบอลลูนที่ขอบหน้า
คลิกทีเดียวเคอร์เซอร์ไปที่ข้อความนั้นเลย ไม่ต้องไล่หาเอง

ใช้คู่กับรายงาน ไม่ได้แทนกัน:
    ไฟล์คอมเมนต์ = "จุดไหนต้องแก้"   ·   รายงาน = "สรุปทั้งเล่มมีอะไรบ้าง กี่จุด"

ใช้:
    python3 scripts/annotate_docx.py "<ไฟล์>.docx" auto.json deep.json apa7.json \
        -o "<ไฟล์>_annotated.docx"

การหาจุดปัก (ไล่ตามลำดับ)
-------------------------
1. `tbl2:r1:c3:¶1` ใน `anchor`/`location` → ย่อหน้านั้นในเซลล์ตาราง
2. `¶345` / `para 345` ใน `anchor`/`location` → ย่อหน้าลำดับที่ 345 ของ body
3. ข้อความที่ยกมาใน `snippet`/`issue`/`correct` → ค้นหาข้อความนั้นในเล่ม
   **ถ้า location บอกแผ่นไว้ จะเลือกจุดที่อยู่บนแผ่นนั้นก่อน** แล้วค่อยขยายไปทั้งเล่ม
4. หาไม่ได้ (เช่น ระยะขอบ/สไตล์ ซึ่งเป็นปัญหาระดับทั้ง section) → รวมเป็นคอมเมนต์
   ก้อนเดียวปักที่ย่อหน้าแรกของเล่ม พร้อมระบุว่าเป็นปัญหาระดับเอกสาร

ข้อ 1–2 มาจากสคริปต์ตรวจอัตโนมัติซึ่งใส่ ¶ ให้เองอยู่แล้ว — **คนหรือ agent ไม่ต้องเขียน ¶ เอง**
ให้ยกข้อความจริงที่ผิดใส่เครื่องหมายคำพูดใน `issue` (หรือใส่ฟิลด์ `snippet`) แล้วข้อ 3 ทำงานแทน

ต้องใช้ python-docx >= 1.2 (รองรับคอมเมนต์)
    pip install -U python-docx --break-system-packages
"""
import argparse
import json
import os
import re
import sys

try:
    import docx
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
except ImportError:
    sys.exit("ต้องติดตั้ง python-docx ก่อน:  pip install -U python-docx")

SEV_MARK = {"critical": "🔴 ต้องแก้ก่อนส่ง", "major": "🟠 ควรแก้",
            "minor": "🟡 เล็กน้อย", "suggestion": "🟢 ข้อเสนอแนะ"}
DOMAIN_LABEL = {"template": "รูปแบบเทมเพลต", "spelling": "คำผิด/อักขระ",
                "apa7": "การอ้างอิง APA 7"}

TBL_LOC = re.compile(r"tbl(\d+):r(\d+):c(\d+):¶(\d+)", re.I)
PARA_LOC = re.compile(r"(?:\bpara|¶)\s*(\d+)", re.I)
SHEET_LOC = re.compile(r"(?:\[sheet|\[p|แผ่นที่)\s*~?\s*(\d+)", re.I)
# (?<![\d]) กัน 1.5" (หน่วยนิ้ว) ไม่ให้ถูกอ่านเป็นเครื่องหมายคำพูด
QUOTED = re.compile(r"(?<![\d])[\"“'‘]([^\"”'’\n]{6,120})[\"”'’]")
WS = re.compile(r"\s+")


def _require_comment_support():
    if not hasattr(Document(), "add_comment"):
        sys.exit(
            f"python-docx เวอร์ชันนี้ ({getattr(docx, '__version__', '?')}) ยังไม่รองรับคอมเมนต์\n"
            "อัปเดตก่อน:  pip install -U python-docx --break-system-packages")


def load_findings(paths):
    out = []
    for p in paths:
        with open(p, encoding="utf-8") as fh:
            payload = json.load(fh)
        for f in payload.get("findings", []):
            f = dict(f)
            f.setdefault("_source", os.path.basename(p))
            out.append(f)
    return out


def severity_rank(f):
    order = {"critical": 0, "major": 1, "minor": 2, "suggestion": 3}
    return order.get((f.get("severity") or "").lower(), 4)


def comment_text(f):
    """ข้อความในบอลลูนคอมเมนต์ — สั้น อ่านแล้วลงมือแก้ได้ทันที"""
    sev = SEV_MARK.get((f.get("severity") or "").lower(), "•")
    dom = DOMAIN_LABEL.get(f.get("domain"), f.get("domain") or "")
    lines = [f"{sev} · {dom}" + (f" · {f['category']}" if f.get("category") else "")]
    if f.get("issue"):
        lines.append(f"ปัญหา: {f['issue']}")
    if f.get("criterion"):
        lines.append(f"เกณฑ์: {f['criterion']}")
    if f.get("fix"):
        lines.append(f"วิธีแก้: {f['fix']}")
    if f.get("correct"):
        lines.append(f"แก้เป็น: {f['correct']}")
    if not f.get("criterion") and not f.get("fix") and f.get("detail"):
        lines.append(str(f["detail"]))
    return "\n".join(lines)


def _count_breaks(el):
    total = 0
    for node in el.iter():
        if node.tag == qn("w:lastRenderedPageBreak"):
            total += 1
        elif node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
            total += 1
    return total


def build_index(doc):
    """คืน (body_paragraphs, table_paragraph_map, text_index)

    table_paragraph_map[(t, r, c, p)] = Paragraph
    text_index = list[(normalized_text, Paragraph, sheet)] ตามลำดับเอกสาร

    ``sheet`` มาจากมาร์กเกอร์แบ่งหน้าที่ Word บันทึกไว้ (นับปกเป็นแผ่นที่ 1) ใช้กรองผล
    ค้นข้อความให้อยู่บนแผ่นเดียวกับที่ finding ระบุ — ข้อความที่ซ้ำได้ทั้งเล่ม เช่น
    ชื่อหัวข้อหรือชื่อผู้แต่งที่อ้างหลายที่ จะได้ไม่ไปปักผิดจุด
    """
    body = doc.paragraphs
    tbl_map, text_index = {}, []
    P, TBL = qn("w:p"), qn("w:tbl")
    page, body_i, table_i = 1, 0, 0

    def add(p, sheet):
        t = WS.sub(" ", (p.text or "")).strip()
        if t:
            text_index.append((t, p, sheet))

    for child in doc.element.body.iterchildren():
        if child.tag == P:
            if body_i < len(body):
                add(body[body_i], page)
            body_i += 1
            page += _count_breaks(child)
        elif child.tag == TBL:
            table_i += 1
            table = Table(child, doc._body)
            seen, local_page = set(), page
            for ri, row in enumerate(table.rows, 1):
                for ci, cell in enumerate(row.cells, 1):
                    if cell._tc in seen:
                        continue
                    seen.add(cell._tc)
                    for pi, p in enumerate(cell.paragraphs, 1):
                        tbl_map[(table_i, ri, ci, pi)] = p
                        add(p, local_page)
                        local_page += _count_breaks(p._p)
            page += _count_breaks(child)
    return body, tbl_map, text_index


def needles(f):
    """ข้อความที่ยกมาใน finding — ใช้ค้นหาจุดปักเมื่อไม่มี ¶

    เรียงจากที่น่าเชื่อถือที่สุด: ฟิลด์ snippet/quote ที่ตั้งใจใส่มา → ข้อความในเครื่องหมาย
    คำพูดใน issue (สิ่งที่พบจริง) → correct/detail
    """
    out = []
    for field in ("snippet", "quote", "text"):
        v = WS.sub(" ", str(f.get(field) or "")).strip()
        if len(v) >= 4:
            out.append(v)
    for field in ("issue", "correct", "detail"):
        for q in QUOTED.findall(str(f.get(field) or "")):
            v = WS.sub(" ", q).strip()
            if len(v) >= 6:
                out.append(v)
    return out


def anchor_for(f, body, tbl_map, text_index, used):
    """คืน (paragraph, how) — how ∈ {'para', 'table', 'text'} หรือ (None, None)"""
    loc = " ".join(str(f.get(k) or "") for k in ("anchor", "location"))

    m = TBL_LOC.search(loc)
    if m:
        key = tuple(int(x) for x in m.groups())
        p = tbl_map.get(key)
        if p is not None and p.runs:
            return p, "table"

    m = PARA_LOC.search(loc)
    if m:
        i = int(m.group(1))
        if 0 <= i < len(body) and body[i].runs:
            return body[i], "para"

    # ไม่มี ¶ (finding ที่คน/agent เขียน) → ค้นจากข้อความที่ยกมา
    # ถ้า location บอกแผ่นไว้ ให้เลือกจุดบนแผ่นนั้นก่อน แล้วค่อยเอาจุดที่ใกล้แผ่นนั้นที่สุด
    m = SHEET_LOC.search(loc)
    sheet = int(m.group(1)) if m else None
    for needle in needles(f):
        best = None
        for text, p, pg in text_index:
            if needle not in text or id(p) in used or not p.runs:
                continue
            if sheet is None or pg == sheet:
                return p, "text"
            if best is None or abs(pg - sheet) < best[1]:
                best = (p, abs(pg - sheet))
        if best is not None:
            return best[0], "text"
    return None, None


def first_anchorable(body, tbl_map):
    for p in body:
        if p.runs:
            return p
    for p in tbl_map.values():
        if p.runs:
            return p
    return None


def annotate(docx_path, json_paths, out_path, author="TULIBS Checker",
             initials="TU", max_doc_level=40, include_info=False):
    _require_comment_support()
    findings = load_findings(json_paths)
    skipped_info = 0
    if not include_info:
        keep = [f for f in findings
                if (f.get("severity") or "").lower() not in ("info", "note")]
        skipped_info = len(findings) - len(keep)
        findings = keep
    findings.sort(key=severity_rank)
    doc = Document(docx_path)
    body, tbl_map, text_index = build_index(doc)

    used = set()
    stats = {"para": 0, "table": 0, "text": 0, "document": 0}
    doc_level = []

    for f in findings:
        p, how = anchor_for(f, body, tbl_map, text_index, used)
        if p is None:
            doc_level.append(f)
            stats["document"] += 1
            continue
        try:
            doc.add_comment(runs=p.runs, text=comment_text(f),
                            author=author, initials=initials)
        except Exception as exc:                                  # noqa: BLE001
            doc_level.append(f)
            stats["document"] += 1
            print(f"⚠ ปักคอมเมนต์ที่ {f.get('location') or '-'} ไม่สำเร็จ ({exc}) "
                  "→ ย้ายไปคอมเมนต์รวมหัวเล่ม")
            continue
        used.add(id(p))
        stats[how] += 1

    if doc_level:
        head = first_anchorable(body, tbl_map)
        if head is not None:
            shown = doc_level[:max_doc_level]
            lines = [f"⚠ ปัญหาระดับทั้งเล่ม/ทั้ง section — {len(doc_level)} รายการ",
                     "(แก้ที่การตั้งค่าเอกสารหรือสไตล์ ไม่ได้ผูกกับย่อหน้าใดย่อหน้าหนึ่ง)", ""]
            for i, f in enumerate(shown, 1):
                loc = f.get("location") or "ทั้งเอกสาร"
                lines.append(f"{i}. [{loc}] {comment_text(f)}")
                lines.append("")
            if len(doc_level) > len(shown):
                lines.append(f"… และอีก {len(doc_level) - len(shown)} รายการ ดูในไฟล์รายงาน")
            doc.add_comment(runs=head.runs, text="\n".join(lines),
                            author=author, initials=initials)

    doc.save(out_path)
    total = len(findings)
    pinned = stats["para"] + stats["table"] + stats["text"]
    print("=" * 70)
    print(f"แทรกคอมเมนต์แล้ว: {out_path}")
    print(f"finding ทั้งหมด {total} รายการ")
    print(f"  · ปักตรงย่อหน้า      {stats['para']}")
    print(f"  · ปักในเซลล์ตาราง    {stats['table']}")
    print(f"  · ปักด้วยการค้นข้อความ {stats['text']}")
    print(f"  · ระดับทั้งเล่ม/section {stats['document']} (รวมเป็นคอมเมนต์เดียวที่ต้นเล่ม)")
    if skipped_info:
        print(f"  · ข้าม {skipped_info} รายการที่เป็นข้อมูลประกอบ (ใส่ --include-info ถ้าต้องการ)")
    if total:
        print(f"ปักตรงจุดได้ {pinned}/{total} = {pinned / total:.0%}")
    print("เปิดไฟล์ใน Word → แท็บ Review → Show Comments เพื่อไล่แก้ทีละจุด")
    return stats


def main():
    ap = argparse.ArgumentParser(
        description="แทรกคอมเมนต์ Word ลงไฟล์วิทยานิพนธ์ตรงจุดที่ต้องแก้")
    ap.add_argument("docx", help="ไฟล์ต้นฉบับของนักศึกษา (.docx)")
    ap.add_argument("json", nargs="+", help="ไฟล์ผลตรวจ (auto.json deep.json apa7.json ...)")
    ap.add_argument("-o", "--out", default=None,
                    help="ไฟล์ปลายทาง (ค่าเริ่มต้น: <ชื่อเดิม>_annotated.docx)")
    ap.add_argument("--author", default="TULIBS Checker")
    ap.add_argument("--initials", default="TU")
    ap.add_argument("--include-info", action="store_true",
                    help="แทรกรายการระดับ info/note ด้วย (ค่าเริ่มต้นข้าม เพราะไม่ใช่จุดที่ต้องแก้)")
    a = ap.parse_args()

    if not os.path.exists(a.docx):
        sys.exit(f"ไม่พบไฟล์: {a.docx}")
    for j in a.json:
        if not os.path.exists(j):
            sys.exit(f"ไม่พบไฟล์ผลตรวจ: {j}")
    out = a.out or (os.path.splitext(a.docx)[0] + "_annotated.docx")
    annotate(a.docx, a.json, out, a.author, a.initials,
             include_info=a.include_info)


if __name__ == "__main__":
    main()
