#!/usr/bin/env python3
"""Regression tests for the rules that historically fired on correct text.

Run: python3 scripts/test_false_positives.py

A false positive costs more than a miss: the student burns time "fixing" text
that was already right, and stops trusting the report. Every rule that fires on
free-form prose or on URLs belongs here with the shapes it must NOT flag.
"""
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import check_deep as cd  # noqa: E402

FAIL = []


def check(cond, msg):
    if not cond:
        FAIL.append(msg)


# --- in-text citation matcher: must not invent author names from prose -------
SHOULD_NOT_MATCH = [
    "Table 3 shows the change in GDP (2019) across provinces.",
    "The results were published later that year (2019).",
    "This framework was revised in Thailand (2020) and elsewhere.",
    "Figure 2 (2021) illustrates the sampling frame.",
    "According to the survey (2018), most respondents agreed.",
    "Data collection ran from January (2019) to March.",
    "ผลการศึกษาในช่วง Section 4 (2020) แสดงให้เห็นว่า",
]
SHOULD_MATCH = [
    ("Smith (2019) argued that the model holds.", "Smith", "2019"),
    ("Recent work (Johnson, 2020) supports this.", "Johnson", "2020"),
    ("Nguyen et al. (2021) replicated the finding.", "Nguyen", "2021"),
    ("As shown by Garcia & Lopez (2018), the effect is small.", "Garcia", "2018"),
    ("แนวคิดนี้ได้รับการยืนยันโดย Anderson (2017) ในงานวิจัยล่าสุด", "Anderson", "2017"),
]

for text in SHOULD_NOT_MATCH:
    hits = [m.group(1) for m in cd._INTEXT_RE.finditer(text)
            if m.group(1).lower() not in cd._NOT_A_SURNAME]
    check(not hits, f"in-text: false positive {hits} in {text!r}")

for text, name, year in SHOULD_MATCH:
    hits = [(m.group(1), m.group(2)) for m in cd._INTEXT_RE.finditer(text)
            if m.group(1).lower() not in cd._NOT_A_SURNAME]
    check((name, year) in hits, f"in-text: missed {name} ({year}) in {text!r}")

# --- hyphen-in-number-range: must not touch URLs or DOIs ---------------------
URL_CASES = [
    "Ministry of Health. (2020). Annual report. https://example.org/11-2019/report.pdf",
    "Jackson, L. M. (2019). The psychology of prejudice. https://doi.org/10.1037/0000168-000",
    "OECD. (2021). Education at a glance. https://doi.org/10.1787/b35a14e5-en",
]
for t in URL_CASES:
    flagged = [m.group(0) for m in
               re.finditer(r"(?<![\d\-–])(\d{1,4})\s*-\s*(\d{1,4})(?![\d\-–])", t)
               if not cd._inside_url(t, m.start())
               and int(m.group(2)) >= int(m.group(1))]
    check(not flagged, f"page-range: false positive {flagged} in {t!r}")

REAL_RANGE = "Grady, J. S. (2019). Emotions in storybooks. Media Culture, 8(3), 207-217."
flagged = [m.group(0) for m in
           re.finditer(r"(?<![\d\-–])(\d{1,4})\s*-\s*(\d{1,4})(?![\d\-–])", REAL_RANGE)
           if not cd._inside_url(REAL_RANGE, m.start())
           and int(m.group(2)) >= int(m.group(1))]
check(flagged == ["207-217"], f"page-range: missed real hyphen range, got {flagged}")

# --- publisher location: only after the title, never inside it ---------------
TITLE_LOOKALIKE = "Chen, W. (2019). Bangkok, Thailand: A case study of urban transit. Routledge."
REAL_LOCATION = "Chen, W. (2019). Urban transit systems. Cambridge, MA: Harvard University Press."
LOC_RE = re.compile(
    r"(?:^|[\.\?!]\s)\s*([A-Z][a-zA-Z\-' ]{2,30},\s*(?:[A-Z]{2}|[A-Z][a-z]+))\s*:\s*[A-Z]")
check(not LOC_RE.search(cd._after_title_span(TITLE_LOOKALIKE)),
      "publisher-location: false positive on a title that starts with a city name")
check(LOC_RE.search(cd._after_title_span(REAL_LOCATION)),
      "publisher-location: missed a real 'City, ST: Publisher'")

# --- ป้ายรายการหน้าบทคัดย่อ: ต้องรับรูปที่เขียนต่างกันได้ ---------------------
# เล่มจริงที่ผ่านการตรวจใช้ป้ายนี้อย่างน้อย 3 แบบ ถ้าจำแบบเดียวจะฟ้องผิดเกือบทุกเล่ม
FACULTY_KEYS = dict(cd.ABSTRACT_LABELS_EN)["Faculty/University"]
for label in ("Department/Faculty/University", "Major Field Faculty University",
              "Faculty University"):
    check(any(k.lower() in label.lower() for k in FACULTY_KEYS),
          f"abstract-table: ป้าย {label!r} ควรถือว่าถูกต้อง")

# --- บันไดการเยื้อง: ค่าที่ต่างกันแค่เศษปัดต้องไม่ถือว่าไม่สม่ำเสมอ -----------
check(cd._INDENT_TOL >= 0.05,
      "indent ladder: tolerance แคบเกินไป จะฟ้อง 0.49\" กับ 0.50\" ว่าต่างกัน")

# --- บรรทัดว่างที่เป็นการจัดหน้าปกติ ต้องไม่ถูกนับเป็น "Enter คั่นย่อหน้า" -----
# ไฟล์ตัวอย่างทางการเว้นบรรทัดก่อนหัวข้อและก่อน "คำสำคัญ:" เป็นเรื่องปกติ กฎจึงต้อง
# ดูทั้งบรรทัดบนและล่างว่าเป็นร้อยแก้วจริงทั้งคู่ ไม่ใช่เห็นบรรทัดว่างแล้วฟ้อง
LONG = "ก" * 150
for boundary in ("คำสำคัญ: การละทิ้งตะกร้าสินค้า, พฤติกรรมผู้บริโภค",
                 "2.1 แนวคิดและทฤษฎีที่เกี่ยวข้องกับพฤติกรรมผู้บริโภคออนไลน์",
                 "ตารางที่ 4.1 แสดงผลการวิเคราะห์ข้อมูลเบื้องต้นของกลุ่มตัวอย่าง",
                 "A STUDY OF THE EFFECTS OF LIFESTYLE ON PURCHASE INTENTION"):
    check(not cd._is_prose(boundary),
          f"blank-line: ถือว่า {boundary[:34]!r} เป็นร้อยแก้ว จะทำให้ฟ้องบรรทัดว่างที่ถูกต้อง")
check(cd._is_prose(LONG), "blank-line: ย่อหน้ายาวปกติต้องนับเป็นร้อยแก้ว")

# --- หน้าปก: ต้องไม่เรียกร้องปกภาษาที่ไฟล์นั้นไม่มี --------------------------
# ไฟล์ปกอังกฤษล้วนมีอักษรไทยได้ (หัว "ภาคผนวก ค") — ถ้าเช็คแค่ "มีอักษรไทย"
# จะไปฟ้องว่าปกไทยไม่ครบทั้งที่ไฟล์นั้นไม่มีปกไทยตั้งแต่แรก
check(len(cd.COVER_TH) >= 4 and len(cd.COVER_EN) >= 4,
      "cover: รายการบรรทัดบนปกหายไป")

# เทมเพลตทางการทั้งสามไฟล์ไม่มีบรรทัดลิขสิทธิ์ — เคยใส่เป็นกฎแล้วฟ้องผิดครบทุกเล่ม
for name, _ in cd.COVER_TH + cd.COVER_EN:
    check("ลิขสิทธิ์" not in name and "COPYRIGHT" not in name.upper(),
          "cover: เอาบรรทัดลิขสิทธิ์/COPYRIGHT กลับเข้ามาเป็นเกณฑ์ — "
          "เทมเพลตทางการ 3 ไฟล์ไม่มีบรรทัดนี้เลย")

# --- ตรวจสองทิศทาง: ขาดบรรทัดที่ต้องมี + มีบรรทัดที่เทมเพลตตัดออกแล้ว ----------
DEP = [c for _, c in cd.COVER_DEPRECATED]
check("ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์" in DEP,
      "cover: บรรทัดที่เทมเพลตตัดออกแล้วหายไปจากรายการตรวจ — เล่มที่ก๊อปปกรุ่นเก่ามาจะผ่านฉลุย")

# ประโยคที่ "พูดถึง" บรรทัดนี้ (เช่นในคู่มือหรือรายงานผลตรวจ) ต้องไม่ถูกนับว่าเป็นหน้าปกผิด
MENTIONS = "ปี พ.ศ. ต้องเว้นวรรค/ขึ้นบรรทัดใหม่ก่อน 'ลิขสิทธิ์ของมหาวิทยาลัยธรรมศาสตร์' ตามเกณฑ์"
for pattern, canonical in cd.COVER_DEPRECATED:
    if re.search(pattern, MENTIONS, re.I):
        check(len(MENTIONS) > len(canonical) + 15,
              "cover-deprecated: ประโยคที่อ้างถึงบรรทัดนี้ต้องยาวพอที่จะถูกกรองออก")

# แม่แบบหัวข้อต้องตรงกับเทมเพลต
check(cd.HEADING_STYLE_BY_LEVEL[2] == "TU_Main Heading"
      and cd.HEADING_STYLE_BY_LEVEL[3] == "TU_Sub-heading 1"
      and cd.HEADING_STYLE_BY_LEVEL[4] == "TU_Sub-heading 2",
      "heading style: แม่แบบระดับหัวข้อไม่ตรงกับเทมเพลต")

# --- ตรวจกับไฟล์ตัวอย่างทางการ: กฎใหม่ต้องไม่ฟ้องอะไรเลย --------------------
# --- ประเภทงาน: เทมเพลตเดียวกัน เปลี่ยนแค่คำเรียก ---------------------------
for blob, want in [("การค้นคว้าอิสระนี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร", "is"),
                   ("AN INDEPENDENT STUDY SUBMITTED IN PARTIAL FULFILLMENT OF", "is"),
                   ("สารนิพนธ์นี้เป็นส่วนหนึ่งของการศึกษา", "term-paper"),
                   ("A DISSERTATION SUBMITTED IN PARTIAL FULFILLMENT", "dissertation"),
                   ("วิทยานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร", "thesis")]:
    got = cd.detect_work_type(blob)[0]
    check(got == want, f"work-type: {blob[:34]!r} ควรเป็น {want} แต่ได้ {got}")

# ปกอังกฤษของดุษฎีนิพนธ์ที่หอสมุดใช้ **มี** COPYRIGHT OF THAMMASAT UNIVERSITY จริง
# ต่างจากปกไทยที่ไม่มี — ใส่เป็นบรรทัดต้องห้ามเมื่อไรคือฟ้องตัวอย่างที่ถูกต้อง
for _, canonical in cd.COVER_DEPRECATED:
    check("COPYRIGHT" not in canonical.upper(),
          "cover-deprecated: COPYRIGHT OF THAMMASAT UNIVERSITY มีอยู่จริงในปกอังกฤษของหอสมุด "
          "ห้ามใส่เป็นบรรทัดต้องห้าม")

# --- ตรวจกับไฟล์ตัวอย่างทางการ: กฎใหม่ต้องไม่ฟ้องอะไรเลย --------------------
# เฉพาะไฟล์ที่ **หอสมุดทำเอง** เท่านั้นที่ใช้เป็นเกณฑ์ "ต้องได้ 0"
# โฟลเดอร์ "อ้างอิงคำเรียกประเภทงาน" เป็นผลงานนักศึกษาที่ใช้ดูคำเรียก (การค้นคว้าอิสระ /
# DISSERTATION) ไม่ใช่ไฟล์ต้นแบบ — เอามาตั้งเป็นเกณฑ์ไม่ได้ เพราะตัวมันเองก็มีจุดที่ผิดรูปแบบ
SAMPLE_DIRS = [Path(__file__).resolve().parent.parent / d
               for d in ("ตัวอย่างการเขียนและวางเนื้อหา",)]
for SAMPLES in SAMPLE_DIRS:
    if SAMPLES.is_dir():
        try:
            from docx import Document
        except ImportError:
            Document = None
        if Document is not None:
            for sample in sorted(SAMPLES.glob("*.docx")):
                doc = Document(sample)
                found = []
                cd.check_front_matter_tables(doc, [None] * len(doc.paragraphs), found)
                cd.check_heading_indent_ladder(doc, [None] * len(doc.paragraphs), found)
                cd.check_cover_elements(doc, [None] * len(doc.paragraphs), found)
                cd.check_heading_styles(doc, [None] * len(doc.paragraphs), found)
                check(not found,
                      f"ตัวอย่างทางการ {sample.name!r} ถูกฟ้อง {len(found)} จุด "
                      f"({found[0]['issue'][:70] if found else ''}) — ตัวอย่างของหอสมุดคือเกณฑ์ "
                      f"ถ้ากฎฟ้องตัวอย่าง แปลว่ากฎผิด ไม่ใช่ตัวอย่างผิด")

# --- กฎใหม่ต้อง "จับได้จริง" ไม่ใช่แค่ไม่ฟ้องเทมเพลต -----------------------
# กฎที่เงียบทั้งกับไฟล์ถูกและไฟล์ผิดจะผ่านด่าน "เทมเพลตต้องได้ 0" ไปได้สบาย ๆ
# ทั้งที่ไม่ทำงานเลย — จึงต้องมีไฟล์ที่จงใจทำให้ผิดมายืนยันอีกด้านหนึ่ง
#
# สร้างจากเทมเพลตทางการแล้วแก้ XML สองจุด: (1) สลับสไตล์หัวข้อใหญ่ของบทหลัง
# ให้เป็นของบทที่ 1 (2) ตั้งย่อหน้าเนื้อความหนึ่งย่อหน้าเป็น justify
_FIXTURE = (Path(__file__).resolve().parent.parent / "fixtures"
            / "TULIBS_Thesis-template-Thai_rev_2024.docx")
if _FIXTURE.exists():
    try:
        from docx import Document as _Document
    except ImportError:
        _Document = None
    if _Document is not None:
        import re as _re
        import tempfile
        import zipfile

        with zipfile.ZipFile(_FIXTURE) as _zin:
            _xml = _zin.read("word/document.xml").decode("utf-8")
            _broken = _xml.replace('w:val="TUMainHeadingChapter3"',
                                   'w:val="TUMainHeadingChapter1"', 1)
            _m = _re.search(r'<w:pPr><w:pStyle w:val="TUParagraphNormal"/>', _broken)
            if _m:
                _broken = _broken[:_m.end()] + '<w:jc w:val="both"/>' + _broken[_m.end():]
            _tmp = Path(tempfile.gettempdir()) / "tu_rules_musfire.docx"
            with zipfile.ZipFile(_tmp, "w", zipfile.ZIP_DEFLATED) as _zout:
                for _item in _zin.infolist():
                    _zout.writestr(_item,
                                   _broken.encode("utf-8")
                                   if _item.filename == "word/document.xml"
                                   else _zin.read(_item.filename))

        _doc = _Document(_tmp)
        _pages = [None] * len(_doc.paragraphs)
        _hits = []
        cd.check_main_heading_chapter_match(_doc, _pages, _hits)
        check(any("ข้ามบท" in f["issue"] or "สไตล์ของบทที่" in f["issue"] for f in _hits),
              "must-fire: สไตล์หัวข้อใหญ่ข้ามบทไม่ถูกจับ — check_main_heading_chapter_match "
              "อ่านลำดับสไตล์ผิดหรือถูกปิดไป")
        _hits = []
        cd.check_body_justified(_doc, _pages, _hits)
        check(any("justify" in f["issue"] for f in _hits),
              "must-fire: ย่อหน้า justify ไม่ถูกจับ — check_body_justified ไม่ทำงาน")

if FAIL:
    print(f"FAILED ({len(FAIL)}):")
    for f in FAIL:
        print("  -", f)
    sys.exit(1)
print("all false-positive regression tests passed")
